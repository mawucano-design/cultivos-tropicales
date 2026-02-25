# app.py - Versión con visualización NDVI+NDRE en lugar de RGB
# CORREGIDO: YOLO sin OpenCV, DEM real SRTM 30m con OpenTopography, mapas Folium interactivos
# FIX: Separación de dependencias y manejo robusto de curvas de nivel
# AÑADIDO: Fuente alternativa Open Topo Data API (sin API Key)
# MEJORADO: Visualización de curvas de nivel y mapa de pendientes (imshow)
# MODIFICADO: Reemplazo de Qwen por DeepSeek para análisis con IA
import streamlit as st
import geopandas as gpd
import pandas as pd
import numpy as np
import tempfile
import os
import zipfile
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
from matplotlib.tri import Triangulation
import matplotlib.patches as mpatches
from matplotlib.colors import LinearSegmentedColormap
from mpl_toolkits.mplot3d import Axes3D
import io
from shapely.geometry import Polygon, LineString, Point
from shapely.geometry import mapping
import math
import warnings
import xml.etree.ElementTree as ET
import json
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import geojson
import requests
import contextily as ctx
# ===== IMPORTACIÓN DE MÓDULOS IA (AHORA CON DEEPSEEK) =====
from modules.ia_integration import (
    preparar_resumen_zonas,
    generar_analisis_fertilidad,
    generar_analisis_riesgo_hidrico,
    generar_recomendaciones_integradas
)

# ===== SOLUCIÓN PARA ERROR libGL.so.1 =====
# Configurar matplotlib para usar backend no interactivo
import matplotlib
matplotlib.use('Agg')  # Usar backend no interactivo

# Configurar variables de entorno para evitar problemas con OpenGL
os.environ['OPENCV_IO_ENABLE_OPENEXR'] = '1'
os.environ['QT_QPA_PLATFORM'] = 'offscreen'

# ===== DEPENDENCIAS OPCIONALES: SEPARADAS PARA MEJOR CONTROL =====
FOLIUM_OK = False
RASTERIO_OK = False
SKIMAGE_OK = False
try:
    import folium
    from folium.plugins import Fullscreen
    from branca.colormap import LinearColormap
    FOLIUM_OK = True
except ImportError:
    st.warning("⚠️ Folium no instalado. Los mapas interactivos no estarán disponibles.")

try:
    import rasterio
    from rasterio.mask import mask
    RASTERIO_OK = True
except ImportError:
    st.warning("⚠️ Rasterio no instalado. No se podrá descargar DEM real, se usará DEM sintético.")

try:
    from skimage import measure
    SKIMAGE_OK = True
except ImportError:
    st.warning("⚠️ scikit-image no instalado. No se generarán curvas de nivel.")

# Variable que indica si se pueden generar curvas (necesita skimage)
CURVAS_OK = SKIMAGE_OK

try:
    from streamlit_folium import folium_static
    FOLIUM_STATIC_OK = True
except ImportError:
    FOLIUM_STATIC_OK = False

# ===== CONFIGURACIÓN DE IA (DEEPSEEK) =====
DEEPSEEK_API_KEY = st.secrets.get("DEEPSEEK_API_KEY", os.getenv("DEEPSEEK_API_KEY"))
if not DEEPSEEK_API_KEY:
    st.warning("⚠️ No se encontró API Key de DeepSeek. La IA no estará disponible.")
else:
    # Establecer variable de entorno para que el módulo la tome
    os.environ["DEEPSEEK_API_KEY"] = DEEPSEEK_API_KEY

# ===== IMPORTACIONES GOOGLE EARTH ENGINE (NO MODIFICAR) =====
try:
    import ee
    GEE_AVAILABLE = True
except ImportError:
    GEE_AVAILABLE = False
    st.warning("⚠️ Google Earth Engine no está instalado. Para usar datos satelitales reales, instala con: pip install earthengine-api")

warnings.filterwarnings('ignore')

# === INICIALIZACIÓN SEGURA DE GOOGLE EARTH ENGINE (NO MODIFICAR) ===
def inicializar_gee():
    """Inicializa GEE con Service Account desde secrets de Streamlit Cloud"""
    if not GEE_AVAILABLE:
        return False
    
    try:
        # Intentar con Service Account desde secrets (Streamlit Cloud)
        gee_secret = os.environ.get('GEE_SERVICE_ACCOUNT')
        if gee_secret:
            try:
                credentials_info = json.loads(gee_secret.strip())
                credentials = ee.ServiceAccountCredentials(
                    credentials_info['client_email'],
                    key_data=json.dumps(credentials_info)
                )
                ee.Initialize(credentials, project='ee-mawucano25')
                st.session_state.gee_authenticated = True
                st.session_state.gee_project = 'ee-mawucano25'
                print("✅ GEE inicializado con Service Account")
                return True
            except Exception as e:
                print(f"⚠️ Error con Service Account: {str(e)}")
        
        # Fallback: autenticación local (desarrollo en tu Linux)
        try:
            ee.Initialize(project='ee-mawucano25')
            st.session_state.gee_authenticated = True
            st.session_state.gee_project = 'ee-mawucano25'
            print("✅ GEE inicializado localmente")
            return True
        except Exception as e:
            print(f"⚠️ Error inicialización local: {str(e)}")
            
        st.session_state.gee_authenticated = False
        return False
        
    except Exception as e:
        st.session_state.gee_authenticated = False
        print(f"❌ Error crítico GEE: {str(e)}")
        return False

# Ejecutar inicialización al inicio (ANTES de cualquier uso de ee.*)
if 'gee_authenticated' not in st.session_state:
    st.session_state.gee_authenticated = False
    st.session_state.gee_project = ''
    if GEE_AVAILABLE:
        inicializar_gee()

# ===== FUNCIONES YOLO PARA DETECCIÓN DE PLAGAS/ENFERMEDADES (VERSIÓN PIL - SIN OpenCV) =====
def cargar_modelo_yolo(modelo_path='yolo_plagas.pt'):
    """
    Carga el modelo YOLO para detección.
    Si falla la carga real, devuelve un modelo de demostración.
    """
    try:
        # Intentar importar ultralytics
        try:
            from ultralytics import YOLO
        except ImportError:
            st.warning("⚠️ Ultralytics no instalado. Usando simulador YOLO.")
            return _crear_modelo_demo()

        # Intentar cargar modelo real
        try:
            if os.path.exists(modelo_path):
                modelo = YOLO(modelo_path)
                st.success(f"✅ Modelo YOLO personalizado cargado: {modelo_path}")
                return modelo
            else:
                modelo = YOLO('yolov8n.pt')
                st.info("ℹ️ Usando modelo YOLO de demostración (yolov8n.pt)")
                return modelo
        except Exception as e:
            st.warning(f"⚠️ No se pudo cargar modelo YOLO real: {e}. Usando simulador.")
            return _crear_modelo_demo()

    except Exception as e:
        st.error(f"❌ Error crítico en YOLO: {str(e)}")
        return _crear_modelo_demo()

def _crear_modelo_demo():
    """Crea un modelo de demostración para simular detecciones."""
    class ModeloDemo:
        def __init__(self):
            self.names = {
                0: 'Plaga_Gusano',
                1: 'Enfermedad_Roya',
                2: 'Deficiencia_Nutricional',
                3: 'Plaga_Pulgón',
                4: 'Enfermedad_Oídio'
            }
        def __call__(self, img, conf=0.5):
            # Solo para compatibilidad, no se usa en predicción real
            return None
    return ModeloDemo()

def detectar_plagas_yolo(imagen_path, modelo, confianza_minima=0.5):
    """
    Ejecuta detección de plagas/enfermedades con YOLO.
    - Utiliza PIL para cargar y dibujar (sin OpenCV).
    - Si el modelo es de demostración, genera detecciones aleatorias realistas.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
        import numpy as np

        # --- 1. Cargar imagen de forma segura ---
        if isinstance(imagen_path, BytesIO):
            imagen_path.seek(0)
            img_pil = Image.open(imagen_path).convert('RGB')
        elif isinstance(imagen_path, str):
            img_pil = Image.open(imagen_path).convert('RGB')
        else:
            # Si es otro tipo de buffer (ej. UploadedFile)
            img_pil = Image.open(imagen_path).convert('RGB')

        img_np = np.array(img_pil)  # Para el modelo YOLO

        # --- 2. Determinar si es modelo demo o real ---
        es_demo = hasattr(modelo, 'names') and not hasattr(modelo, 'predict')

        # --- 3. Obtener detecciones ---
        detecciones = []

        if es_demo:
            # ========== MODELO DE DEMOSTRACIÓN ==========
            altura, ancho = img_np.shape[:2]
            np.random.seed(int(datetime.now().timestamp()))
            n_detecciones = np.random.randint(3, 8)

            for _ in range(n_detecciones):
                x1 = np.random.randint(0, ancho - 100)
                y1 = np.random.randint(0, altura - 100)
                ancho_bbox = np.random.randint(50, 200)
                alto_bbox = np.random.randint(50, 200)
                x2 = min(x1 + ancho_bbox, ancho)
                y2 = min(y1 + alto_bbox, altura)

                clase_id = np.random.choice(list(modelo.names.keys()))
                confianza = np.random.uniform(confianza_minima, 0.95)

                detecciones.append({
                    'clase': modelo.names[clase_id],
                    'confianza': confianza,
                    'bbox': [x1, y1, x2, y2],
                    'area': (x2 - x1) * (y2 - y1)
                })
        else:
            # ========== MODELO REAL DE YOLO ==========
            from ultralytics import YOLO

            # Ejecutar predicción
            resultados = modelo(img_np, conf=confianza_minima)

            for r in resultados:
                if r.boxes is not None:
                    boxes = r.boxes.xyxy.cpu().numpy()
                    confs = r.boxes.conf.cpu().numpy()
                    cls_ids = r.boxes.cls.cpu().numpy().astype(int)

                    for box, conf, cls_id in zip(boxes, confs, cls_ids):
                        x1, y1, x2, y2 = map(int, box)
                        nombre_clase = modelo.names[cls_id]
                        detecciones.append({
                            'clase': nombre_clase,
                            'confianza': float(conf),
                            'bbox': [x1, y1, x2, y2],
                            'area': (x2 - x1) * (y2 - y1)
                        })

        # --- 4. Dibujar bounding boxes con PIL ---
        draw = ImageDraw.Draw(img_pil)
        try:
            font = ImageFont.truetype("arial.ttf", 14)
        except:
            font = ImageFont.load_default()

        for det in detecciones:
            x1, y1, x2, y2 = det['bbox']
            conf = det['confianza']
            label = f"{det['clase']} {conf:.2f}"

            # Color según tipo (demo o real)
            if es_demo:
                color = 'green' if 'Plaga' in det['clase'] else 'red'
            else:
                color = 'lime'

            draw.rectangle([x1, y1, x2, y2], outline=color, width=3)
            draw.text((x1, y1 - 16), label, fill=color, font=font)

        return detecciones, np.array(img_pil)

    except Exception as e:
        st.error(f"❌ Error en detección YOLO: {str(e)}")
        return [], None

def analizar_imagen_dron(gdf, fecha_analisis):
    """
    Simula una imagen de dron usando PIL (sin OpenCV).
    """
    try:
        from PIL import Image, ImageDraw
        import numpy as np

        ancho, altura = 1000, 800
        img_pil = Image.new('RGB', (ancho, altura), color=(100, 150, 100))
        draw = ImageDraw.Draw(img_pil)

        np.random.seed(int(fecha_analisis.timestamp()))
        num_anomalias = np.random.randint(5, 15)

        for _ in range(num_anomalias):
            x = np.random.randint(0, ancho)
            y = np.random.randint(0, altura)
            radio = np.random.randint(20, 60)

            tipo = np.random.choice(['plaga', 'enfermedad', 'deficiencia'])
            if tipo == 'plaga':
                color = (255, 0, 0)      # Rojo
            elif tipo == 'enfermedad':
                color = (0, 0, 255)      # Azul
            else:
                color = (255, 255, 0)    # Amarillo

            draw.ellipse([x-radio, y-radio, x+radio, y+radio], fill=color)

        imagen_bytes = BytesIO()
        img_pil.save(imagen_bytes, format='JPEG')
        imagen_bytes.seek(0)
        return imagen_bytes

    except Exception as e:
        st.error(f"❌ Error generando imagen de dron: {str(e)}")
        return None

def generar_reporte_plagas(detecciones, cultivo):
    """
    Genera un reporte detallado de las plagas/enfermedades detectadas.
    """
    try:
        if not detecciones:
            return "✅ No se detectaron plagas/enfermedades significativas."

        conteo_plagas = {}
        areas_plagas = {}

        for det in detecciones:
            clase = det['clase']
            conteo_plagas[clase] = conteo_plagas.get(clase, 0) + 1
            areas_plagas[clase] = areas_plagas.get(clase, 0) + det['area']

        reporte = f"## 🦠 REPORTE DE PLAGAS/ENFERMEDADES - {cultivo}\n\n"
        reporte += f"**Total de detecciones:** {len(detecciones)}\n\n"
        reporte += "**Distribución por tipo:**\n"
        for clase, conteo in conteo_plagas.items():
            porcentaje = (conteo / len(detecciones)) * 100
            area_prom = areas_plagas[clase] / conteo
            reporte += f"- **{clase}**: {conteo} detecciones ({porcentaje:.1f}%), área promedio: {area_prom:.0f} px²\n"

        # Recomendaciones específicas por cultivo
        reporte += "\n**🧪 RECOMENDACIONES ESPECÍFICAS:**\n"
        if cultivo in ['TRIGO', 'MAIZ', 'SORGO']:
            if any('roya' in clase.lower() for clase in conteo_plagas.keys()):
                reporte += "- **Fungicida**: Aplicar Triazol (0.5-1.0 L/ha) cada 15 días\n"
            if any('gusano' in clase.lower() for clase in conteo_plagas.keys()):
                reporte += "- **Insecticida**: Lambda-cialotrina (0.2-0.3 L/ha) en detección temprana\n"
            if any('pulgón' in clase.lower() for clase in conteo_plagas.keys()):
                reporte += "- **Control biológico**: Liberar Aphidius colemani (parásito de pulgones)\n"
        elif cultivo in ['VID', 'OLIVO', 'ALMENDRO']:
            if any('oidio' in clase.lower() for clase in conteo_plagas.keys()):
                reporte += "- **Azufre micronizado**: 3-5 kg/ha aplicado preventivamente\n"
            if any('mosca' in clase.lower() for clase in conteo_plagas.keys()):
                reporte += "- **Trampas amarillas**: 50-100 trampas/ha + Spinosad (0.3 L/ha)\n"
            if any('polilla' in clase.lower() for clase in conteo_plagas.keys()):
                reporte += "- **Confusión sexual**: Difusores de feromonas (500-1000 unidades/ha)\n"

        # Recomendaciones generales
        reporte += "\n**📋 RECOMENDACIONES GENERALES:**\n"
        reporte += "- **Monitoreo**: Revisar cada 7-10 días durante períodos críticos\n"
        reporte += "- **Umbrales**: Actuar cuando >5% de plantas muestren síntomas\n"
        reporte += "- **Rotación**: Alternar modos de acción para evitar resistencias\n"
        reporte += "- **Registro**: Documentar todas las aplicaciones y resultados\n"

        # Acciones inmediatas según severidad
        reporte += "\n**⚠️ ACCIONES INMEDIATAS:**\n"
        if len(detecciones) > 20:
            reporte += "- **ALERTA ROJA**: Incidencia crítica. Aplicación urgente requerida\n"
            reporte += "- **Contactar**: Asesor técnico para plan de emergencia\n"
        elif len(detecciones) > 10:
            reporte += "- **ALERTA AMARILLA**: Incidencia media. Aplicar en 48 horas\n"
            reporte += "- **Aumentar**: Frecuencia de monitoreo a cada 5 días\n"
        else:
            reporte += "- **SITUACIÓN CONTROLADA**: Continuar monitoreo rutinario\n"
            reporte += "- **Preventivo**: Aplicar tratamiento preventivo en próximos 15 días\n"

        # Métricas de gravedad
        reporte += "\n**📊 MÉTRICAS DE GRAVEDAD:**\n"
        severidad_total = sum(d['area'] * d['confianza'] for d in detecciones)
        reporte += f"- Índice de Severidad: {severidad_total:.0f}\n"
        reporte += f"- Área Total Afectada: {sum(d['area'] for d in detecciones):.0f} px²\n"

        return reporte

    except Exception as e:
        return f"❌ Error generando reporte: {str(e)}"

# ===== NUEVAS FUNCIONES PARA MAPAS DE POTENCIAL DE COSECHA =====
def crear_mapa_potencial_cosecha(gdf_completo, cultivo):
    """Crear mapa de potencial de cosecha"""
    try:
        gdf_plot = gdf_completo.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        
        # Usar colores para potencial (verde = alto, rojo = bajo)
        cmap = LinearSegmentedColormap.from_list('potencial', ['#ff4444', '#ffff44', '#44ff44'])
        
        # Obtener valores de potencial
        potenciales = gdf_plot['proy_rendimiento_sin_fert']
        vmin, vmax = potenciales.min(), potenciales.max()
        
        for idx, row in gdf_plot.iterrows():
            valor = row['proy_rendimiento_sin_fert']
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            color = cmap(valor_norm)
            
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='black', linewidth=1.5, alpha=0.7)
            
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.0f}kg", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='black', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9))
        
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.7)
        except:
            pass
        
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} POTENCIAL DE COSECHA - {cultivo}',
                     fontsize=16, fontweight='bold', pad=20)
        ax.set_xlabel('Longitud')
        ax.set_ylabel('Latitud')
        ax.grid(True, alpha=0.3)
        
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label('Rendimiento Potencial (kg/ha)', fontsize=12, fontweight='bold')
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando mapa de potencial de cosecha: {str(e)}")
        return None

def crear_mapa_potencial_con_recomendaciones(gdf_completo, cultivo):
    """Crear mapa de potencial de cosecha con recomendaciones aplicadas"""
    try:
        gdf_plot = gdf_completo.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        
        # Colores para potencial mejorado
        cmap = LinearSegmentedColormap.from_list('potencial_mejorado', ['#ffaa44', '#ffff44', '#44ff44', '#00aa00'])
        
        # Obtener valores de potencial con recomendaciones
        potenciales = gdf_plot['proy_rendimiento_con_fert']
        incrementos = gdf_plot['proy_incremento_esperado']
        vmin, vmax = potenciales.min(), potenciales.max()
        
        for idx, row in gdf_plot.iterrows():
            valor = row['proy_rendimiento_con_fert']
            incremento = row['proy_incremento_esperado']
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            color = cmap(valor_norm)
            
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='black', linewidth=1.5, alpha=0.7)
            
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.0f}kg\n+{incremento:.1f}%", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=7, color='black', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9))
        
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.7)
        except:
            pass
        
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} POTENCIAL CON RECOMENDACIONES - {cultivo}',
                     fontsize=16, fontweight='bold', pad=20)
        ax.set_xlabel('Longitud')
        ax.set_ylabel('Latitud')
        ax.grid(True, alpha=0.3)
        
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label('Rendimiento Mejorado (kg/ha)', fontsize=12, fontweight='bold')
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando mapa de potencial con recomendaciones: {str(e)}")
        return None

def crear_grafico_comparativo_potencial(gdf_completo, cultivo):
    """Crear gráfico comparativo de potencial vs potencial con recomendaciones"""
    try:
        fig, ax = plt.subplots(1, 1, figsize=(14, 7))
        
        zonas = gdf_completo['id_zona'].astype(str).tolist()
        sin_fert = gdf_completo['proy_rendimiento_sin_fert'].tolist()
        con_fert = gdf_completo['proy_rendimiento_con_fert'].tolist()
        incrementos = gdf_completo['proy_incremento_esperado'].tolist()
        
        x = np.arange(len(zonas))
        width = 0.35
        
        bars1 = ax.bar(x - width/2, sin_fert, width, label='Sin Fertilización', color='#ff9999', alpha=0.8)
        bars2 = ax.bar(x + width/2, con_fert, width, label='Con Fertilización', color='#66b3ff', alpha=0.8)
        
        # Agregar línea de incremento porcentual
        ax2 = ax.twinx()
        ax2.plot(x, incrementos, 'g-', marker='o', linewidth=2, markersize=6, label='Incremento %')
        ax2.set_ylabel('Incremento (%)', color='green', fontsize=12)
        ax2.tick_params(axis='y', labelcolor='green')
        ax2.set_ylim(0, max(incrementos) * 1.2)
        
        ax.set_xlabel('Zona', fontsize=12)
        ax.set_ylabel('Rendimiento (kg/ha)', fontsize=12)
        ax.set_title(f'COMPARATIVO DE POTENCIAL DE COSECHA - {cultivo}', fontsize=14, fontweight='bold', pad=20)
        ax.set_xticks(x)
        ax.set_xticklabels(zonas, rotation=45)
        ax.legend(loc='upper left')
        ax2.legend(loc='upper right')
        
        # Agregar valores en las barras
        for bar1, bar2 in zip(bars1, bars2):
            height1 = bar1.get_height()
            height2 = bar2.get_height()
            ax.text(bar1.get_x() + bar1.get_width()/2., height1 + max(sin_fert)*0.01,
                   f'{height1:.0f}', ha='center', va='bottom', fontsize=8, rotation=90)
            ax.text(bar2.get_x() + bar2.get_width()/2., height2 + max(con_fert)*0.01,
                   f'{height2:.0f}', ha='center', va='bottom', fontsize=8, rotation=90)
        
        # Estadísticas
        stats_text = f"""
        Estadísticas:
        • Rendimiento promedio sin fertilización: {np.mean(sin_fert):.0f} kg/ha
        • Rendimiento promedio con fertilización: {np.mean(con_fert):.0f} kg/ha
        • Incremento promedio: {np.mean(incrementos):.1f}%
        • Máximo incremento: {max(incrementos):.1f}% (Zona {zonas[incrementos.index(max(incrementos))]})
        """
        
        ax.text(0.02, 0.98, stats_text, transform=ax.transAxes, fontsize=10,
                verticalalignment='top',
                bbox=dict(boxstyle="round,pad=0.5", facecolor='lightyellow', alpha=0.9))
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando gráfico comparativo: {str(e)}")
        return None

# ===== NUEVA FUNCIÓN: VISUALIZACIÓN NDVI + NDRE GEE (INTERACTIVA) =====
def visualizar_indices_gee(gdf, satelite, fecha_inicio, fecha_fin):
    """Genera visualización NDVI + NDRE interactiva con iframes"""
    if not GEE_AVAILABLE or not st.session_state.gee_authenticated:
        return None, "❌ Google Earth Engine no está autenticado"
    
    try:
        # Obtener bounding box de la parcela
        bounds = gdf.total_bounds
        min_lon, min_lat, max_lon, max_lat = bounds
        
        # Expandir ligeramente el área para asegurar cobertura
        min_lon -= 0.001
        max_lon += 0.001
        min_lat -= 0.001
        max_lat += 0.001
        
        # Crear geometría
        geometry = ee.Geometry.Rectangle([min_lon, min_lat, max_lon, max_lat])
        
        # Formatear fechas
        start_date = fecha_inicio.strftime('%Y-%m-%d')
        end_date = fecha_fin.strftime('%Y-%m-%d')
        
        # Seleccionar colección según satélite
        if satelite == 'SENTINEL-2_GEE':
            collection = ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED')
            ndvi_bands = ['B8', 'B4']
            ndre_bands = ['B8', 'B5']
            title = "Sentinel-2 NDVI + NDRE"
            
        elif satelite == 'LANDSAT-8_GEE':
            collection = ee.ImageCollection('LANDSAT/LC08/C02/T1_L2')
            ndvi_bands = ['SR_B5', 'SR_B4']
            ndre_bands = ['SR_B5', 'SR_B6']
            title = "Landsat 8 NDVI + NDRE"
            
        elif satelite == 'LANDSAT-9_GEE':
            collection = ee.ImageCollection('LANDSAT/LC09/C02/T1_L2')
            ndvi_bands = ['SR_B5', 'SR_B4']
            ndre_bands = ['SR_B5', 'SR_B6']
            title = "Landsat 9 NDVI + NDRE"
            
        else:
            return None, "⚠️ Satélite no soportado para visualización de índices"
        
        # Filtrar colección
        try:
            filtered = (collection
                       .filterBounds(geometry)
                       .filterDate(start_date, end_date)
                       .filter(ee.Filter.lt('CLOUDY_PIXEL_PERCENTAGE', 60)))
            
            # Verificar si hay imágenes
            count = filtered.size().getInfo()
            if count == 0:
                return None, f"⚠️ No hay imágenes disponibles para {start_date} - {end_date}"
            
            # Tomar la imagen con menos nubes
            image = filtered.sort('CLOUDY_PIXEL_PERCENTAGE').first()
            
            if image is None:
                return None, "❌ Error: La imagen obtenida es nula"
            
            # Calcular NDVI
            ndvi = image.normalizedDifference(ndvi_bands).rename('NDVI')
            
            # Calcular NDRE
            ndre = image.normalizedDifference(ndre_bands).rename('NDRE')
            
            # Obtener información de la imagen
            image_id = image.get('system:index').getInfo()
            
            cloud_percent_ee = image.get('CLOUDY_PIXEL_PERCENTAGE')
            cloud_percent = cloud_percent_ee.getInfo() if cloud_percent_ee else 0
            
            fecha_imagen_ee = image.get('system:time_start')
            fecha_imagen = fecha_imagen_ee.getInfo() if fecha_imagen_ee else None
            
            if fecha_imagen:
                fecha_str = datetime.fromtimestamp(fecha_imagen / 1000).strftime('%Y-%m-%d')
                title += f" - {fecha_str}"
            
            # Parámetros de visualización
            ndvi_vis_params = {
                'min': -0.2,
                'max': 0.8,
                'palette': ['red', 'yellow', 'green']
            }
            
            ndre_vis_params = {
                'min': -0.1,
                'max': 0.6,
                'palette': ['blue', 'white', 'green']
            }
            
            # Generar URLs de los mapas
            ndvi_map_id_dict = ndvi.getMapId(ndvi_vis_params)
            ndre_map_id_dict = ndre.getMapId(ndre_vis_params)
            
            if not ndvi_map_id_dict or 'mapid' not in ndvi_map_id_dict:
                return None, "❌ Error generando mapa NDVI"
            
            if not ndre_map_id_dict or 'mapid' not in ndre_map_id_dict:
                return None, "❌ Error generando mapa NDRE"
            
            # Usar URLs de tiles de Earth Engine
            ndvi_mapid = ndvi_map_id_dict['mapid']
            ndre_mapid = ndre_map_id_dict['mapid']
            
            # Si hay token, agregarlo como parámetro
            ndvi_token = ndvi_map_id_dict.get('token', '')
            ndre_token = ndre_map_id_dict.get('token', '')
            
            ndvi_token_param = f"?token={ndvi_token}" if ndvi_token else ""
            ndre_token_param = f"?token={ndre_token}" if ndre_token else ""
            
            # Crear HTML con iframes
            html = f"""
            <div style="display: flex; flex-wrap: wrap; gap: 20px; margin-bottom: 20px;">
                <div style="flex: 1; min-width: 300px; border: 2px solid #3b82f6; border-radius: 10px; overflow: hidden;">
                    <h4 style="text-align: center; background: linear-gradient(135deg, #ff4444, #ffff44, #44ff44); color: #000; padding: 10px; margin: 0;">🌱 MAPA NDVI</h4>
                    <iframe
                        width="100%"
                        height="400"
                        src="https://earthengine.googleapis.com/v1alpha/{ndvi_mapid}/tiles/{{z}}/{{x}}/{{y}}{ndvi_token_param}"
                        frameborder="0"
                        allowfullscreen
                        style="display: block;"
                    ></iframe>
                    <div style="background: #f0f9ff; padding: 8px; border-top: 1px solid #3b82f6;">
                        <p style="margin: 5px 0; font-size: 0.8em;">
                            <strong>Escala:</strong> -0.2 (rojo) a 0.8 (verde)
                        </p>
                    </div>
                </div>
                
                <div style="flex: 1; min-width: 300px; border: 2px solid #10b981; border-radius: 10px; overflow: hidden;">
                    <h4 style="text-align: center; background: linear-gradient(135deg, #0000ff, #ffffff, #00ff00); color: #000; padding: 10px; margin: 0;">🌿 MAPA NDRE</h4>
                    <iframe
                        width="100%"
                        height="400"
                        src="https://earthengine.googleapis.com/v1alpha/{ndre_mapid}/tiles/{{z}}/{{x}}/{{y}}{ndre_token_param}"
                        frameborder="0"
                        allowfullscreen
                        style="display: block;"
                    ></iframe>
                    <div style="background: #f0f9ff; padding: 8px; border-top: 1px solid #10b981;">
                        <p style="margin: 5px 0; font-size: 0.8em;">
                            <strong>Escala:</strong> -0.1 (azul) a 0.6 (verde)
                        </p>
                    </div>
                </div>
            </div>
            
            <div style="background: #f8fafc; padding: 15px; border-radius: 8px; margin-top: 15px; border: 1px solid #e2e8f0;">
                <h4 style="margin-top: 0; color: #3b82f6;">📊 INFORMACIÓN DE LOS ÍNDICES</h4>
                <div style="display: flex; flex-wrap: wrap; gap: 20px;">
                    <div style="flex: 1; min-width: 200px;">
                        <h5 style="color: #3b82f6; margin-bottom: 8px;">🌱 NDVI (Índice de Vegetación de Diferencia Normalizada)</h5>
                        <ul style="margin: 0; padding-left: 20px; font-size: 0.9em;">
                            <li><strong>Rango saludable:</strong> 0.3 - 0.8</li>
                            <li><strong>Valores bajos (&lt;0.2):</strong> Suelo desnudo, estrés hídrico</li>
                            <li><strong>Valores medios (0.3-0.5):</strong> Vegetación moderada</li>
                            <li><strong>Valores altos (&gt;0.6):</strong> Vegetación densa y saludable</li>
                        </ul>
                    </div>
                    
                    <div style="flex: 1; min-width: 200px;">
                        <h5 style="color: #10b981; margin-bottom: 8px;">🌿 NDRE (Índice de Borde Rojo Normalizado)</h5>
                        <ul style="margin: 0; padding-left: 20px; font-size:0.9em;">
                            <li><strong>Rango saludable:</strong> 0.2 - 0.5</li>
                            <li><strong>Sensibilidad:</strong> Clorofila en capas internas</li>
                            <li><strong>Uso:</strong> Monitoreo de nitrógeno</li>
                            <li><strong>Ventaja:</strong> Menos saturación en vegetación densa</li>
                        </ul>
                    </div>
                </div>
                
                <div style="margin-top: 15px; padding: 10px; background: #e0f2fe; border-radius: 5px; border-left: 4px solid #3b82f6;">
                    <p style="margin: 0; font-size: 0.85em;">
                        <strong>ℹ️ Información técnica:</strong> {title} | Nubes: {cloud_percent}% | ID: {image_id} | 
                        <strong>Interpretación:</strong> Compara ambos índices para detectar estrés temprano
                    </p>
                </div>
            </div>
            """
            
            return html, f"✅ {title}"
            
        except Exception as e:
            error_msg = str(e)
            if "Parameter 'object' is required" in error_msg:
                return None, f"❌ No se encontró imagen para el período {start_date} - {end_date}"
            else:
                return None, f"❌ Error GEE: {error_msg}"
        
    except Exception as e:
        return None, f"❌ Error general: {str(e)}"


# ===== MODIFICACIÓN DE LA FUNCIÓN visualizar_indices_gee_estatico =====
def visualizar_indices_gee_estatico(gdf, satelite, fecha_inicio, fecha_fin):
    """Versión mejorada que devuelve las imágenes en bytes para descarga"""
    if not GEE_AVAILABLE or not st.session_state.gee_authenticated:
        return None, "❌ Google Earth Engine no está autenticado"
    
    try:
        # Obtener bounding box de la parcela
        bounds = gdf.total_bounds
        min_lon, min_lat, max_lon, max_lat = bounds
        
        # Crear geometría
        geometry = ee.Geometry.Rectangle([min_lon, min_lat, max_lon, max_lat])
        
        # Formatear fechas
        start_date = fecha_inicio.strftime('%Y-%m-%d')
        end_date = fecha_fin.strftime('%Y-%m-%d')
        
        # Seleccionar colección según satélite
        if satelite == 'SENTINEL-2_GEE':
            collection = ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED')
            ndvi_bands = ['B8', 'B4']
            ndre_bands = ['B8', 'B5']
            title = "Sentinel-2"
            
        elif satelite == 'LANDSAT-8_GEE':
            collection = ee.ImageCollection('LANDSAT/LC08/C02/T1_L2')
            ndvi_bands = ['SR_B5', 'SR_B4']
            ndre_bands = ['SR_B5', 'SR_B6']
            title = "Landsat 8"
            
        elif satelite == 'LANDSAT-9_GEE':
            collection = ee.ImageCollection('LANDSAT/LC09/C02/T1_L2')
            ndvi_bands = ['SR_B5', 'SR_B4']
            ndre_bands = ['SR_B5', 'SR_B6']
            title = "Landsat 9"
            
        else:
            return None, "⚠️ Satélite no soportado"
        
        # Filtrar colección
        filtered = (collection
                   .filterBounds(geometry)
                   .filterDate(start_date, end_date)
                   .filter(ee.Filter.lt('CLOUDY_PIXEL_PERCENTAGE', 60)))
        
        # Verificar si hay imágenes
        count = filtered.size().getInfo()
        if count == 0:
            return None, f"⚠️ No hay imágenes disponibles para el período {start_date} - {end_date}"
        
        # Tomar la imagen con menos nubes
        image = filtered.sort('CLOUDY_PIXEL_PERCENTAGE').first()
        
        # Calcular índices
        ndvi = image.normalizedDifference(ndvi_bands).rename('NDVI')
        ndre = image.normalizedDifference(ndre_bands).rename('NDRE')
        
        # Generar URLs de miniaturas (thumbnails) estáticas
        try:
            # Parámetros comunes
            region_params = {
                'dimensions': 800,
                'region': geometry,
                'format': 'png'
            }
            
            # Configuración específica para cada índice
            ndvi_thumbnail_url = ndvi.getThumbURL({
                'min': -0.2,
                'max': 0.8,
                'palette': ['red', 'yellow', 'green'],
                **region_params
            })
            
            ndre_thumbnail_url = ndre.getThumbURL({
                'min': -0.1,
                'max': 0.6,
                'palette': ['blue', 'white', 'green'],
                **region_params
            })
            
            # Descargar las imágenes
            import requests
            
            ndvi_response = requests.get(ndvi_thumbnail_url)
            ndre_response = requests.get(ndre_thumbnail_url)
            
            if ndvi_response.status_code != 200 or ndre_response.status_code != 200:
                return None, f"❌ Error descargando imágenes: {ndvi_response.status_code}, {ndre_response.status_code}"
            
            # Convertir a bytes
            ndvi_bytes = BytesIO(ndvi_response.content)
            ndre_bytes = BytesIO(ndre_response.content)
            
            return {
                'ndvi_bytes': ndvi_bytes,
                'ndre_bytes': ndre_bytes,
                'title': title,
                'image_date': image.get('system:time_start').getInfo() if image.get('system:time_start') else None,
                'cloud_percent': image.get('CLOUDY_PIXEL_PERCENTAGE').getInfo() if image.get('CLOUDY_PIXEL_PERCENTAGE') else 0,
                'image_id': image.get('system:index').getInfo() if image.get('system:index') else 'N/A'
            }, f"✅ {title} - Imágenes descargadas correctamente"
            
        except Exception as e:
            return None, f"❌ Error generando imágenes estáticas: {str(e)}"
        
    except Exception as e:
        return None, f"❌ Error: {str(e)}"

# ===== FUNCIONES MODIFICADAS PARA EXPORTACIÓN TIFF/GeoTIFF =====
import rasterio
from rasterio.transform import from_origin
from rasterio.crs import CRS
from PIL import Image

def exportar_mapa_tiff(buffer_png, gdf, nombre_base, cultivo):
    """Exporta un mapa PNG a formato TIFF/GeoTIFF con georreferenciación"""
    try:
        # Cargar la imagen PNG
        img = Image.open(buffer_png)
        
        # Obtener bounds de la parcela
        gdf_proj = gdf.to_crs(epsg=3857)  # Web Mercator para cálculos
        bounds = gdf_proj.total_bounds
        
        # Calcular transformación affine
        width, height = img.size
        transform = from_origin(bounds[0], bounds[3], 
                              (bounds[2] - bounds[0]) / width,
                              (bounds[3] - bounds[1]) / height)
        
        # Convertir imagen a array numpy
        img_array = np.array(img)
        
        # Si la imagen es RGBA, convertir a RGB
        if img_array.shape[2] == 4:
            img_array = img_array[:, :, :3]
        
        # Transponer para formato rasterio (bandas, altura, ancho)
        if len(img_array.shape) == 3:
            img_array = np.transpose(img_array, (2, 0, 1))
        else:
            img_array = np.expand_dims(img_array, axis=0)
        
        # Crear archivo TIFF en memoria
        tiff_buffer = BytesIO()
        
        with rasterio.open(
            tiff_buffer,
            'w',
            driver='GTiff',
            height=height,
            width=width,
            count=img_array.shape[0],
            dtype=img_array.dtype,
            crs=CRS.from_epsg(3857),
            transform=transform,
            compress='lzw'  # Compresión para reducir tamaño
        ) as dst:
            dst.write(img_array)
        
        tiff_buffer.seek(0)
        
        # Generar nombre de archivo
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"{nombre_base}_{cultivo}_{timestamp}.tiff"
        
        return tiff_buffer, nombre_archivo
        
    except Exception as e:
        st.error(f"❌ Error exportando a TIFF: {str(e)}")
        return None, None

def crear_boton_descarga_tiff(buffer_png, gdf, nombre_archivo, texto_boton="📥 Descargar TIFF", cultivo=""):
    """Crear botón de descarga para archivos TIFF/GeoTIFF"""
    if buffer_png and gdf is not None:
        # Exportar a TIFF
        tiff_buffer, tiff_filename = exportar_mapa_tiff(buffer_png, gdf, nombre_archivo, cultivo)
        
        if tiff_buffer:
            st.download_button(
                label=texto_boton,
                data=tiff_buffer,
                file_name=tiff_filename,
                mime="image/tiff"
            )
    else:
        st.warning("No hay datos para exportar")

# ===== FUNCIONES DE CURVAS DE NIVEL (MODIFICADAS) =====

def obtener_dem_opentopography(gdf, api_key=None):
    """
    Descarga DEM SRTM 1 arc-seg (30m) desde OpenTopography.
    Retorna (dem_array, meta, transform) o (None, None, None) si falla.
    Requiere rasterio.
    """
    if not RASTERIO_OK:
        st.warning("⚠️ Rasterio no instalado. No se puede descargar DEM real.")
        return None, None, None

    # 1. Obtener API Key (prioridad: argumento > variable entorno > secret)
    if api_key is None:
        api_key = os.environ.get("OPENTOPOGRAPHY_API_KEY", None)
    if not api_key:
        st.warning("⚠️ No se encontró API Key de OpenTopography. Se usará DEM sintético.")
        st.info("📌 Obtén una API Key gratuita en: https://opentopography.org/")
        return None, None, None

    try:
        # 2. Obtener bounding box y validar que esté dentro de la cobertura SRTM (latitudes entre -60 y 60)
        bounds = gdf.total_bounds
        west, south, east, north = bounds

        # Verificar límites
        if south < -60 or north > 60:
            st.warning("⚠️ El área está fuera de la cobertura de SRTM (latitudes > 60° o < -60°). Usando DEM sintético.")
            return None, None, None

        lon_span = east - west
        lat_span = north - south
        west = max(west - 0.05 * lon_span, -180)
        east = min(east + 0.05 * lon_span, 180)
        south = max(south - 0.05 * lat_span, -60)
        north = min(north + 0.05 * lat_span, 60)

        params = {
            "demtype": "SRTMGL1",
            "south": south,
            "north": north,
            "west": west,
            "east": east,
            "outputFormat": "GTiff",
            "API_Key": api_key
        }

        url = "https://portal.opentopography.org/API/globaldem"
        
        with st.spinner("🛰️ Descargando DEM desde OpenTopography..."):
            response = requests.get(url, params=params, timeout=60)
            
        if response.status_code == 403:
            st.error("❌ API Key inválida o no autorizada.")
            return None, None, None
        elif response.status_code == 404:
            st.error("❌ No se encontraron datos SRTM para esta área.")
            return None, None, None
        elif response.status_code != 200:
            st.error(f"❌ Error en OpenTopography: HTTP {response.status_code}")
            return None, None, None

        dem_bytes = BytesIO(response.content)
        with rasterio.open(dem_bytes) as src:
            geom = [mapping(gdf.unary_union)]
            out_image, out_transform = mask(src, geom, crop=True, nodata=-32768, all_touched=True)
            out_meta = src.meta.copy()
            out_meta.update({
                "driver": "GTiff",
                "height": out_image.shape[1],
                "width": out_image.shape[2],
                "transform": out_transform,
                "nodata": -32768
            })

        dem_array = out_image.squeeze()
        dem_array = np.ma.masked_where(dem_array <= -32768, dem_array)
        
        if dem_array.mask.all() if isinstance(dem_array, np.ma.MaskedArray) else np.all(dem_array <= -32768):
            st.warning("⚠️ El DEM descargado no contiene datos válidos dentro del polígono.")
            return None, None, None

        st.success("✅ DEM SRTM 30m descargado y recortado exitosamente.")
        return dem_array, out_meta, out_transform

    except requests.exceptions.Timeout:
        st.error("❌ Tiempo de espera agotado al conectar con OpenTopography.")
        return None, None, None
    except Exception as e:
        st.error(f"❌ Error inesperado al obtener DEM: {str(e)[:200]}")
        return None, None, None

def obtener_dem_opentopodata_api(gdf, dataset="srtm30m"):
    """
    Obtiene DEM desde la API pública Open Topo Data.
    Datasets disponibles: srtm30m, srtm90m, aster30m, eudem25m, etc.
    Límite gratuito: 1000 consultas/día, 100 puntos/consulta.
    Retorna (dem_array, meta, transform) compatible con el resto del código.
    """
    if not RASTERIO_OK:
        st.warning("⚠️ Rasterio no instalado. No se puede procesar DEM desde Open Topo Data.")
        return None, None, None

    try:
        bounds = gdf.total_bounds
        minx, miny, maxx, maxy = bounds

        # Definir resolución aproximada para grilla (máximo 50x50 para cumplir límite de 100 puntos)
        nx = 50
        ny = 50
        x_vals = np.linspace(minx, maxx, nx)
        y_vals = np.linspace(miny, maxy, ny)

        # Construir lista de ubicaciones (lat,lon) para la API
        locations = []
        for y in y_vals:
            for x in x_vals:
                locations.append(f"{y:.6f},{x:.6f}")

        # Dividir en lotes de 100 (límite de la API)
        batch_size = 100
        all_elevations = []
        for i in range(0, len(locations), batch_size):
            batch = locations[i:i+batch_size]
            loc_str = "|".join(batch)
            url = f"https://api.opentopodata.org/v1/{dataset}"
            params = {"locations": loc_str, "interpolation": "cubic"}

            with st.spinner(f"📡 Consultando lote {i//batch_size + 1} de Open Topo Data..."):
                resp = requests.get(url, params=params, timeout=30)

            if resp.status_code != 200:
                st.error(f"Error en API Open Topo Data: HTTP {resp.status_code}")
                return None, None, None

            data = resp.json()
            if data.get('status') != 'OK':
                st.error(f"Error en respuesta: {data.get('error', 'desconocido')}")
                return None, None, None

            elevations = [r['elevation'] for r in data['results']]
            all_elevations.extend(elevations)

        # Reconstruir grilla
        Z = np.array(all_elevations).reshape(ny, nx)
        X, Y = np.meshgrid(x_vals, y_vals)

        # Crear una transformación aproximada (para compatibilidad con código que espera transform)
        # La transform de rasterio: (res_x, 0, minx, 0, -res_y, maxy) si se usa from_origin
        # Pero aquí podemos usar None y luego tratar como DEM sintético
        # Para simplificar, devolvemos None en transform y construiremos X,Y,Z en dem_data
        # Creamos un array enmascarado con NaN fuera del polígono
        points = np.vstack([X.ravel(), Y.ravel()]).T
        mask = gdf.geometry.unary_union.contains([Point(p) for p in points])
        mask = mask.reshape(X.shape)
        Z_masked = Z.copy().astype(float)
        Z_masked[~mask] = np.nan
        dem_array = np.ma.masked_invalid(Z_masked)

        # Meta información básica
        meta = {
            'driver': 'GTiff',
            'height': ny,
            'width': nx,
            'count': 1,
            'crs': CRS.from_epsg(4326),
            'transform': None  # No tenemos transform real, lo manejaremos aparte
        }

        st.success(f"✅ DEM obtenido de Open Topo Data ({dataset}) - {nx}x{ny} puntos")
        return dem_array, meta, None  # transform = None

    except Exception as e:
        st.error(f"❌ Error obteniendo DEM de Open Topo Data: {str(e)}")
        return None, None, None

def generar_curvas_nivel_reales(dem_array, transform, intervalo=10, polygon=None):
    """
    Genera curvas de nivel a partir de un DEM real (array) y su transform.
    Opcionalmente filtra curvas que intersecten el polígono de la parcela.
    Requiere scikit-image.
    """
    if dem_array is None or not SKIMAGE_OK:
        return []

    # Enmascarar nodata
    if isinstance(dem_array, np.ma.MaskedArray):
        valid_mask = ~dem_array.mask
        data = dem_array.data.astype(float)
        data[~valid_mask] = np.nan
    else:
        data = dem_array.astype(float)
        valid_mask = data > -32768
        data[~valid_mask] = np.nan

    if not np.any(valid_mask):
        st.warning("⚠️ El DEM no contiene datos válidos para generar curvas.")
        return []

    vmin = np.nanmin(data)
    vmax = np.nanmax(data)
    if np.isnan(vmin) or np.isnan(vmax):
        return []

    niveles = np.arange(np.floor(vmin / intervalo) * intervalo,
                        np.ceil(vmax / intervalo) * intervalo + intervalo,
                        intervalo)

    # Si el rango es muy pequeño, usar un intervalo más fino
    if vmax - vmin < intervalo * 2 and len(niveles) < 3:
        intervalo_ajustado = (vmax - vmin) / 5
        niveles = np.arange(vmin, vmax + intervalo_ajustado, intervalo_ajustado)
        st.info(f"ℹ️ Terreno muy plano: se usó intervalo de {intervalo_ajustado:.1f} m en lugar de {intervalo} m")

    # Rellenar con un valor muy negativo para que find_contours no se salga
    data_filled = np.where(valid_mask, data, -9999)

    contours = []
    for nivel in niveles:
        try:
            for contour in measure.find_contours(data_filled, nivel):
                coords = []
                valid_contour = True
                for row, col in contour:
                    r, c = int(round(row)), int(round(col))
                    # Verificar que el punto esté dentro del array y sea válido
                    if not (0 <= r < data.shape[0] and 0 <= c < data.shape[1]) or not valid_mask[r, c]:
                        valid_contour = False
                        break
                    x, y = transform * (col, row)
                    coords.append((x, y))
                if valid_contour and len(coords) >= 3:
                    line = LineString(coords)
                    # Filtrar líneas muy cortas y opcionalmente por polígono
                    if line.length > 0.01:
                        if polygon is None or line.intersects(polygon):
                            contours.append((line, nivel))
        except Exception:
            continue
    if contours:
        st.info(f"✅ Generadas {len(contours)} curvas de nivel (intervalo {intervalo} m)")
    else:
        st.warning("⚠️ No se generaron curvas de nivel. El terreno puede ser muy plano o el DEM no tiene variación.")
    return contours

def generar_curvas_nivel_simuladas(gdf, intervalo=10):
    """
    Genera curvas de nivel sintéticas cuando no hay DEM real.
    También puede usarse para datos provenientes de Open Topo Data (X,Y,Z ya definidos).
    Requiere scikit-image.
    """
    if not SKIMAGE_OK:
        return []
    from scipy.ndimage import gaussian_filter
    bounds = gdf.total_bounds
    minx, miny, maxx, maxy = bounds
    n = 200  # Mayor resolución para más detalle
    x = np.linspace(minx, maxx, n)
    y = np.linspace(miny, maxy, n)
    X, Y = np.meshgrid(x, y)

    # Semilla reproducible basada en la ubicación
    seed = int((minx + miny) * 1e6) % (2**32)
    rng = np.random.RandomState(seed)

    # Generar relieve con varias ondas
    Z = rng.randn(n, n) * 30
    Z = gaussian_filter(Z, sigma=8)
    # Añadir gradiente y colinas
    Z = 50 + Z + 0.01 * (X - minx) * 111000 + 0.005 * (Y - miny) * 111000
    for _ in range(5):
        cx = rng.uniform(minx, maxx)
        cy = rng.uniform(miny, maxy)
        r = rng.uniform(0.001, 0.008)
        h = rng.uniform(30, 100)
        Z += h * np.exp(-((X-cx)**2 + (Y-cy)**2) / (2*r**2))

    # Enmascarar fuera del polígono
    points = np.vstack([X.ravel(), Y.ravel()]).T
    mask = gdf.geometry.unary_union.contains([Point(p) for p in points])
    mask = mask.reshape(X.shape)
    Z[~mask] = np.nan

    # Rellenar NaN con valor muy bajo para find_contours
    Z_filled = np.where(np.isnan(Z), -9999, Z)

    vmin = np.nanmin(Z)
    vmax = np.nanmax(Z)
    if np.isnan(vmin) or np.isnan(vmax):
        return []

    niveles = np.arange(vmin, vmax + intervalo, intervalo)
    if len(niveles) < 2:
        return []

    contours = []
    polygon = gdf.geometry.unary_union
    for nivel in niveles:
        try:
            for contour in measure.find_contours(Z_filled, nivel):
                coords = []
                for row, col in contour:
                    r, c = int(round(row)), int(round(col))
                    if r < 0 or r >= n or c < 0 or c >= n or np.isnan(Z[r, c]):
                        continue
                    lon = minx + (c / n) * (maxx - minx)
                    lat = miny + (r / n) * (maxy - miny)
                    coords.append((lon, lat))
                if len(coords) >= 3:
                    line = LineString(coords)
                    if line.length > 0.01 and line.intersects(polygon):
                        contours.append((line, nivel))
        except Exception:
            continue
    if contours:
        st.info(f"✅ Generadas {len(contours)} curvas de nivel sintéticas (intervalo {intervalo} m)")
    else:
        st.warning("⚠️ No se generaron curvas de nivel sintéticas.")
    return contours

def mapa_curvas_coloreadas(gdf_original, curvas_con_elevacion):
    """
    Crea un mapa Folium interactivo con las curvas de nivel coloreadas por elevación.
    Requiere folium.
    """
    if not FOLIUM_OK:
        st.error("Folium no está instalado. No se puede generar el mapa interactivo.")
        return None

    centroide = gdf_original.geometry.unary_union.centroid
    m = folium.Map(location=[centroide.y, centroide.x], zoom_start=15, tiles=None, control_scale=True)

    # Capas base
    folium.TileLayer(
        'https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
        attr='Esri, Maxar, Earthstar Geographics',
        name='Satélite Esri',
        overlay=False,
        control=True
    ).add_to(m)
    folium.TileLayer(
        'https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png',
        attr='OpenStreetMap',
        name='OpenStreetMap',
        overlay=False,
        control=True
    ).add_to(m)

    # Parcela
    folium.GeoJson(
        gdf_original.to_json(),
        name='Parcela',
        style_function=lambda x: {'color': 'blue', 'fillOpacity': 0.1, 'weight': 2}
    ).add_to(m)

    # Curvas de nivel
    if curvas_con_elevacion:
        elevaciones = [e for _, e in curvas_con_elevacion]
        vmin = min(elevaciones)
        vmax = max(elevaciones)
        colormap = LinearColormap(
            colors=['green', 'yellow', 'orange', 'brown'],
            vmin=vmin, vmax=vmax,
            caption='Elevación (m.s.n.m)'
        )
        colormap.add_to(m)

        for line, elev in curvas_con_elevacion:
            folium.GeoJson(
                gpd.GeoSeries(line).to_json(),
                style_function=lambda x, e=elev: {'color': colormap(e), 'weight': 1.5, 'opacity': 0.9},
                tooltip=f'Elevación: {elev:.0f} m'
            ).add_to(m)
    else:
        folium.Marker(
            [centroide.y, centroide.x],
            popup='No se generaron curvas de nivel',
            icon=folium.Icon(color='red')
        ).add_to(m)

    folium.LayerControl(collapsed=False).add_to(m)
    from folium.plugins import Fullscreen
    Fullscreen().add_to(m)
    return m

def generar_dem_sintetico_fallback(gdf, resolucion=10.0):
    """
    Función de respaldo para obtener X, Y, Z cuando no hay DEM real.
    No requiere rasterio ni skimage.
    """
    bounds = gdf.total_bounds
    minx, miny, maxx, maxy = bounds

    num_cells_x = int((maxx - minx) * 111000 / resolucion)
    num_cells_y = int((maxy - miny) * 111000 / resolucion)
    num_cells_x = max(50, min(num_cells_x, 200))
    num_cells_y = max(50, min(num_cells_y, 200))

    x = np.linspace(minx, maxx, num_cells_x)
    y = np.linspace(miny, maxy, num_cells_y)
    X, Y = np.meshgrid(x, y)

    centroid = gdf.geometry.unary_union.centroid
    seed_value = int(centroid.x * 10000 + centroid.y * 10000) % (2**32)
    rng = np.random.RandomState(seed_value)

    elevacion_base = rng.uniform(100, 300)
    slope_x = rng.uniform(-0.001, 0.001)
    slope_y = rng.uniform(-0.001, 0.001)

    Z = elevacion_base + slope_x * (X - minx) + slope_y * (Y - miny)
    n_hills = rng.randint(3, 7)
    for _ in range(n_hills):
        cx = rng.uniform(minx, maxx)
        cy = rng.uniform(miny, maxy)
        r = rng.uniform(0.001, 0.005)
        h = rng.uniform(20, 80)
        Z += h * np.exp(-((X-cx)**2 + (Y-cy)**2) / (2*r**2))

    # enmascarar fuera de la parcela
    points = np.vstack([X.flatten(), Y.flatten()]).T
    mask = gdf.geometry.unary_union.contains([Point(p) for p in points])
    mask = mask.reshape(X.shape)
    Z[~mask] = np.nan

    return X, Y, Z, bounds


# ===== INICIALIZACIÓN DE VARIABLES DE SESIÓN =====
if 'reporte_completo' not in st.session_state:
    st.session_state.reporte_completo = None
if 'geojson_data' not in st.session_state:
    st.session_state.geojson_data = None
if 'nombre_geojson' not in st.session_state:
    st.session_state.nombre_geojson = ""
if 'nombre_reporte' not in st.session_state:
    st.session_state.nombre_reporte = ""
if 'resultados_todos' not in st.session_state:
    st.session_state.resultados_todos = {}
if 'analisis_completado' not in st.session_state:
    st.session_state.analisis_completado = False
if 'mapas_generados' not in st.session_state:
    st.session_state.mapas_generados = {}
if 'dem_data' not in st.session_state:
    st.session_state.dem_data = {}
if 'gee_authenticated' not in st.session_state:
    st.session_state.gee_authenticated = False
if 'gee_project' not in st.session_state:
    st.session_state.gee_project = ''
if 'modelo_yolo' not in st.session_state:
    st.session_state.modelo_yolo = None
if 'curvas_nivel' not in st.session_state:
    st.session_state.curvas_nivel = None

# ===== ESTILOS PERSONALIZADOS - VERSIÓN COMPATIBLE CON STREAMLIT CLOUD =====
st.markdown("""
<style>
/* === FONDO GENERAL OSCURO ELEGANTE === */
.stApp {
    background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
    color: #ffffff;
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}

/* === BANNER HERO SIN IMÁGENES EXTERNAS (100% CSS) === */
.hero-banner {
    background: linear-gradient(145deg, rgba(15, 23, 42, 0.95), rgba(30, 41, 59, 0.98)),
                radial-gradient(circle at 20% 30%, rgba(59, 130, 246, 0.15), transparent 40%),
                radial-gradient(circle at 80% 70%, rgba(16, 185, 129, 0.1), transparent 45%);
    padding: 2.5em 1.5em;
    border-radius: 20px;
    margin-bottom: 2em;
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.4);
    border: 1px solid rgba(59, 130, 246, 0.3);
    position: relative;
    overflow: hidden;
    text-align: center;
}

.hero-banner::before {
    content: '';
    position: absolute;
    top: -50%;
    left: -50%;
    width: 200%;
    height: 200%;
    background: radial-gradient(circle, rgba(59, 130, 246, 0.08) 0%, transparent 70%);
    z-index: 0;
}

.hero-content {
    position: relative;
    z-index: 2;
    padding: 1.5em;
}

.hero-title {
    color: #ffffff;
    font-size: 2.8em;
    font-weight: 800;
    margin-bottom: 0.5em;
    text-shadow: 0 4px 12px rgba(0, 0, 0, 0.5);
    background: linear-gradient(135deg, #ffffff 0%, #60a5fa 50%, #3b82f6 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: -0.5px;
}

.hero-subtitle {
    color: #cbd5e1;
    font-size: 1.2em;
    font-weight: 400;
    max-width: 700px;
    margin: 0 auto;
    line-height: 1.6;
    opacity: 0.95;
}

/* === DECORACIÓN DEL BANNER (cultivos abstractos) === */
.hero-banner::after {
    content: '🌾 🌾 🌾 🌾 🌾 🌾 🌾 🌾 🌾 🌾';
    position: absolute;
    bottom: -15px;
    left: 0;
    right: 0;
    font-size: 1.8em;
    letter-spacing: 12px;
    color: rgba(255, 255, 255, 0.15);
    text-align: center;
    z-index: 1;
    transform: scale(1.2);
}

/* === SIDEBAR: FONDO BLANCO CON TEXTO NEGRO === */
[data-testid="stSidebar"] {
    background: #ffffff !important;
    border-right: 1px solid #e2e8f0 !important;
    box-shadow: 2px 0 15px rgba(0, 0, 0, 0.08) !important;
}

/* Texto general del sidebar en NEGRO */
[data-testid="stSidebar"] *,
[data-testid="stSidebar"] .stMarkdown,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stText,
[data-testid="stSidebar"] .stTitle,
[data-testid="stSidebar"] .stSubheader { 
    color: #000000 !important;
    text-shadow: none !important;
}

/* Título del sidebar elegante */
.sidebar-title {
    font-size: 1.4em;
    font-weight: 800;
    margin: 1.5em 0 1em 0;
    text-align: center;
    padding: 14px;
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%);
    border-radius: 16px;
    color: #ffffff !important;
    box-shadow: 0 4px 12px rgba(59, 130, 246, 0.25);
    border: 1px solid rgba(255, 255, 255, 0.2);
    letter-spacing: 0.5px;
}

/* Widgets del sidebar */
[data-testid="stSidebar"] .stSelectbox,
[data-testid="stSidebar"] .stDateInput,
[data-testid="stSidebar"] .stSlider {
    background: rgba(255, 255, 255, 0.95) !important;
    backdrop-filter: blur(8px);
    border-radius: 12px;
    padding: 12px;
    margin: 8px 0;
    border: 1px solid #d1d5db !important;
    box-shadow: 0 2px 6px rgba(0, 0, 0, 0.05) !important;
}

/* Botones premium */
.stButton > button {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: white !important;
    border: none !important;
    padding: 0.8em 1.5em !important;
    border-radius: 12px !important;
    font-weight: 700 !important;
    font-size: 1em !important;
    box-shadow: 0 4px 12px rgba(59, 130, 246, 0.35) !important;
    transition: all 0.25s ease !important;
    text-transform: uppercase !important;
    letter-spacing: 0.5px !important;
}

.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 18px rgba(59, 130, 246, 0.45) !important;
    background: linear-gradient(135deg, #4f8df8 0%, #2d5fe8 100%) !important;
}

/* === PESTAÑAS === */
.stTabs [data-baseweb="tab-list"] {
    background: rgba(30, 41, 59, 0.7) !important;
    backdrop-filter: blur(10px) !important;
    padding: 8px 16px !important;
    border-radius: 16px !important;
    border: 1px solid rgba(59, 130, 246, 0.3) !important;
    margin-top: 1.5em !important;
    gap: 6px !important;
}

.stTabs [data-baseweb="tab"] {
    color: #94a3b8 !important;
    font-weight: 600 !important;
    padding: 10px 20px !important;
    border-radius: 12px !important;
    background: rgba(15, 23, 42, 0.6) !important;
    transition: all 0.25s ease !important;
    border: 1px solid rgba(56, 189, 248, 0.2) !important;
}

.stTabs [data-baseweb="tab"]:hover {
    color: #ffffff !important;
    background: rgba(59, 130, 246, 0.2) !important;
    border-color: rgba(59, 130, 246, 0.4) !important;
    transform: translateY(-1px) !important;
}

.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: #ffffff !important;
    font-weight: 700 !important;
    border: none !important;
    box-shadow: 0 4px 15px rgba(59, 130, 246, 0.4) !important;
}

/* === MÉTRICAS === */
div[data-testid="metric-container"] {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.9), rgba(15, 23, 42, 0.95)) !important;
    backdrop-filter: blur(10px) !important;
    border-radius: 18px !important;
    padding: 22px !important;
    box-shadow: 0 6px 20px rgba(0, 0, 0, 0.35) !important;
    border: 1px solid rgba(59, 130, 246, 0.25) !important;
    transition: all 0.3s ease !important;
}

div[data-testid="metric-container"]:hover {
    transform: translateY(-4px) !important;
    box-shadow: 0 10px 25px rgba(59, 130, 246, 0.3) !important;
    border-color: rgba(59, 130, 246, 0.45) !important;
}

div[data-testid="metric-container"] label,
div[data-testid="metric-container"] div,
div[data-testid="metric-container"] [data-testid="stMetricValue"] { 
    color: #ffffff !important;
    font-weight: 600 !important;
}

div[data-testid="metric-container"] [data-testid="stMetricValue"] {
    font-size: 2.3em !important;
    font-weight: 800 !important;
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
    background-clip: text !important;
}

/* === DATAFRAMES === */
.dataframe {
    background: rgba(15, 23, 42, 0.85) !important;
    backdrop-filter: blur(8px) !important;
    border-radius: 14px !important;
    border: 1px solid rgba(255, 255, 255, 0.12) !important;
    color: #e2e8f0 !important;
    font-size: 0.95em !important;
}

.dataframe th {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: #ffffff !important;
    font-weight: 700 !important;
    padding: 14px 16px !important;
}

.dataframe td {
    color: #cbd5e1 !important;
    padding: 12px 16px !important;
    border-bottom: 1px solid rgba(255, 255, 255, 0.08) !important;
}

/* === FOOTER === */
.footer-divider {
    margin: 2.5em 0 1.5em 0;
    border-top: 1px solid rgba(59, 130, 246, 0.3);
}

.footer-content {
    background: rgba(15, 23, 42, 0.92);
    backdrop-filter: blur(12px);
    border-radius: 16px;
    padding: 1.8em;
    border: 1px solid rgba(59, 130, 246, 0.2);
    margin-top: 1.5em;
}

.footer-copyright {
    text-align: center;
    color: #94a3b8;
    padding: 1.2em 0 0.8em 0;
    font-size: 0.95em;
    border-top: 1px solid rgba(255, 255, 255, 0.08);
    margin-top: 1.5em;
}
</style>
""", unsafe_allow_html=True)

# ===== BANNER HERO CORREGIDO (100% CSS - SIN IMÁGENES EXTERNAS) =====
st.markdown("""
<div class="hero-banner">
    <div class="hero-content">
        <h1 class="hero-title">🌾 ANALIZADOR MULTI-CULTIVO SATELITAL</h1>
        <p class="hero-subtitle">Potenciado con Google Earth Engine, NASA POWER y datos SRTM para agricultura de precisión</p>
    </div>
</div>
""", unsafe_allow_html=True)

# ===== CONFIGURACIÓN DE SATÉLITES DISPONIBLES =====
SATELITES_DISPONIBLES = {
    'SENTINEL-2_GEE': {
        'nombre': 'Sentinel-2 (Google Earth Engine)',
        'resolucion': '10m',
        'revisita': '5 días',
        'bandas': ['B2', 'B3', 'B4', 'B8', 'B5', 'B11', 'B12'],
        'indices': ['NDVI', 'NDRE', 'NDWI', 'EVI', 'SAVI', 'MSAVI'],
        'icono': '🌍',
        'requerimiento': 'Google Earth Engine'
    },
    'LANDSAT-8_GEE': {
        'nombre': 'Landsat 8 (Google Earth Engine)',
        'resolucion': '30m',
        'revisita': '16 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B6', 'B7'],
        'indices': ['NDVI', 'NDRE', 'NDWI', 'EVI', 'SAVI', 'MSAVI'],
        'icono': '🌍',
        'requerimiento': 'Google Earth Engine'
    },
    'LANDSAT-9_GEE': {
        'nombre': 'Landsat 9 (Google Earth Engine)',
        'resolucion': '30m',
        'revisita': '16 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B6', 'B7'],
        'indices': ['NDVI', 'NDRE', 'NDWI', 'EVI', 'SAVI', 'MSAVI'],
        'icono': '🌍',
        'requerimiento': 'Google Earth Engine'
    },
    'SENTINEL-2': {
        'nombre': 'Sentinel-2 (Simulado)',
        'resolucion': '10m',
        'revisita': '5 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B8', 'B11'],
        'indices': ['NDVI', 'NDRE', 'GNDVI', 'OSAVI', 'MCARI'],
        'icono': '🛰️'
    },
    'LANDSAT-8': {
        'nombre': 'Landsat 8 (Simulado)',
        'resolucion': '30m',
        'revisita': '16 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B6', 'B7'],
        'indices': ['NDVI', 'NDRE', 'NDWI', 'EVI', 'SAVI', 'MSAVI'],
        'icono': '🛰️'
    },
    'DATOS_SIMULADOS': {
        'nombre': 'Datos Simulados',
        'resolucion': '10m',
        'revisita': '5 días',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B8'],
        'indices': ['NDVI', 'NDRE', 'GNDVI'],
        'icono': '🔬'
    }
}

# ===== CONFIGURACIÓN VARIEDADES CULTIVOS (ACTUALIZADO CON NUEVOS CULTIVOS) =====
VARIEDADES_CULTIVOS = {
    'TRIGO': [
        'ACA 303', 'ACA 315', 'Baguette Premium 11', 'Baguette Premium 13',
        'Biointa 1005', 'Biointa 2004', 'Klein Don Enrique', 'Klein Guerrero',
        'Buck Meteoro', 'Buck Poncho', 'SY 110', 'SY 200'
    ],
    'MAIZ': [
        'DK 72-10', 'DK 73-20', 'Pioneer 30F53', 'Pioneer 30F35',
        'Syngenta AG 6800', 'Syngenta AG 8088', 'Dow 2A610', 'Dow 2B710',
        'Nidera 8710', 'Nidera 8800', 'Morgan 360', 'Morgan 390'
    ],
    'SORGO': [
        'Advanta AS 5405', 'Advanta AS 5505', 'Pioneer 84G62', 'Pioneer 85G96',
        'DEKALB 53-67', 'DEKALB 55-00', 'MACER S-10', 'MACER S-15',
        'Sorgocer 105', 'Sorgocer 110', 'Río IV 100', 'Río IV 110'
    ],
    'SOJA': [
        'DM 53i52', 'DM 58i62', 'Nidera 49X', 'Nidera 52X',
        'Don Mario 49X', 'Don Mario 52X', 'SYNGENTA 4.9i', 'SYNGENTA 5.2i',
        'Biosoys 4.9', 'Biosoys 5.2', 'ACA 49', 'ACA 52'
    ],
    'GIRASOL': [
        'ACA 884', 'ACA 887', 'Nidera 7120', 'Nidera 7150',
        'Syngenta 390', 'Syngenta 410', 'Pioneer 64A15', 'Pioneer 65A25',
        'Advanta G 100', 'Advanta G 110', 'Biosun 400', 'Biosun 420'
    ],
    'MANI': [
        'ASEM 400', 'ASEM 500', 'Granoleico', 'Guasu',
        'Florman INTA', 'Elena', 'Colorado Irradiado', 'Overo Colorado',
        'Runner 886', 'Runner 890', 'Tegua', 'Virginia 98R'
    ],
    'VID': [
        'Malbec', 'Cabernet Sauvignon', 'Merlot', 'Syrah', 'Chardonnay',
        'Torrontés', 'Bonarda', 'Tempranillo', 'Sangiovese', 'Pinot Noir',
        'Chenin', 'Sauvignon Blanc', 'Viognier', 'Carménère', 'Petit Verdot'
    ],
    'OLIVO': [
        'Arbequina', 'Picual', 'Manzanilla', 'Hojiblanca', 'Cornicabra',
        'Empeltre', 'Frantoio', 'Leccino', 'Coratina', 'Picholine',
        'Kalamata', 'Mission', 'Ascolano', 'Barnea', 'Arbosana'
    ],
    'ALMENDRO': [
        'Non Pareil', 'Carmel', 'Butte', 'Padre', 'Mission',
        'Fritz', 'Monterey', 'Price', 'Aldrich', 'Wood Colony',
        'Peerless', 'Thompson', 'Livingston', 'Sonora', 'Winters'
    ],
    'BANANO': [
        'Cavendish', 'Gros Michel', 'Plátano', 'Manzano', 'Rojo',
        'Morado', 'Baby Banana', 'Blue Java', 'Goldfinger', 'Pisang Awak',
        'Mysore', 'Saba', 'Lakatan', 'Señorita', 'Dwarf Cavendish'
    ],
    'CAFE': [
        'Arabica', 'Robusta', 'Liberica', 'Excelsa', 'Typica',
        'Bourbon', 'Caturra', 'Catuai', 'Mundo Novo', 'Maragogipe',
        'Geisha', 'Pacamara', 'SL-28', 'SL-34', 'Kona'
    ],
    'CACAO': [
        'Forastero', 'Criollo', 'Trinitario', 'Nacional', 'Amelonado',
        'Contamana', 'Marañón', 'Porcelana', 'Chuao', 'Carenero',
        'Ocumare', 'Cundeamor', 'ICS-95', 'UF-613', 'TSH-565'
    ],
    'PALMA_ACEITERA': [
        'Tenera', 'Dura', 'Pisifera', 'DxP', 'Yangambi',
        'AVROS', 'La Mé', 'Ekona', 'Calabar', 'NIFOR',
        'MARDI', 'CIRAD', 'ASD Costa Rica', 'Dami', 'Socfindo'
    ]
}

# ===== CONFIGURACIÓN PARÁMETROS CULTIVOS (ACTUALIZADO) =====
PARAMETROS_CULTIVOS = {
    'TRIGO': {
        'NITROGENO': {'min': 100, 'max': 180},
        'FOSFORO': {'min': 40, 'max': 80},
        'POTASIO': {'min': 90, 'max': 150},
        'MATERIA_ORGANICA_OPTIMA': 3.5,
        'HUMEDAD_OPTIMA': 0.28,
        'NDVI_OPTIMO': 0.75,
        'NDRE_OPTIMO': 0.40,
        'RENDIMIENTO_OPTIMO': 4500,
        'COSTO_FERTILIZACION': 350,
        'PRECIO_VENTA': 0.25,
        'VARIEDADES': VARIEDADES_CULTIVOS['TRIGO'],
        'ZONAS_ARGENTINA': ['Pampeana', 'Noroeste', 'Noreste']
    },
    'MAIZ': {
        'NITROGENO': {'min': 150, 'max': 250},
        'FOSFORO': {'min': 50, 'max': 90},
        'POTASIO': {'min': 120, 'max': 200},
        'MATERIA_ORGANICA_OPTIMA': 3.8,
        'HUMEDAD_OPTIMA': 0.32,
        'NDVI_OPTIMO': 0.80,
        'NDRE_OPTIMO': 0.45,
        'RENDIMIENTO_OPTIMO': 8500,
        'COSTO_FERTILIZACION': 550,
        'PRECIO_VENTA': 0.20,
        'VARIEDADES': VARIEDADES_CULTIVOS['MAIZ'],
        'ZONAS_ARGENTINA': ['Pampeana', 'Noroeste', 'Noreste', 'Cuyo']
    },
    'SORGO': {
        'NITROGENO': {'min': 80, 'max': 140},
        'FOSFORO': {'min': 35, 'max': 65},
        'POTASIO': {'min': 100, 'max': 180},
        'MATERIA_ORGANICA_OPTIMA': 3.0,
        'HUMEDAD_OPTIMA': 0.25,
        'NDVI_OPTIMO': 0.70,
        'NDRE_OPTIMO': 0.35,
        'RENDIMIENTO_OPTIMO': 5000,
        'COSTO_FERTILIZACION': 300,
        'PRECIO_VENTA': 0.18,
        'VARIEDADES': VARIEDADES_CULTIVOS['SORGO'],
        'ZONAS_ARGENTINA': ['Pampeana', 'Noroeste', 'Noreste']
    },
    'SOJA': {
        'NITROGENO': {'min': 20, 'max': 40},
        'FOSFORO': {'min': 45, 'max': 85},
        'POTASIO': {'min': 140, 'max': 220},
        'MATERIA_ORGANICA_OPTIMA': 3.5,
        'HUMEDAD_OPTIMA': 0.30,
        'NDVI_OPTIMO': 0.78,
        'NDRE_OPTIMO': 0.42,
        'RENDIMIENTO_OPTIMO': 3200,
        'COSTO_FERTILIZACION': 400,
        'PRECIO_VENTA': 0.45,
        'VARIEDADES': VARIEDADES_CULTIVOS['SOJA'],
        'ZONAS_ARGENTINA': ['Pampeana', 'Noroeste', 'Noreste']
    },
    'GIRASOL': {
        'NITROGENO': {'min': 70, 'max': 120},
        'FOSFORO': {'min': 40, 'max': 75},
        'POTASIO': {'min': 110, 'max': 190},
        'MATERIA_ORGANICA_OPTIMA': 3.2,
        'HUMEDAD_OPTIMA': 0.26,
        'NDVI_OPTIMO': 0.72,
        'NDRE_OPTIMO': 0.38,
        'RENDIMIENTO_OPTIMO': 2800,
        'COSTO_FERTILIZACION': 320,
        'PRECIO_VENTA': 0.35,
        'VARIEDADES': VARIEDADES_CULTIVOS['GIRASOL'],
        'ZONAS_ARGENTINA': ['Pampeana', 'Noroeste', 'Noreste']
    },
    'MANI': {
        'NITROGENO': {'min': 15, 'max': 30},
        'FOSFORO': {'min': 50, 'max': 90},
        'POTASIO': {'min': 80, 'max': 140},
        'MATERIA_ORGANICA_OPTIMA': 2.8,
        'HUMEDAD_OPTIMA': 0.22,
        'NDVI_OPTIMO': 0.68,
        'NDRE_OPTIMO': 0.32,
        'RENDIMIENTO_OPTIMO': 3800,
        'COSTO_FERTILIZACION': 380,
        'PRECIO_VENTA': 0.60,
        'VARIEDADES': VARIEDADES_CULTIVOS['MANI'],
        'ZONAS_ARGENTINA': ['Córdoba', 'San Luis', 'La Pampa']
    },
    'VID': {
        'NITROGENO': {'min': 60, 'max': 120},
        'FOSFORO': {'min': 30, 'max': 70},
        'POTASIO': {'min': 150, 'max': 250},
        'MATERIA_ORGANICA_OPTIMA': 2.5,
        'HUMEDAD_OPTIMA': 0.35,
        'NDVI_OPTIMO': 0.65,
        'NDRE_OPTIMO': 0.35,
        'RENDIMIENTO_OPTIMO': 15000,
        'COSTO_FERTILIZACION': 800,
        'PRECIO_VENTA': 0.80,
        'VARIEDADES': VARIEDADES_CULTIVOS['VID'],
        'ZONAS_ARGENTINA': ['Mendoza', 'San Juan', 'La Rioja', 'Salta']
    },
    'OLIVO': {
        'NITROGENO': {'min': 40, 'max': 100},
        'FOSFORO': {'min': 20, 'max': 50},
        'POTASIO': {'min': 100, 'max': 200},
        'MATERIA_ORGANICA_OPTIMA': 2.0,
        'HUMEDAD_OPTIMA': 0.25,
        'NDVI_OPTIMO': 0.60,
        'NDRE_OPTIMO': 0.30,
        'RENDIMIENTO_OPTIMO': 8000,
        'COSTO_FERTILIZACION': 600,
        'PRECIO_VENTA': 1.20,
        'VARIEDADES': VARIEDADES_CULTIVOS['OLIVO'],
        'ZONAS_ARGENTINA': ['La Rioja', 'Catamarca', 'San Juan', 'Mendoza']
    },
    'ALMENDRO': {
        'NITROGENO': {'min': 80, 'max': 160},
        'FOSFORO': {'min': 40, 'max': 80},
        'POTASIO': {'min': 120, 'max': 200},
        'MATERIA_ORGANICA_OPTIMA': 2.2,
        'HUMEDAD_OPTIMA': 0.30,
        'NDVI_OPTIMO': 0.62,
        'NDRE_OPTIMO': 0.32,
        'RENDIMIENTO_OPTIMO': 3000,
        'COSTO_FERTILIZACION': 700,
        'PRECIO_VENTA': 4.50,
        'VARIEDADES': VARIEDADES_CULTIVOS['ALMENDRO'],
        'ZONAS_ARGENTINA': ['Río Negro', 'Neuquén', 'Mendoza', 'San Juan']
    },
    'BANANO': {
        'NITROGENO': {'min': 200, 'max': 350},
        'FOSFORO': {'min': 60, 'max': 120},
        'POTASIO': {'min': 300, 'max': 500},
        'MATERIA_ORGANICA_OPTIMA': 4.0,
        'HUMEDAD_OPTIMA': 0.45,
        'NDVI_OPTIMO': 0.78,
        'NDRE_OPTIMO': 0.40,
        'RENDIMIENTO_OPTIMO': 40000,
        'COSTO_FERTILIZACION': 1200,
        'PRECIO_VENTA': 0.30,
        'VARIEDADES': VARIEDADES_CULTIVOS['BANANO'],
        'ZONAS_ARGENTINA': ['Formosa', 'Misiones', 'Corrientes']
    },
    'CAFE': {
        'NITROGENO': {'min': 100, 'max': 200},
        'FOSFORO': {'min': 40, 'max': 80},
        'POTASIO': {'min': 150, 'max': 250},
        'MATERIA_ORGANICA_OPTIMA': 3.5,
        'HUMEDAD_OPTIMA': 0.40,
        'NDVI_OPTIMO': 0.70,
        'NDRE_OPTIMO': 0.38,
        'RENDIMIENTO_OPTIMO': 2000,
        'COSTO_FERTILIZACION': 900,
        'PRECIO_VENTA': 3.50,
        'VARIEDADES': VARIEDADES_CULTIVOS['CAFE'],
        'ZONAS_ARGENTINA': ['Misiones', 'Corrientes', 'Jujuy']
    },
    'CACAO': {
        'NITROGENO': {'min': 80, 'max': 150},
        'FOSFORO': {'min': 30, 'max': 60},
        'POTASIO': {'min': 120, 'max': 200},
        'MATERIA_ORGANICA_OPTIMA': 4.0,
        'HUMEDAD_OPTIMA': 0.50,
        'NDVI_OPTIMO': 0.72,
        'NDRE_OPTIMO': 0.38,
        'RENDIMIENTO_OPTIMO': 1500,
        'COSTO_FERTILIZACION': 850,
        'PRECIO_VENTA': 5.00,
        'VARIEDADES': VARIEDADES_CULTIVOS['CACAO'],
        'ZONAS_ARGENTINA': ['Misiones', 'Corrientes', 'Formosa']
    },
    'PALMA_ACEITERA': {
        'NITROGENO': {'min': 150, 'max': 250},
        'FOSFORO': {'min': 50, 'max': 100},
        'POTASIO': {'min': 200, 'max': 350},
        'MATERIA_ORGANICA_OPTIMA': 3.8,
        'HUMEDAD_OPTIMA': 0.55,
        'NDVI_OPTIMO': 0.75,
        'NDRE_OPTIMO': 0.42,
        'RENDIMIENTO_OPTIMO': 20000,
        'COSTO_FERTILIZACION': 1100,
        'PRECIO_VENTA': 0.40,
        'VARIEDADES': VARIEDADES_CULTIVOS['PALMA_ACEITERA'],
        'ZONAS_ARGENTINA': ['Formosa', 'Chaco', 'Misiones']
    }
}

# ===== CONFIGURACIÓN TEXTURA SUELO ÓPTIMA (ACTUALIZADO) =====
TEXTURA_SUELO_OPTIMA = {
    'TRIGO': {
        'textura_optima': 'Franco-arcilloso',
        'arena_optima': 35,
        'limo_optima': 40,
        'arcilla_optima': 25,
        'densidad_aparente_optima': 1.35,
        'porosidad_optima': 0.48
    },
    'MAIZ': {
        'textura_optima': 'Franco',
        'arena_optima': 45,
        'limo_optima': 35,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.30,
        'porosidad_optima': 0.50
    },
    'SORGO': {
        'textura_optima': 'Franco-arenoso',
        'arena_optima': 55,
        'limo_optima': 30,
        'arcilla_optima': 15,
        'densidad_aparente_optima': 1.40,
        'porosidad_optima': 0.45
    },
    'SOJA': {
        'textura_optima': 'Franco',
        'arena_optima': 40,
        'limo_optima': 40,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.25,
        'porosidad_optima': 0.52
    },
    'GIRASOL': {
        'textura_optima': 'Franco-arcilloso',
        'arena_optima': 30,
        'limo_optima': 45,
        'arcilla_optima': 25,
        'densidad_aparente_optima': 1.32,
        'porosidad_optima': 0.49
    },
    'MANI': {
        'textura_optima': 'Franco-arenoso',
        'arena_optima': 60,
        'limo_optima': 25,
        'arcilla_optima': 15,
        'densidad_aparente_optima': 1.38,
        'porosidad_optima': 0.46
    },
    'VID': {
        'textura_optima': 'Franco-arenoso',
        'arena_optima': 50,
        'limo_optima': 30,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.40,
        'porosidad_optima': 0.50
    },
    'OLIVO': {
        'textura_optima': 'Franco-arcilloso',
        'arena_optima': 40,
        'limo_optima': 35,
        'arcilla_optima': 25,
        'densidad_aparente_optima': 1.35,
        'porosidad_optima': 0.48
    },
    'ALMENDRO': {
        'textura_optima': 'Franco',
        'arena_optima': 45,
        'limo_optima': 35,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.38,
        'porosidad_optima': 0.47
    },
    'BANANO': {
        'textura_optima': 'Franco-arcilloso',
        'arena_optima': 35,
        'limo_optima': 40,
        'arcilla_optima': 25,
        'densidad_aparente_optima': 1.20,
        'porosidad_optima': 0.55
    },
    'CAFE': {
        'textura_optima': 'Franco',
        'arena_optima': 40,
        'limo_optima': 40,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.25,
        'porosidad_optima': 0.52
    },
    'CACAO': {
        'textura_optima': 'Franco-arcilloso',
        'arena_optima': 30,
        'limo_optima': 45,
        'arcilla_optima': 25,
        'densidad_aparente_optima': 1.15,
        'porosidad_optima': 0.56
    },
    'PALMA_ACEITERA': {
        'textura_optima': 'Franco',
        'arena_optima': 45,
        'limo_optima': 35,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.30,
        'porosidad_optima': 0.51
    }
}

# ===== ICONOS Y COLORES PARA CULTIVOS (ACTUALIZADO) =====
ICONOS_CULTIVOS = {
    'TRIGO': '🌾',
    'MAIZ': '🌽',
    'SORGO': '🌾',
    'SOJA': '🫘',
    'GIRASOL': '🌻',
    'MANI': '🥜',
    'VID': '🍇',
    'OLIVO': '🫒',
    'ALMENDRO': '🌰',
    'BANANO': '🍌',
    'CAFE': '☕',
    'CACAO': '🍫',
    'PALMA_ACEITERA': '🌴'
}

COLORES_CULTIVOS = {
    'TRIGO': '#FFD700',
    'MAIZ': '#F4A460',
    'SORGO': '#8B4513',
    'SOJA': '#228B22',
    'GIRASOL': '#FFD700',
    'MANI': '#D2691E',
    'VID': '#8B0000',
    'OLIVO': '#808000',
    'ALMENDRO': '#D2B48C',
    'BANANO': '#FFD700',
    'CAFE': '#8B4513',
    'CACAO': '#4A2C2A',
    'PALMA_ACEITERA': '#32CD32'
}

PALETAS_GEE = {
    'FERTILIDAD': ['#d73027', '#f46d43', '#fdae61', '#fee08b', '#d9ef8b', '#a6d96a', '#66bd63', '#1a9850', '#006837'],
    'NITROGENO': ['#00ff00', '#80ff00', '#ffff00', '#ff8000', '#ff0000'],
    'FOSFORO': ['#0000ff', '#4040ff', '#8080ff', '#c0c0ff', '#ffffff'],
    'POTASIO': ['#4B0082', '#6A0DAD', '#8A2BE2', '#9370DB', '#D8BFD8'],
    'TEXTURA': ['#8c510a', '#d8b365', '#f6e8c3', '#c7eae5', '#5ab4ac', '#01665e'],
    'ELEVACION': ['#006837', '#1a9850', '#66bd63', '#a6d96a', '#d9ef8b', '#ffffbf', '#fee08b', '#fdae61', '#f46d43', '#d73027'],
    'PENDIENTE': ['#4daf4a', '#a6d96a', '#ffffbf', '#fdae61', '#f46d43', '#d73027']
}

# ===== FUNCIÓN MEJORADA PARA MOSTRAR INFORMACIÓN DEL CULTIVO =====
def mostrar_info_cultivo(cultivo):
    """Muestra información específica del cultivo seleccionado"""
    if cultivo in PARAMETROS_CULTIVOS:
        params = PARAMETROS_CULTIVOS[cultivo]
        zonas = params.get('ZONAS_ARGENTINA', [])
        
        es_argentino = any("Internacional" not in zona for zona in zonas)
        
        st.markdown(f"""
        <div class="cultivo-card">
            <h3>{ICONOS_CULTIVOS[cultivo]} {cultivo} - Información {'Argentina' if es_argentino else 'Internacional'}</h3>
            <p><strong>Región principal:</strong> {', '.join(zonas)}</p>
            <p><strong>Variedades comunes:</strong></p>
            <ul>
        """, unsafe_allow_html=True)
        
        for variedad in params.get('VARIEDADES', [])[:5]:
            st.markdown(f"<li>{variedad}</li>", unsafe_allow_html=True)
        
        if len(params.get('VARIEDADES', [])) > 5:
            st.markdown(f"<li>... y {len(params.get('VARIEDADES', [])) - 5} más</li>", unsafe_allow_html=True)
        
        if not es_argentino:
            st.markdown("""
            </ul>
            <div style="background: rgba(59, 130, 246, 0.15); padding: 12px; border-radius: 8px; margin-top: 15px;">
                <p style="margin: 0; font-size: 0.9em; color: #60a5fa;">
                    💡 <strong>Nota:</strong> Este cultivo no es típico de Argentina. 
                    Los parámetros están adaptados para zonas productoras internacionales.
                </p>
            </div>
        """, unsafe_allow_html=True)
        else:
            st.markdown("</ul>", unsafe_allow_html=True)
        
        st.markdown("</div>", unsafe_allow_html=True)

# ===== SIDEBAR MEJORADO (INTERFAZ VISUAL) =====
with st.sidebar:
    st.markdown('<div class="sidebar-title">⚙️ CONFIGURACIÓN</div>', unsafe_allow_html=True)
    
    CULTIVOS_TOTALES = [
        "TRIGO", "MAIZ", "SORGO", "SOJA", "GIRASOL", "MANI",
        "VID", "OLIVO", "ALMENDRO", "BANANO", "CACAO", "CAFE", "PALMA_ACEITERA"
    ]
    
    cultivo = st.selectbox("Cultivo:", CULTIVOS_TOTALES)
    
    mostrar_info_cultivo(cultivo)

    variedades = VARIEDADES_CULTIVOS.get(cultivo, [])
    if variedades:
        variedad = st.selectbox(
            "Variedad/Cultivar:",
            ["No especificada"] + variedades,
            help="Selecciona la variedad o cultivar específico"
        )
    else:
        variedad = "No especificada"
        st.caption(f"ℹ️ Sin variedades predefinidas para {cultivo}")
    
    st.subheader("🌍 Google Earth Engine")
    if st.session_state.gee_authenticated:
        st.success(f"✅ Autenticado\nProyecto: {st.session_state.gee_project}")
    else:
        st.error("❌ No autenticado\nUsando datos simulados")
    
    st.subheader("🛰️ Fuente de Datos Satelitales")
    
    opciones_satelites = []
    if GEE_AVAILABLE and st.session_state.gee_authenticated:
        opciones_satelites.extend(["SENTINEL-2_GEE", "LANDSAT-8_GEE", "LANDSAT-9_GEE"])
    opciones_satelites.extend(["SENTINEL-2", "LANDSAT-8", "DATOS_SIMULADOS"])
    
    satelite_seleccionado = st.selectbox(
        "Satélite:",
        opciones_satelites,
        help="Selecciona la fuente de datos satelitales",
        index=0
    )
    
    if satelite_seleccionado in SATELITES_DISPONIBLES:
        info_satelite = SATELITES_DISPONIBLES[satelite_seleccionado]
        st.caption(f"{info_satelite['icono']} {info_satelite['nombre']} - {info_satelite['resolucion']}")
        if 'requerimiento' in info_satelite:
            st.caption(f"Requerimiento: {info_satelite['requerimiento']}")
    
    st.subheader("📊 Índice de Vegetación")
    if satelite_seleccionado in SATELITES_DISPONIBLES:
        indices_disponibles = SATELITES_DISPONIBLES[satelite_seleccionado]['indices']
        indice_seleccionado = st.selectbox("Índice:", indices_disponibles)

    st.subheader("📅 Rango Temporal")
    fecha_fin = st.date_input("Fecha fin", datetime.now())
    fecha_inicio = st.date_input("Fecha inicio", datetime.now() - timedelta(days=30))

    st.subheader("🎯 División de Parcela")
    n_divisiones = st.slider("Número de zonas de manejo:", min_value=16, max_value=48, value=32)

    st.subheader("🏔️ Configuración Curvas de Nivel")
    intervalo_curvas = st.slider("Intervalo entre curvas (metros):", 1.0, 20.0, 5.0, 1.0)
    resolucion_dem = st.slider("Resolución DEM (metros):", 5.0, 50.0, 10.0, 5.0)

    # No hay selector de proveedor, solo DeepSeek

    st.subheader("📤 Subir Parcela")
    uploaded_file = st.file_uploader("Subir archivo de tu parcela", type=['zip', 'kml', 'kmz'],
                                     help="Formatos aceptados: Shapefile (.zip), KML (.kml), KMZ (.kmz)")

# ===== FUNCIONES AUXILIARES - CORREGIDAS PARA EPSG:4326 =====
def validar_y_corregir_crs(gdf):
    if gdf is None or len(gdf) == 0:
        return gdf
    try:
        if gdf.crs is None:
            gdf = gdf.set_crs('EPSG:4326', inplace=False)
            st.info("ℹ️ Se asignó EPSG:4326 al archivo (no tenía CRS)")
        elif str(gdf.crs).upper() != 'EPSG:4326':
            original_crs = str(gdf.crs)
            gdf = gdf.to_crs('EPSG:4326')
            st.info(f"ℹ️ Transformado de {original_crs} a EPSG:4326")
        return gdf
    except Exception as e:
        st.warning(f"⚠️ Error al corregir CRS: {str(e)}")
        return gdf

def calcular_superficie(gdf):
    try:
        if gdf is None or len(gdf) == 0:
            return 0.0
        gdf = validar_y_corregir_crs(gdf)
        bounds = gdf.total_bounds
        if bounds[0] < -180 or bounds[2] > 180 or bounds[1] < -90 or bounds[3] > 90:
            area_grados2 = gdf.geometry.area.sum()
            area_m2 = area_grados2 * 111000 * 111000
            return area_m2 / 10000
        gdf_projected = gdf.to_crs('EPSG:3857')
        area_m2 = gdf_projected.geometry.area.sum()
        return area_m2 / 10000
    except Exception as e:
        try:
            return gdf.geometry.area.sum() / 10000
        except:
            return 0.0

def dividir_parcela_en_zonas(gdf, n_zonas):
    if len(gdf) == 0:
        return gdf
    gdf = validar_y_corregir_crs(gdf)
    parcela_principal = gdf.iloc[0].geometry
    bounds = parcela_principal.bounds
    minx, miny, maxx, maxy = bounds
    sub_poligonos = []
    n_cols = math.ceil(math.sqrt(n_zonas))
    n_rows = math.ceil(n_zonas / n_cols)
    width = (maxx - minx) / n_cols
    height = (maxy - miny) / n_rows
    for i in range(n_rows):
        for j in range(n_cols):
            if len(sub_poligonos) >= n_zonas:
                break
            cell_minx = minx + (j * width)
            cell_maxx = minx + ((j + 1) * width)
            cell_miny = miny + (i * height)
            cell_maxy = miny + ((i + 1) * height)
            cell_poly = Polygon([(cell_minx, cell_miny), (cell_maxx, cell_miny), (cell_maxx, cell_maxy), (cell_minx, cell_maxy)])
            intersection = parcela_principal.intersection(cell_poly)
            if not intersection.is_empty and intersection.area > 0:
                sub_poligonos.append(intersection)
    if sub_poligonos:
        nuevo_gdf = gpd.GeoDataFrame({'id_zona': range(1, len(sub_poligonos) + 1), 'geometry': sub_poligonos}, crs='EPSG:4326')
        return nuevo_gdf
    else:
        return gdf

# ===== FUNCIONES PARA CARGAR ARCHIVOS =====
def cargar_shapefile_desde_zip(zip_file):
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            with zipfile.ZipFile(zip_file, 'r') as zip_ref:
                zip_ref.extractall(tmp_dir)
            shp_files = [f for f in os.listdir(tmp_dir) if f.endswith('.shp')]
            if shp_files:
                shp_path = os.path.join(tmp_dir, shp_files[0])
                gdf = gpd.read_file(shp_path)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
            else:
                st.error("❌ No se encontró ningún archivo .shp en el ZIP")
                return None
    except Exception as e:
        st.error(f"❌ Error cargando shapefile desde ZIP: {str(e)}")
        return None

def parsear_kml_manual(contenido_kml):
    try:
        root = ET.fromstring(contenido_kml)
        namespaces = {'kml': 'http://www.opengis.net/kml/2.2'}
        polygons = []
        for polygon_elem in root.findall('.//kml:Polygon', namespaces):
            coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
            if coords_elem is not None and coords_elem.text:
                coord_text = coords_elem.text.strip()
                coord_list = []
                for coord_pair in coord_text.split():
                    parts = coord_pair.split(',')
                    if len(parts) >= 2:
                        lon = float(parts[0])
                        lat = float(parts[1])
                        coord_list.append((lon, lat))
                if len(coord_list) >= 3:
                    polygons.append(Polygon(coord_list))
        if not polygons:
            for multi_geom in root.findall('.//kml:MultiGeometry', namespaces):
                for polygon_elem in multi_geom.findall('.//kml:Polygon', namespaces):
                    coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
                    if coords_elem is not None and coords_elem.text:
                        coord_text = coords_elem.text.strip()
                        coord_list = []
                        for coord_pair in coord_text.split():
                            parts = coord_pair.split(',')
                            if len(parts) >= 2:
                                lon = float(parts[0])
                                lat = float(parts[1])
                                coord_list.append((lon, lat))
                        if len(coord_list) >= 3:
                            polygons.append(Polygon(coord_list))
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        else:
            for placemark in root.findall('.//kml:Placemark', namespaces):
                for elem_name in ['Polygon', 'LineString', 'Point', 'LinearRing']:
                    elem = placemark.find(f'.//kml:{elem_name}', namespaces)
                    if elem is not None:
                        coords_elem = elem.find('.//kml:coordinates', namespaces)
                        if coords_elem is not None and coords_elem.text:
                            coord_text = coords_elem.text.strip()
                            coord_list = []
                            for coord_pair in coord_text.split():
                                parts = coord_pair.split(',')
                                if len(parts) >= 2:
                                    lon = float(parts[0])
                                    lat = float(parts[1])
                                    coord_list.append((lon, lat))
                            if len(coord_list) >= 3:
                                polygons.append(Polygon(coord_list))
                            break
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        return None
    except Exception as e:
        st.error(f"❌ Error parseando KML manualmente: {str(e)}")
        return None

def cargar_kml(kml_file):
    try:
        if kml_file.name.endswith('.kmz'):
            with tempfile.TemporaryDirectory() as tmp_dir:
                with zipfile.ZipFile(kml_file, 'r') as zip_ref:
                    zip_ref.extractall(tmp_dir)
                kml_files = [f for f in os.listdir(tmp_dir) if f.endswith('.kml')]
                if kml_files:
                    kml_path = os.path.join(tmp_dir, kml_files[0])
                    with open(kml_path, 'r', encoding='utf-8') as f:
                        contenido = f.read()
                    gdf = parsear_kml_manual(contenido)
                    if gdf is not None:
                        return gdf
                    else:
                        try:
                            gdf = gpd.read_file(kml_path)
                            gdf = validar_y_corregir_crs(gdf)
                            return gdf
                        except:
                            st.error("❌ No se pudo cargar el archivo KML/KMZ")
                            return None
                else:
                    st.error("❌ No se encontró ningún archivo .kml en el KMZ")
                    return None
        else:
            contenido = kml_file.read().decode('utf-8')
            gdf = parsear_kml_manual(contenido)
            if gdf is not None:
                return gdf
            else:
                kml_file.seek(0)
                gdf = gpd.read_file(kml_file)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
    except Exception as e:
        st.error(f"❌ Error cargando archivo KML/KMZ: {str(e)}")
        return None

def cargar_archivo_parcela(uploaded_file):
    try:
        if uploaded_file.name.endswith('.zip'):
            gdf = cargar_shapefile_desde_zip(uploaded_file)
        elif uploaded_file.name.endswith(('.kml', '.kmz')):
            gdf = cargar_kml(uploaded_file)
        else:
            st.error("❌ Formato de archivo no soportado")
            return None
        
        if gdf is not None:
            gdf = validar_y_corregir_crs(gdf)
            gdf = gdf.explode(ignore_index=True)
            gdf = gdf[gdf.geometry.geom_type.isin(['Polygon', 'MultiPolygon'])]
            if len(gdf) == 0:
                st.error("❌ No se encontraron polígonos en el archivo")
                return None
            geometria_unida = gdf.unary_union
            gdf_unido = gpd.GeoDataFrame([{'geometry': geometria_unida}], crs='EPSG:4326')
            gdf_unido = validar_y_corregir_crs(gdf_unido)
            st.info(f"✅ Se unieron {len(gdf)} polígono(s) en una sola geometría.")
            gdf_unido['id_zona'] = 1
            return gdf_unido
        return gdf
    except Exception as e:
        st.error(f"❌ Error cargando archivo: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

# ===== FUNCIONES PARA DATOS SATELITALES =====
def descargar_datos_landsat8(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    try:
        datos_simulados = {
            'indice': indice,
            'valor_promedio': 0.65 + np.random.normal(0, 0.1),
            'fuente': 'Landsat-8',
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'id_escena': f"LC08_{np.random.randint(1000000, 9999999)}",
            'cobertura_nubes': f"{np.random.randint(0, 15)}%",
            'resolucion': '30m'
        }
        return datos_simulados
    except Exception as e:
        st.error(f"❌ Error procesando Landsat 8: {str(e)}")
        return None

def descargar_datos_sentinel2(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    try:
        datos_simulados = {
            'indice': indice,
            'valor_promedio': 0.72 + np.random.normal(0, 0.08),
            'fuente': 'Sentinel-2',
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'id_escena': f"S2A_{np.random.randint(1000000, 9999999)}",
            'cobertura_nubes': f"{np.random.randint(0, 10)}%",
            'resolucion': '10m'
        }
        return datos_simulados
    except Exception as e:
        st.error(f"❌ Error procesando Sentinel-2: {str(e)}")
        return None

def generar_datos_simulados(gdf, cultivo, indice='NDVI'):
    datos_simulados = {
        'indice': indice,
        'valor_promedio': PARAMETROS_CULTIVOS[cultivo]['NDVI_OPTIMO'] * 0.8 + np.random.normal(0, 0.1),
        'fuente': 'Simulación',
        'fecha': datetime.now().strftime('%Y-%m-%d'),
        'resolucion': '10m'
    }
    return datos_simulados

# ===== FUNCIONES GOOGLE EARTH ENGINE =====
def obtener_datos_sentinel2_gee(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    """Obtener datos reales de Sentinel-2 usando Google Earth Engine con manejo robusto"""
    if not GEE_AVAILABLE or not st.session_state.gee_authenticated:
        st.warning("⚠️ GEE no disponible o no autenticado")
        return None
    
    try:
        if gdf is None or len(gdf) == 0:
            st.error("❌ El área de estudio no es válida")
            return None
        
        bounds = gdf.total_bounds
        min_lon, min_lat, max_lon, max_lat = bounds
        
        if (abs(max_lon - min_lon) < 0.0001 or abs(max_lat - min_lat) < 0.0001):
            st.warning("⚠️ El área de estudio es muy pequeña. Ampliando bounding box.")
            min_lon -= 0.001
            max_lon += 0.001
            min_lat -= 0.001
            max_lat += 0.001
        
        geometry = ee.Geometry.Rectangle([min_lon, min_lat, max_lon, max_lat])
        start_date = fecha_inicio.strftime('%Y-%m-%d')
        end_date = fecha_fin.strftime('%Y-%m-%d')
        
        if fecha_inicio > fecha_fin:
            st.error("❌ La fecha de inicio debe ser anterior a la fecha de fin")
            start_date, end_date = end_date, start_date
            st.info("ℹ️ Se intercambiaron las fechas automáticamente")
        
        collection = (ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED')
                     .filterBounds(geometry)
                     .filterDate(start_date, end_date)
                     .filter(ee.Filter.lt('CLOUDY_PIXEL_PERCENTAGE', 60)))
        
        collection_size = collection.size().getInfo()
        
        if collection_size == 0:
            st.warning(f"⚠️ No se encontraron imágenes Sentinel-2 para:")
            st.warning(f"   - Área: [{min_lon:.4f}, {min_lat:.4f}, {max_lon:.4f}, {max_lat:.4f}]")
            st.warning(f"   - Período: {start_date} a {end_date}")
            st.info("🔄 Intentando con filtro de nubes más permisivo (<80%)...")
            collection = (ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED')
                         .filterBounds(geometry)
                         .filterDate(start_date, end_date)
                         .filter(ee.Filter.lt('CLOUDY_PIXEL_PERCENTAGE', 80)))
            collection_size = collection.size().getInfo()
            if collection_size == 0:
                st.error("❌ No hay imágenes disponibles incluso con filtro permisivo")
                return None
            else:
                st.success(f"✅ Encontradas {collection_size} imágenes con filtro permisivo")
        
        image = collection.sort('CLOUDY_PIXEL_PERCENTAGE').first()
        if image is None:
            st.error("❌ Error crítico: La imagen seleccionada es nula")
            return None
        
        image_id = image.get('system:index').getInfo()
        cloud_percent = image.get('CLOUDY_PIXEL_PERCENTAGE').getInfo()
        image_date = image.get('system:time_start').getInfo()
        
        if image_date:
            image_date_str = datetime.fromtimestamp(image_date / 1000).strftime('%Y-%m-%d')
            st.info(f"📅 Imagen seleccionada: {image_id} ({image_date_str}) - Nubes: {cloud_percent}%")
        
        try:
            if indice == 'NDVI':
                ndvi = image.normalizedDifference(['B8', 'B4']).rename('NDVI')
                index_image = ndvi
            elif indice == 'NDWI':
                ndwi = image.normalizedDifference(['B3', 'B8']).rename('NDWI')
                index_image = ndwi
            elif indice == 'EVI':
                evi = image.expression(
                    '2.5 * ((NIR - RED) / (NIR + 6 * RED - 7.5 * BLUE + 1))',
                    {'NIR': image.select('B8'), 'RED': image.select('B4'), 'BLUE': image.select('B2')}
                ).rename('EVI')
                index_image = evi
            elif indice == 'NDRE':
                ndre = image.normalizedDifference(['B8', 'B5']).rename('NDRE')
                index_image = ndre
            elif indice == 'SAVI':
                savi = image.expression(
                    '((NIR - RED) / (NIR + RED + 0.5)) * (1.5)',
                    {'NIR': image.select('B8'), 'RED': image.select('B4')}
                ).rename('SAVI')
                index_image = savi
            elif indice == 'MSAVI':
                msavi = image.expression(
                    '(2 * NIR + 1 - sqrt(pow((2 * NIR + 1), 2) - 8 * (NIR - RED))) / 2',
                    {'NIR': image.select('B8'), 'RED': image.select('B4')}
                ).rename('MSAVI')
                index_image = msavi
            else:
                ndvi = image.normalizedDifference(['B8', 'B4']).rename('NDVI')
                index_image = ndvi
                indice = 'NDVI'
        except Exception as e:
            st.error(f"❌ Error calculando índice {indice}: {str(e)}")
            try:
                ndvi = image.normalizedDifference(['B8', 'B4']).rename('NDVI')
                index_image = ndvi
                indice = 'NDVI'
                st.info("ℹ️ Usando NDVI como índice por defecto")
            except:
                st.error("❌ Error crítico: No se pudo calcular ningún índice")
                return None
        
        try:
            stats = index_image.reduceRegion(
                reducer=ee.Reducer.mean().combine(
                    reducer2=ee.Reducer.minMax(),
                    sharedInputs=True
                ).combine(
                    reducer2=ee.Reducer.stdDev(),
                    sharedInputs=True
                ),
                geometry=geometry,
                scale=10,
                bestEffort=True,
                maxPixels=1e9
            )
            
            stats_dict = stats.getInfo()
            
            if not stats_dict:
                st.warning("⚠️ No se pudieron obtener estadísticas de la imagen")
                valor_promedio = 0.6
                valor_min = 0.3
                valor_max = 0.9
                valor_std = 0.1
            else:
                valor_promedio = stats_dict.get(f'{indice}_mean', 0.6)
                valor_min = stats_dict.get(f'{indice}_min', 0.3)
                valor_max = stats_dict.get(f'{indice}_max', 0.9)
                valor_std = stats_dict.get(f'{indice}_stdDev', 0.1)
                
        except Exception as e:
            st.warning(f"⚠️ Error obteniendo estadísticas: {str(e)}")
            valor_promedio = 0.6 + np.random.normal(0, 0.1)
            valor_min = max(0.1, valor_promedio - 0.3)
            valor_max = min(0.95, valor_promedio + 0.3)
            valor_std = 0.1
        
        return {
            'indice': indice,
            'valor_promedio': valor_promedio,
            'valor_min': valor_min,
            'valor_max': valor_max,
            'valor_std': valor_std,
            'fuente': f'Sentinel-2 (Google Earth Engine) - {image_id}' if image_id else 'Sentinel-2 (GEE)',
            'fecha_descarga': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'fecha_imagen': image_date_str if 'image_date_str' in locals() else 'N/A',
            'resolucion': '10m',
            'estado': 'exitosa',
            'cobertura_nubes': f"{cloud_percent}%" if cloud_percent else 'N/A',
            'nota': f"Imágenes encontradas: {collection_size}" if collection_size else 'Sin imágenes'
        }
        
    except Exception as e:
        st.error(f"❌ Error obteniendo datos de Google Earth Engine: {str(e)}")
        st.info("💡 Usando datos simulados como alternativa")
        return None

def obtener_datos_landsat_gee(gdf, fecha_inicio, fecha_fin, dataset='LANDSAT/LC08/C02/T1_L2', indice='NDVI'):
    if not GEE_AVAILABLE or not st.session_state.gee_authenticated:
        return None
    try:
        bounds = gdf.total_bounds
        min_lon, min_lat, max_lon, max_lat = bounds
        geometry = ee.Geometry.Rectangle([min_lon, min_lat, max_lon, max_lat])
        start_date = fecha_inicio.strftime('%Y-%m-%d')
        end_date = fecha_fin.strftime('%Y-%m-%d')
        
        if 'LC08' in dataset or 'LANDSAT/LC08' in dataset:
            red_band = 'SR_B4'
            nir_band = 'SR_B5'
            red_edge_band = 'SR_B6'
            blue_band = 'SR_B2'
        elif 'LC09' in dataset:
            red_band = 'SR_B4'
            nir_band = 'SR_B5'
            red_edge_band = 'SR_B6'
            blue_band = 'SR_B2'
        else:
            red_band = 'SR_B4'
            nir_band = 'SR_B5'
            red_edge_band = 'SR_B6'
            blue_band = 'SR_B2'
        
        collection = (ee.ImageCollection(dataset)
                     .filterBounds(geometry)
                     .filterDate(start_date, end_date)
                     .filter(ee.Filter.lt('CLOUD_COVER', 20)))
        
        image = collection.sort('CLOUD_COVER').first()
        if image is None:
            st.warning("⚠️ No se encontraron imágenes Landsat para el período y área seleccionados")
            return None
        
        if indice == 'NDVI':
            ndvi = image.normalizedDifference([nir_band, red_band]).rename('NDVI')
            index_image = ndvi
        elif indice == 'NDRE':
            ndre = image.normalizedDifference([nir_band, red_edge_band]).rename('NDRE')
            index_image = ndre
        elif indice == 'NDWI':
            ndwi = image.normalizedDifference(['SR_B3', nir_band]).rename('NDWI')
            index_image = ndwi
        elif indice == 'EVI':
            evi = image.expression(
                '2.5 * ((NIR - RED) / (NIR + 6 * RED - 7.5 * BLUE + 1))',
                {'NIR': image.select(nir_band), 'RED': image.select(red_band), 'BLUE': image.select(blue_band)}
            ).rename('EVI')
            index_image = evi
        elif indice == 'SAVI':
            savi = image.expression(
                '((NIR - RED) / (NIR + RED + 0.5)) * (1.5)',
                {'NIR': image.select(nir_band), 'RED': image.select(red_band)}
            ).rename('SAVI')
            index_image = savi
        elif indice == 'MSAVI':
            msavi = image.expression(
                '(2 * NIR + 1 - sqrt(pow((2 * NIR + 1), 2) - 8 * (NIR - RED))) / 2',
                {'NIR': image.select(nir_band), 'RED': image.select(red_band)}
            ).rename('MSAVI')
            index_image = msavi
        else:
            ndvi = image.normalizedDifference([nir_band, red_band]).rename('NDVI')
            index_image = ndvi
            indice = 'NDVI'
        
        stats = index_image.reduceRegion(
            reducer=ee.Reducer.mean().combine(
                reducer2=ee.Reducer.minMax(),
                sharedInputs=True
            ).combine(
                reducer2=ee.Reducer.stdDev(),
                sharedInputs=True
            ),
            geometry=geometry,
            scale=30,
            bestEffort=True
        )
        
        stats_dict = stats.getInfo()
        if not stats_dict:
            st.warning("⚠️ No se pudieron obtener estadísticas de la imagen")
            return None
        
        valor_promedio = stats_dict.get(f'{indice}_mean', 0)
        valor_min = stats_dict.get(f'{indice}_min', 0)
        valor_max = stats_dict.get(f'{indice}_max', 0)
        valor_std = stats_dict.get(f'{indice}_stdDev', 0)
        
        fecha_imagen_ee = image.get('system:time_start')
        fecha_imagen = fecha_imagen_ee.getInfo() if fecha_imagen_ee else None
        if fecha_imagen:
            fecha_imagen = datetime.fromtimestamp(fecha_imagen / 1000).strftime('%Y-%m-%d')
        
        if 'LC08' in dataset:
            nombre_satelite = 'Landsat 8'
        elif 'LC09' in dataset:
            nombre_satelite = 'Landsat 9'
        else:
            nombre_satelite = 'Landsat'
        
        cloud_cover_ee = image.get('CLOUD_COVER')
        cloud_cover = cloud_cover_ee.getInfo() if cloud_cover_ee else 'N/A'
        
        return {
            'indice': indice,
            'valor_promedio': valor_promedio,
            'valor_min': valor_min,
            'valor_max': valor_max,
            'valor_std': valor_std,
            'fuente': f'{nombre_satelite} (Google Earth Engine)',
            'fecha_descarga': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'fecha_imagen': fecha_imagen,
            'resolucion': '30m',
            'estado': 'exitosa',
            'cobertura_nubes': f"{cloud_cover}%" if cloud_cover != 'N/A' else 'N/A'
        }
        
    except Exception as e:
        st.error(f"❌ Error obteniendo datos de Landsat desde GEE: {str(e)}")
        return None

def descargar_datos_satelitales_gee(gdf, fecha_inicio, fecha_fin, satelite, indice='NDVI'):
    if satelite == 'SENTINEL-2_GEE':
        return obtener_datos_sentinel2_gee(gdf, fecha_inicio, fecha_fin, indice)
    elif satelite == 'LANDSAT-8_GEE':
        return obtener_datos_landsat_gee(gdf, fecha_inicio, fecha_fin, 'LANDSAT/LC08/C02/T1_L2', indice)
    elif satelite == 'LANDSAT-9_GEE':
        return obtener_datos_landsat_gee(gdf, fecha_inicio, fecha_fin, 'LANDSAT/LC09/C02/T1_L2', indice)
    else:
        return None

# ===== FUNCIÓN PARA OBTENER DATOS DE NASA POWER =====
def obtener_datos_nasa_power(gdf, fecha_inicio, fecha_fin):
    try:
        centroid = gdf.geometry.unary_union.centroid
        lat = round(centroid.y, 4)
        lon = round(centroid.x, 4)
        start = fecha_inicio.strftime("%Y%m%d")
        end = fecha_fin.strftime("%Y%m%d")
        params = {
            'parameters': 'ALLSKY_SFC_SW_DWN,WS2M,T2M,PRECTOTCORR',
            'community': 'RE',
            'longitude': lon,
            'latitude': lat,
            'start': start,
            'end': end,
            'format': 'JSON'
        }
        url = "https://power.larc.nasa.gov/api/temporal/daily/point"
        response = requests.get(url, params=params, timeout=15)
        data = response.json()
        if 'properties' not in data or 'parameter' not in data['properties']:
            return None
        series = data['properties']['parameter']
        df_power = pd.DataFrame({
            'fecha': pd.to_datetime(list(series['ALLSKY_SFC_SW_DWN'].keys())),
            'radiacion_solar': list(series['ALLSKY_SFC_SW_DWN'].values()),
            'viento_2m': list(series['WS2M'].values()),
            'temperatura': list(series['T2M'].values()),
            'precipitacion': list(series['PRECTOTCORR'].values())
        })
        df_power = df_power.replace(-999, np.nan).dropna()
        if df_power.empty:
            return None
        return df_power
    except Exception as e:
        return None

# ===== FUNCIONES DE ANÁLISIS COMPLETOS =====
def analizar_fertilidad_actual(gdf_dividido, cultivo, datos_satelitales):
    n_poligonos = len(gdf_dividido)
    resultados = []
    gdf_centroids = gdf_dividido.copy()
    gdf_centroids['centroid'] = gdf_dividido.geometry.centroid
    gdf_centroids['x'] = gdf_centroids.centroid.x
    gdf_centroids['y'] = gdf_centroids.centroid.y
    x_coords = gdf_centroids['x'].tolist()
    y_coords = gdf_centroids['y'].tolist()
    x_min, x_max = min(x_coords), max(x_coords)
    y_min, y_max = min(y_coords), max(y_coords)
    params = PARAMETROS_CULTIVOS[cultivo]
    valor_base_satelital = datos_satelitales.get('valor_promedio', 0.6) if datos_satelitales else 0.6
    for idx, row in gdf_centroids.iterrows():
        x_norm = (row['x'] - x_min) / (x_max - x_min) if x_max != x_min else 0.5
        y_norm = (row['y'] - y_min) / (y_max - y_min) if y_max != y_min else 0.5
        patron_espacial = (x_norm * 0.6 + y_norm * 0.4)
        
        base_mo = params['MATERIA_ORGANICA_OPTIMA'] * 0.7
        variabilidad_mo = patron_espacial * (params['MATERIA_ORGANICA_OPTIMA'] * 0.6)
        materia_organica = base_mo + variabilidad_mo + np.random.normal(0, 0.2)
        materia_organica = max(0.5, min(8.0, materia_organica))
        
        base_humedad = params['HUMEDAD_OPTIMA'] * 0.8
        variabilidad_humedad = patron_espacial * (params['HUMEDAD_OPTIMA'] * 0.4)
        humedad_suelo = base_humedad + variabilidad_humedad + np.random.normal(0, 0.05)
        humedad_suelo = max(0.1, min(0.8, humedad_suelo))
        
        ndvi_base = valor_base_satelital * 0.8
        ndvi_variacion = patron_espacial * (valor_base_satelital * 0.4)
        ndvi = ndvi_base + ndvi_variacion + np.random.normal(0, 0.06)
        ndvi = max(0.1, min(0.9, ndvi))
        
        ndre_base = params['NDRE_OPTIMO'] * 0.7
        ndre_variacion = patron_espacial * (params['NDRE_OPTIMO'] * 0.4)
        ndre = ndre_base + ndre_variacion + np.random.normal(0, 0.04)
        ndre = max(0.05, min(0.7, ndre))
        
        ndwi = 0.2 + np.random.normal(0, 0.08)
        ndwi = max(0, min(1, ndwi))
        
        npk_actual = (ndvi * 0.4) + (ndre * 0.3) + ((materia_organica / 8) * 0.2) + (humedad_suelo * 0.1)
        npk_actual = max(0, min(1, npk_actual))
        
        resultados.append({
            'materia_organica': round(materia_organica, 2),
            'humedad_suelo': round(humedad_suelo, 3),
            'ndvi': round(ndvi, 3),
            'ndre': round(ndre, 3),
            'ndwi': round(ndwi, 3),
            'npk_actual': round(npk_actual, 3)
        })

    return resultados

def analizar_recomendaciones_npk(indices, cultivo):
    recomendaciones_n = []
    recomendaciones_p = []
    recomendaciones_k = []
    params = PARAMETROS_CULTIVOS[cultivo]

    for idx in indices:
        ndre = idx['ndre']
        materia_organica = idx['materia_organica']
        humedad_suelo = idx['humedad_suelo']
        ndvi = idx['ndvi']
        
        factor_n = ((1 - ndre) * 0.6 + (1 - ndvi) * 0.4)
        n_recomendado = (factor_n * (params['NITROGENO']['max'] - params['NITROGENO']['min']) + params['NITROGENO']['min'])
        n_recomendado = max(params['NITROGENO']['min'] * 0.8, min(params['NITROGENO']['max'] * 1.2, n_recomendado))
        recomendaciones_n.append(round(n_recomendado, 1))
        
        factor_p = ((1 - (materia_organica / 8)) * 0.7 + (1 - humedad_suelo) * 0.3)
        p_recomendado = (factor_p * (params['FOSFORO']['max'] - params['FOSFORO']['min']) + params['FOSFORO']['min'])
        p_recomendado = max(params['FOSFORO']['min'] * 0.8, min(params['FOSFORO']['max'] * 1.2, p_recomendado))
        recomendaciones_p.append(round(p_recomendado, 1))
        
        factor_k = ((1 - ndre) * 0.4 + (1 - humedad_suelo) * 0.4 + (1 - (materia_organica / 8)) * 0.2)
        k_recomendado = (factor_k * (params['POTASIO']['max'] - params['POTASIO']['min']) + params['POTASIO']['min'])
        k_recomendado = max(params['POTASIO']['min'] * 0.8, min(params['POTASIO']['max'] * 1.2, k_recomendado))
        recomendaciones_k.append(round(k_recomendado, 1))

    return recomendaciones_n, recomendaciones_p, recomendaciones_k

def analizar_costos(gdf_dividido, cultivo, recomendaciones_n, recomendaciones_p, recomendaciones_k):
    costos = []
    params = PARAMETROS_CULTIVOS[cultivo]
    precio_n = 1.2
    precio_p = 2.5
    precio_k = 1.8

    for i in range(len(gdf_dividido)):
        costo_n = recomendaciones_n[i] * precio_n
        costo_p = recomendaciones_p[i] * precio_p
        costo_k = recomendaciones_k[i] * precio_k
        costo_total = costo_n + costo_p + costo_k + params['COSTO_FERTILIZACION']
        
        costos.append({
            'costo_nitrogeno': round(costo_n, 2),
            'costo_fosforo': round(costo_p, 2),
            'costo_potasio': round(costo_k, 2),
            'costo_total': round(costo_total, 2)
        })

    return costos

def analizar_proyecciones_cosecha(gdf_dividido, cultivo, indices):
    proyecciones = []
    params = PARAMETROS_CULTIVOS[cultivo]
    for idx in indices:
        npk_actual = idx['npk_actual']
        ndvi = idx['ndvi']
        
        rendimiento_base = params['RENDIMIENTO_OPTIMO'] * npk_actual * 0.7
        incremento = (1 - npk_actual) * 0.4 + (1 - ndvi) * 0.2
        rendimiento_con_fert = rendimiento_base * (1 + incremento)
        
        proyecciones.append({
            'rendimiento_sin_fert': round(rendimiento_base, 0),
            'rendimiento_con_fert': round(rendimiento_con_fert, 0),
            'incremento_esperado': round(incremento * 100, 1)
        })

    return proyecciones

def clasificar_textura_suelo(arena, limo, arcilla):
    try:
        total = arena + limo + arcilla
        if total == 0:
            return "NO_DETERMINADA"
        arena_norm = (arena / total) * 100
        limo_norm = (limo / total) * 100
        arcilla_norm = (arcilla / total) * 100
        if arcilla_norm >= 35:
            return "Franco arcilloso"
        elif arcilla_norm >= 25 and arcilla_norm <= 35 and arena_norm >= 20 and arena_norm <= 45:
            return "Franco arcilloso"
        elif arena_norm >= 55 and arena_norm <= 70 and arcilla_norm >= 10 and arcilla_norm <= 20:
            return "Franco arenoso"
        elif arena_norm >= 40 and arena_norm <= 55 and arcilla_norm >= 20 and arcilla_norm <= 30:
            return "Franco"
        else:
            return "Franco"
    except Exception as e:
        return "NO_DETERMINADA"

def analizar_textura_suelo(gdf_dividido, cultivo):
    gdf_dividido = validar_y_corregir_crs(gdf_dividido)
    params_textura = TEXTURA_SUELO_OPTIMA[cultivo]
    gdf_dividido['area_ha'] = 0.0
    gdf_dividido['arena'] = 0.0
    gdf_dividido['limo'] = 0.0
    gdf_dividido['arcilla'] = 0.0
    gdf_dividido['textura_suelo'] = "NO_DETERMINADA"

    for idx, row in gdf_dividido.iterrows():
        try:
            area_gdf = gpd.GeoDataFrame({'geometry': [row.geometry]}, crs=gdf_dividido.crs)
            area_ha = calcular_superficie(area_gdf)
            if hasattr(area_ha, 'iloc'):
                area_ha = float(area_ha.iloc[0])
            elif hasattr(area_ha, '__len__') and len(area_ha) > 0:
                area_ha = float(area_ha[0])
            else:
                area_ha = float(area_ha)
            
            centroid = row.geometry.centroid if hasattr(row.geometry, 'centroid') else row.geometry.representative_point()
            seed_value = abs(hash(f"{centroid.x:.6f}_{centroid.y:.6f}_{cultivo}_textura")) % (2**32)
            rng = np.random.RandomState(seed_value)
            
            lat_norm = (centroid.y + 90) / 180 if centroid.y else 0.5
            lon_norm = (centroid.x + 180) / 360 if centroid.x else 0.5
            variabilidad_local = 0.15 + 0.7 * (lat_norm * lon_norm)
            
            arena_optima = params_textura['arena_optima']
            limo_optima = params_textura['limo_optima']
            arcilla_optima = params_textura['arcilla_optima']
            
            arena_val = max(5, min(95, rng.normal(
                arena_optima * (0.8 + 0.4 * variabilidad_local),
                arena_optima * 0.15
            )))
            limo_val = max(5, min(95, rng.normal(
                limo_optima * (0.7 + 0.6 * variabilidad_local),
                limo_optima * 0.2
            )))
            arcilla_val = max(5, min(95, rng.normal(
                arcilla_optima * (0.75 + 0.5 * variabilidad_local),
                arcilla_optima * 0.15
            )))
            
            total = arena_val + limo_val + arcilla_val
            arena_pct = (arena_val / total) * 100
            limo_pct = (limo_val / total) * 100
            arcilla_pct = (arcilla_val / total) * 100
            
            textura = clasificar_textura_suelo(arena_pct, limo_pct, arcilla_pct)
            
            gdf_dividido.at[idx, 'area_ha'] = area_ha
            gdf_dividido.at[idx, 'arena'] = float(arena_pct)
            gdf_dividido.at[idx, 'limo'] = float(limo_pct)
            gdf_dividido.at[idx, 'arcilla'] = float(arcilla_pct)
            gdf_dividido.at[idx, 'textura_suelo'] = textura
            
        except Exception as e:
            gdf_dividido.at[idx, 'area_ha'] = 0.0
            gdf_dividido.at[idx, 'arena'] = float(params_textura['arena_optima'])
            gdf_dividido.at[idx, 'limo'] = float(params_textura['limo_optima'])
            gdf_dividido.at[idx, 'arcilla'] = float(params_textura['arcilla_optima'])
            gdf_dividido.at[idx, 'textura_suelo'] = params_textura['textura_optima']

    return gdf_dividido

# ===== FUNCIÓN PARA EJECUTAR TODOS LOS ANÁLISIS =====
def ejecutar_analisis_completo(gdf, cultivo, n_divisiones, satelite, fecha_inicio, fecha_fin,
                               intervalo_curvas=5.0, resolucion_dem=10.0):
    resultados = {
        'exitoso': False,
        'gdf_dividido': None,
        'fertilidad_actual': None,
        'recomendaciones_npk': None,
        'costos': None,
        'proyecciones': None,
        'textura': None,
        'df_power': None,
        'area_total': 0,
        'mapas': {},
        'dem_data': None,
        'curvas_nivel': None,
        'pendientes': None,
        'datos_satelitales': None
    }

    try:
        gdf = validar_y_corregir_crs(gdf)
        area_total = calcular_superficie(gdf)
        resultados['area_total'] = area_total
        
        datos_satelitales = None
        if satelite in ['SENTINEL-2_GEE', 'LANDSAT-8_GEE', 'LANDSAT-9_GEE']:
            datos_satelitales = descargar_datos_satelitales_gee(gdf, fecha_inicio, fecha_fin, satelite, indice_seleccionado)
            if datos_satelitales is None:
                st.warning("⚠️ No se pudieron obtener datos de GEE. Usando datos simulados.")
                datos_satelitales = generar_datos_simulados(gdf, cultivo, indice_seleccionado)
        elif satelite == "SENTINEL-2":
            datos_satelitales = descargar_datos_sentinel2(gdf, fecha_inicio, fecha_fin, indice_seleccionado)
        elif satelite == "LANDSAT-8":
            datos_satelitales = descargar_datos_landsat8(gdf, fecha_inicio, fecha_fin, indice_seleccionado)
        else:
            datos_satelitales = generar_datos_simulados(gdf, cultivo, indice_seleccionado)
        
        resultados['datos_satelitales'] = datos_satelitales
        
        df_power = obtener_datos_nasa_power(gdf, fecha_inicio, fecha_fin)
        resultados['df_power'] = df_power
        
        gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
        resultados['gdf_dividido'] = gdf_dividido
        
        areas_ha_list = []
        for idx, row in gdf_dividido.iterrows():
            area_gdf = gpd.GeoDataFrame({'geometry': [row.geometry]}, crs=gdf_dividido.crs)
            area_ha = calcular_superficie(area_gdf)
            if hasattr(area_ha, 'iloc'):
                area_ha = float(area_ha.iloc[0])
            elif hasattr(area_ha, '__len__') and len(area_ha) > 0:
                area_ha = float(area_ha[0])
            else:
                area_ha = float(area_ha)
            areas_ha_list.append(area_ha)
        
        gdf_dividido['area_ha'] = areas_ha_list
        
        fertilidad_actual = analizar_fertilidad_actual(gdf_dividido, cultivo, datos_satelitales)
        resultados['fertilidad_actual'] = fertilidad_actual
        
        rec_n, rec_p, rec_k = analizar_recomendaciones_npk(fertilidad_actual, cultivo)
        resultados['recomendaciones_npk'] = {'N': rec_n, 'P': rec_p, 'K': rec_k}
        
        costos = analizar_costos(gdf_dividido, cultivo, rec_n, rec_p, rec_k)
        resultados['costos'] = costos
        
        proyecciones = analizar_proyecciones_cosecha(gdf_dividido, cultivo, fertilidad_actual)
        resultados['proyecciones'] = proyecciones
        
        textura = analizar_textura_suelo(gdf_dividido, cultivo)
        resultados['textura'] = textura

        # ----- 6. Análisis DEM y curvas de nivel (PRIORIDAD: REAL > OPENTOPODATA > SINTÉTICO) -----
        try:
            api_key = os.environ.get("OPENTOPOGRAPHY_API_KEY", None)
            dem_array, dem_meta, dem_transform = obtener_dem_opentopography(gdf, api_key)

            # Si falla OpenTopography, intentar con Open Topo Data API
            if dem_array is None:
                st.info("ℹ️ Intentando con fuente alternativa: Open Topo Data API (srtm30m)")
                dem_array, dem_meta, dem_transform = obtener_dem_opentopodata_api(gdf, dataset="srtm30m")

            dem_data = {
                'X': None, 'Y': None, 'Z': None,
                'bounds': None,
                'curvas_nivel': [], 'elevaciones': [],
                'curvas_con_elevacion': [], 'pendientes': None,
                'fuente': 'No disponible'
            }

            if dem_array is not None and not (isinstance(dem_array, np.ma.MaskedArray) and dem_array.mask.all()):
                # Determinar la fuente real
                if dem_transform is not None:
                    # Caso OpenTopography (con transform)
                    st.info("✅ Usando DEM real SRTM 30m (OpenTopography)")
                    dem_data['fuente'] = 'SRTM 30m'

                    height, width = dem_array.shape
                    cols = np.arange(width)
                    rows = np.arange(height)
                    X_grid, Y_grid = np.meshgrid(cols, rows)
                    X_geo = dem_transform[2] + dem_transform[0] * X_grid + dem_transform[1] * Y_grid
                    Y_geo = dem_transform[5] + dem_transform[3] * X_grid + dem_transform[4] * Y_grid

                    # Convertir a float antes de rellenar con NaN
                    if isinstance(dem_array, np.ma.MaskedArray):
                        Z = dem_array.astype(float).filled(np.nan)
                    else:
                        Z = dem_array.astype(float)
                        Z[Z <= -32768] = np.nan

                    dem_data.update({
                        'X': X_geo, 'Y': Y_geo, 'Z': Z,
                        'bounds': gdf.total_bounds
                    })

                    if CURVAS_OK:
                        polygon_union = gdf.geometry.unary_union
                        curvas_con_elev = generar_curvas_nivel_reales(dem_array, dem_transform, intervalo_curvas, polygon=polygon_union)
                        if curvas_con_elev:
                            dem_data['curvas_con_elevacion'] = curvas_con_elev
                            dem_data['curvas_nivel'] = [line for line, _ in curvas_con_elev]
                            dem_data['elevaciones'] = [e for _, e in curvas_con_elev]
                            st.success(f"✅ Generadas {len(curvas_con_elev)} curvas de nivel reales.")

                else:
                    # Caso Open Topo Data (sin transform, tenemos X, Y, Z en la máscara)
                    st.info("✅ Usando DEM de Open Topo Data")
                    dem_data['fuente'] = 'Open Topo Data'
                    
                    # Extraer X, Y, Z desde el array enmascarado
                    # Para Open Topo Data, hemos guardado X, Y en el array? No, necesitamos reconstruirlos.
                    # Vamos a reconstruir la malla a partir de los bounds y la forma del array.
                    height, width = dem_array.shape
                    bounds = gdf.total_bounds
                    minx, miny, maxx, maxy = bounds
                    x_vals = np.linspace(minx, maxx, width)
                    y_vals = np.linspace(miny, maxy, height)
                    X_geo, Y_geo = np.meshgrid(x_vals, y_vals)

                    if isinstance(dem_array, np.ma.MaskedArray):
                        Z = dem_array.astype(float).filled(np.nan)
                    else:
                        Z = dem_array.astype(float)

                    dem_data.update({
                        'X': X_geo, 'Y': Y_geo, 'Z': Z,
                        'bounds': bounds
                    })

                    if CURVAS_OK:
                        # Reutilizar la función de curvas sintéticas pero con nuestros datos reales
                        # Como ya tenemos X,Y,Z podemos usar una función genérica que extraiga contornos
                        # Por simplicidad, usamos generar_curvas_nivel_simuladas pero adaptada
                        # Para no duplicar código, podemos modificar generar_curvas_nivel_simuladas para aceptar X,Y,Z opcionales.
                        # Sin embargo, para mantener el código manejable, usaremos una versión simplificada aquí.
                        # Llamamos a generar_curvas_nivel_simuladas con el gdf, pero eso generaría un nuevo DEM sintético.
                        # Mejor implementamos una función auxiliar que extraiga contornos de (X,Y,Z).
                        curvas_con_elev = extraer_curvas_de_grid(X_geo, Y_geo, Z, intervalo_curvas, gdf.geometry.unary_union)
                        if curvas_con_elev:
                            dem_data['curvas_con_elevacion'] = curvas_con_elev
                            dem_data['curvas_nivel'] = [line for line, _ in curvas_con_elev]
                            dem_data['elevaciones'] = [e for _, e in curvas_con_elev]
                            st.success(f"✅ Generadas {len(curvas_con_elev)} curvas de nivel desde Open Topo Data.")

            else:
                st.info("ℹ️ Usando DEM sintético (fuentes externas no disponibles)")
                dem_data['fuente'] = 'Sintético'
                X, Y, Z, bounds = generar_dem_sintetico_fallback(gdf, resolucion_dem)
                dem_data.update({'X': X, 'Y': Y, 'Z': Z, 'bounds': bounds})

                if CURVAS_OK:
                    curvas_con_elev = generar_curvas_nivel_simuladas(gdf, intervalo_curvas)
                    if curvas_con_elev:
                        dem_data['curvas_con_elevacion'] = curvas_con_elev
                        dem_data['curvas_nivel'] = [line for line, _ in curvas_con_elev]
                        dem_data['elevaciones'] = [e for _, e in curvas_con_elev]

            # Calcular pendientes (si hay datos válidos)
            if dem_data['Z'] is not None and not np.all(np.isnan(dem_data['Z'])):
                Z_grid = dem_data['Z'].astype(float)
                mask_valid = ~np.isnan(Z_grid)

                if np.any(mask_valid):
                    # Obtener resolución espacial en grados
                    if dem_data['fuente'] == 'SRTM 30m' and dem_transform is not None:
                        res_x_deg = abs(dem_transform[0])
                        res_y_deg = abs(dem_transform[4])
                        lat_media = np.nanmean(dem_data['Y'][mask_valid])
                        res_x_m = res_x_deg * 111320 * np.cos(np.radians(lat_media))
                        res_y_m = res_y_deg * 111320
                    else:
                        # Para Open Topo Data o sintético, calcular desde la malla
                        X = dem_data['X']
                        Y = dem_data['Y']
                        dx_deg = X[0,1] - X[0,0]
                        dy_deg = Y[1,0] - Y[0,0]
                        lat_media = np.nanmean(Y[mask_valid])
                        res_x_m = abs(dx_deg) * 111320 * np.cos(np.radians(lat_media))
                        res_y_m = abs(dy_deg) * 111320

                    dy = np.gradient(Z_grid, axis=0) / res_y_m
                    dx = np.gradient(Z_grid, axis=1) / res_x_m
                    pendientes = np.sqrt(dx**2 + dy**2) * 100
                    pendientes[~mask_valid] = np.nan
                    dem_data['pendientes'] = pendientes
                else:
                    dem_data['pendientes'] = None

            resultados['dem_data'] = dem_data

        except Exception as e:
            st.error(f"❌ Error crítico en análisis DEM: {str(e)[:100]}")
            resultados['dem_data'] = None

        # ===== COMBINAR TODOS LOS RESULTADOS EN UN SOLO GeoDataFrame =====
        gdf_completo = gdf_dividido.copy()
        # Fertilidad
        for i, f in enumerate(fertilidad_actual):
            gdf_completo.loc[i, 'fert_npk_actual'] = f['npk_actual']
            gdf_completo.loc[i, 'fert_ndvi'] = f['ndvi']
            gdf_completo.loc[i, 'fert_ndre'] = f['ndre']
            gdf_completo.loc[i, 'fert_ndwi'] = f['ndwi']
            gdf_completo.loc[i, 'fert_materia_organica'] = f['materia_organica']
            gdf_completo.loc[i, 'fert_humedad_suelo'] = f['humedad_suelo']
        # Recomendaciones
        gdf_completo['rec_N'] = rec_n
        gdf_completo['rec_P'] = rec_p
        gdf_completo['rec_K'] = rec_k
        # Costos
        for i, c in enumerate(costos):
            gdf_completo.loc[i, 'costo_costo_nitrogeno'] = c['costo_nitrogeno']
            gdf_completo.loc[i, 'costo_costo_fosforo'] = c['costo_fosforo']
            gdf_completo.loc[i, 'costo_costo_potasio'] = c['costo_potasio']
            gdf_completo.loc[i, 'costo_costo_total'] = c['costo_total']
        # Proyecciones
        for i, p in enumerate(proyecciones):
            gdf_completo.loc[i, 'proy_rendimiento_sin_fert'] = p['rendimiento_sin_fert']
            gdf_completo.loc[i, 'proy_rendimiento_con_fert'] = p['rendimiento_con_fert']
            gdf_completo.loc[i, 'proy_incremento_esperado'] = p['incremento_esperado']
        # Textura (ya está en gdf_dividido, se copia automáticamente)
        resultados['gdf_completo'] = gdf_completo

        resultados['exitoso'] = True
        return resultados

    except Exception as e:
        st.error(f"❌ Error en el análisis completo: {str(e)}")
        import traceback
        traceback.print_exc()
        resultados['exitoso'] = False
        return resultados

# Función auxiliar para extraer curvas de nivel de una grilla regular (X,Y,Z)
def extraer_curvas_de_grid(X, Y, Z, intervalo, polygon=None):
    """
    Extrae curvas de nivel de una grilla regular definida por X, Y, Z.
    X, Y son matrices de coordenadas, Z es matriz de elevaciones (con NaN).
    """
    if not SKIMAGE_OK:
        return []
    from skimage import measure

    Z_filled = np.where(np.isnan(Z), -9999, Z)
    niveles = np.arange(np.nanmin(Z), np.nanmax(Z) + intervalo, intervalo)
    if len(niveles) < 2:
        return []

    ny, nx = Z.shape
    contours = []
    for nivel in niveles:
        try:
            for contour in measure.find_contours(Z_filled, nivel):
                coords = []
                for row, col in contour:
                    r, c = int(round(row)), int(round(col))
                    if r < 0 or r >= ny or c < 0 or c >= nx or np.isnan(Z[r, c]):
                        continue
                    # Interpolar coordenadas (podría ser más preciso, pero aproximado)
                    lon = X[r, c]
                    lat = Y[r, c]
                    coords.append((lon, lat))
                if len(coords) >= 3:
                    line = LineString(coords)
                    if line.length > 0.01 and (polygon is None or line.intersects(polygon)):
                        contours.append((line, nivel))
        except Exception:
            continue
    if contours:
        st.info(f"✅ Generadas {len(contours)} curvas de nivel desde grilla")
    else:
        st.warning("⚠️ No se generaron curvas de nivel desde la grilla.")
    return contours

# ===== FUNCIONES DE VISUALIZACIÓN CON BOTONES DESCARGA =====
def crear_mapa_fertilidad(gdf_completo, cultivo, satelite):
    try:
        gdf_plot = gdf_completo.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        cmap = LinearSegmentedColormap.from_list('fertilidad_gee', PALETAS_GEE['FERTILIDAD'])
        vmin, vmax = 0, 1
        
        for idx, row in gdf_plot.iterrows():
            valor = row['fert_npk_actual']
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = cmap(valor_norm)
            
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='black', linewidth=1.5, alpha=0.7)
            
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.2f}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='black', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9))
        
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.7)
        except:
            pass
        
        info_satelite = SATELITES_DISPONIBLES.get(satelite, SATELITES_DISPONIBLES['DATOS_SIMULADOS'])
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} FERTILIDAD ACTUAL - {cultivo}\n'
                     f'{info_satelite["icono"]} {info_satelite["nombre"]}',
                     fontsize=16, fontweight='bold', pad=20)
        ax.set_xlabel('Longitud')
        ax.set_ylabel('Latitud')
        ax.grid(True, alpha=0.3)
        
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label('Índice de Fertilidad', fontsize=12, fontweight='bold')
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando mapa de fertilidad: {str(e)}")
        return None

def crear_mapa_npk(gdf_completo, cultivo, nutriente='N'):
    try:
        gdf_plot = gdf_completo.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        if nutriente == 'N':
            cmap = LinearSegmentedColormap.from_list('nitrogeno_gee', PALETAS_GEE['NITROGENO'])
            columna = 'rec_N'
            titulo_nut = 'NITRÓGENO'
            vmin = PARAMETROS_CULTIVOS[cultivo]['NITROGENO']['min'] * 0.8
            vmax = PARAMETROS_CULTIVOS[cultivo]['NITROGENO']['max'] * 1.2
        elif nutriente == 'P':
            cmap = LinearSegmentedColormap.from_list('fosforo_gee', PALETAS_GEE['FOSFORO'])
            columna = 'rec_P'
            titulo_nut = 'FÓSFORO'
            vmin = PARAMETROS_CULTIVOS[cultivo]['FOSFORO']['min'] * 0.8
            vmax = PARAMETROS_CULTIVOS[cultivo]['FOSFORO']['max'] * 1.2
        else:
            cmap = LinearSegmentedColormap.from_list('potasio_gee', PALETAS_GEE['POTASIO'])
            columna = 'rec_K'
            titulo_nut = 'POTASIO'
            vmin = PARAMETROS_CULTIVOS[cultivo]['POTASIO']['min'] * 0.8
            vmax = PARAMETROS_CULTIVOS[cultivo]['POTASIO']['max'] * 1.2
        
        for idx, row in gdf_plot.iterrows():
            valor = row[columna]
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = cmap(valor_norm)
            
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='black', linewidth=1.5, alpha=0.7)
            
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.0f}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='black', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9))
        
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.7)
        except:
            pass
        
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} RECOMENDACIONES {titulo_nut} - {cultivo}',
                     fontsize=16, fontweight='bold', pad=20)
        ax.set_xlabel('Longitud')
        ax.set_ylabel('Latitud')
        ax.grid(True, alpha=0.3)
        
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label(f'{titulo_nut} (kg/ha)', fontsize=12, fontweight='bold')
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando mapa NPK: {str(e)}")
        return None

def crear_mapa_texturas(gdf_completo, cultivo):
    try:
        gdf_plot = gdf_completo.to_crs(epsg=3857)
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        colores_textura = {
            'Franco': '#c7eae5',
            'Franco arcilloso': '#5ab4ac',
            'Franco arenoso': '#f6e8c3',
            'NO_DETERMINADA': '#999999'
        }
        
        for idx, row in gdf_plot.iterrows():
            textura = row['textura_suelo']
            color = colores_textura.get(textura, '#999999')
            
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='black', linewidth=1.5, alpha=0.8)
            
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{textura[:10]}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='black', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9))
        
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.6)
        except:
            pass
        
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} MAPA DE TEXTURAS - {cultivo}',
                     fontsize=16, fontweight='bold', pad=20)
        ax.set_xlabel('Longitud')
        ax.set_ylabel('Latitud')
        ax.grid(True, alpha=0.3)
        
        from matplotlib.patches import Patch
        legend_elements = [Patch(facecolor=color, edgecolor='black', label=textura)
                           for textura, color in colores_textura.items()]
        ax.legend(handles=legend_elements, title='Texturas', loc='upper left', bbox_to_anchor=(1.05, 1))
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando mapa de texturas: {str(e)}")
        return None

def crear_grafico_distribucion_costos(costos_n, costos_p, costos_k, otros, costo_total):
    try:
        fig, ax = plt.subplots(figsize=(10, 6))
        categorias = ['Nitrógeno', 'Fósforo', 'Potasio', 'Otros']
        valores = [costos_n, costos_p, costos_k, otros]
        colores = ['#00ff00', '#0000ff', '#4B0082', '#cccccc']
        
        bars = ax.bar(categorias, valores, color=colores, edgecolor='black')
        ax.set_title('Distribución de Costos de Fertilización', fontsize=14, fontweight='bold')
        ax.set_ylabel('USD', fontsize=12)
        ax.set_xlabel('Componente', fontsize=12)
        
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 10,
                   f'${height:.0f}', ha='center', va='bottom', fontweight='bold')
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando gráfico de costos: {str(e)}")
        return None

def crear_grafico_composicion_textura(arena_prom, limo_prom, arcilla_prom, textura_dist):
    try:
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
        composicion = [arena_prom, limo_prom, arcilla_prom]
        labels = ['Arena', 'Limo', 'Arcilla']
        colors_pie = ['#d8b365', '#f6e8c3', '#01665e']
        ax1.pie(composicion, labels=labels, colors=colors_pie, autopct='%1.1f%%', startangle=90)
        ax1.set_title('Composición Promedio del Suelo')
        
        ax2.bar(textura_dist.index, textura_dist.values, 
               color=[PALETAS_GEE['TEXTURA'][i % len(PALETAS_GEE['TEXTURA'])] for i in range(len(textura_dist))])
        ax2.set_title('Distribución de Texturas')
        ax2.set_xlabel('Textura')
        ax2.set_ylabel('Número de Zonas')
        ax2.tick_params(axis='x', rotation=45)
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando gráfico de textura: {str(e)}")
        return None

def crear_grafico_proyecciones_rendimiento(zonas, sin_fert, con_fert):
    try:
        fig, ax = plt.subplots(figsize=(14, 7))
        x = np.arange(len(zonas))
        width = 0.35
        
        bars1 = ax.bar(x - width/2, sin_fert, width, label='Sin Fertilización', 
                      color='#ff9999', edgecolor='darkred', linewidth=1)
        bars2 = ax.bar(x + width/2, con_fert, width, label='Con Fertilización', 
                      color='#66b3ff', edgecolor='darkblue', linewidth=1)
        
        ax.set_xlabel('Zona', fontsize=12)
        ax.set_ylabel('Rendimiento (kg/ha)', fontsize=12)
        ax.set_title('PROYECCIONES DE RENDIMIENTO POR ZONA', fontsize=14, fontweight='bold', pad=20)
        ax.set_xticks(x)
        ax.set_xticklabels(zonas, rotation=45, ha='right')
        ax.legend()
        
        incrementos = [(c-s)/s*100 if s>0 else 0 for s,c in zip(sin_fert, con_fert)]
        
        for i, (bar1, bar2) in enumerate(zip(bars1, bars2)):
            height1 = bar1.get_height()
            height2 = bar2.get_height()
            ax.text(bar1.get_x() + bar1.get_width()/2., height1 + max(sin_fert)*0.01,
                   f'{height1:.0f}', ha='center', va='bottom', fontsize=8, rotation=90)
            ax.text(bar2.get_x() + bar2.get_width()/2., height2 + max(con_fert)*0.01,
                   f'{height2:.0f}', ha='center', va='bottom', fontsize=8, rotation=90)
            if incrementos[i] > 0:
                ax.text(bar2.get_x() + bar2.get_width()/2., height2 * 1.05,
                       f'+{incrementos[i]:.1f}%', ha='center', va='bottom', 
                       fontsize=7, color='green', weight='bold')
        
        if len(zonas) > 1:
            z = np.polyfit(x, sin_fert, 1)
            p = np.poly1d(z)
            ax.plot(x, p(x), "r--", alpha=0.5, label='Tendencia Base')
            z2 = np.polyfit(x, con_fert, 1)
            p2 = np.poly1d(z2)
            ax.plot(x, p2(x), "b--", alpha=0.5, label='Tendencia Mejorada')
        
        stats_text = f"""
        Resumen:
        • Total base: {sum(sin_fert):.0f} kg
        • Total mejorado: {sum(con_fert):.0f} kg
        • Incremento total: {sum(con_fert)-sum(sin_fert):.0f} kg
        • Incremento promedio: {np.mean(incrementos):.1f}%
        """
        
        ax.text(0.02, 0.98, stats_text, transform=ax.transAxes, fontsize=10,
                verticalalignment='top',
                bbox=dict(boxstyle="round,pad=0.5", facecolor='lightyellow', alpha=0.9))
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"❌ Error creando gráfico de proyecciones: {str(e)}")
        return None

# ===== FUNCIONES DE EXPORTACIÓN =====
def exportar_a_geojson(gdf, nombre_base="parcela"):
    try:
        gdf = validar_y_corregir_crs(gdf)
        geojson_data = gdf.to_json()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"{nombre_base}_{timestamp}.geojson"
        return geojson_data, nombre_archivo
    except Exception as e:
        st.error(f"❌ Error exportando a GeoJSON: {str(e)}")
        return None, None

def generar_reporte_completo(resultados, cultivo, satelite, fecha_inicio, fecha_fin,
                             resolucion_dem, intervalo_curvas):
    """
    Genera un informe DOCX sin IA (solo datos y tablas).
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import numpy as np
        from datetime import datetime
        import io

        doc = Document()
        
        # ===== PORTADA =====
        title = doc.add_heading(f'REPORTE DE AMBIENTACIÓN AGRONÓMICA - {cultivo}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # ===== 1. INFORMACIÓN GENERAL =====
        doc.add_heading('1. INFORMACIÓN GENERAL', level=1)
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = 'Cultivo'; info_table.cell(0, 1).text = cultivo
        info_table.cell(1, 0).text = 'Área Total'; info_table.cell(1, 1).text = f'{resultados["area_total"]:.2f} ha'
        info_table.cell(2, 0).text = 'Zonas Analizadas'; info_table.cell(2, 1).text = str(len(resultados['gdf_completo']))
        info_table.cell(3, 0).text = 'Satélite'; info_table.cell(3, 1).text = satelite
        info_table.cell(4, 0).text = 'Período'; info_table.cell(4, 1).text = f'{fecha_inicio.strftime("%d/%m/%Y")} a {fecha_fin.strftime("%d/%m/%Y")}'
        info_table.cell(5, 0).text = 'Fuente Datos'; info_table.cell(5, 1).text = resultados['datos_satelitales']['fuente'] if resultados['datos_satelitales'] else 'N/A'

        # ===== 2. FERTILIDAD ACTUAL =====
        doc.add_heading('2. FERTILIDAD ACTUAL', level=1)
        doc.add_paragraph('**Resumen de parámetros de fertilidad por zona:**')
        fert_table = doc.add_table(rows=1, cols=7)
        fert_table.style = 'Table Grid'
        headers = ['Zona', 'Área (ha)', 'Índice NPK', 'NDVI', 'NDRE', 'Materia Org (%)', 'Humedad']
        for i, header in enumerate(headers): fert_table.cell(0, i).text = header
        for i in range(min(10, len(resultados['gdf_completo']))):
            row = fert_table.add_row().cells
            row[0].text = str(resultados['gdf_completo'].iloc[i]['id_zona'])
            row[1].text = f"{resultados['gdf_completo'].iloc[i]['area_ha']:.2f}"
            row[2].text = f"{resultados['gdf_completo'].iloc[i]['fert_npk_actual']:.3f}"
            row[3].text = f"{resultados['gdf_completo'].iloc[i]['fert_ndvi']:.3f}"
            row[4].text = f"{resultados['gdf_completo'].iloc[i]['fert_ndre']:.3f}"
            row[5].text = f"{resultados['gdf_completo'].iloc[i]['fert_materia_organica']:.1f}"
            row[6].text = f"{resultados['gdf_completo'].iloc[i]['fert_humedad_suelo']:.3f}"
        doc.add_paragraph()

        # ===== 3. RECOMENDACIONES NPK =====
        doc.add_heading('3. RECOMENDACIONES NPK', level=1)
        doc.add_paragraph('**Recomendaciones de fertilización por zona (kg/ha):**')
        npk_table = doc.add_table(rows=1, cols=4)
        npk_table.style = 'Table Grid'
        npk_headers = ['Zona', 'Nitrógeno (N)', 'Fósforo (P)', 'Potasio (K)']
        for i, header in enumerate(npk_headers): npk_table.cell(0, i).text = header
        for i in range(min(10, len(resultados['gdf_completo']))):
            row = npk_table.add_row().cells
            row[0].text = str(resultados['gdf_completo'].iloc[i]['id_zona'])
            row[1].text = f"{resultados['gdf_completo'].iloc[i]['rec_N']:.1f}"
            row[2].text = f"{resultados['gdf_completo'].iloc[i]['rec_P']:.1f}"
            row[3].text = f"{resultados['gdf_completo'].iloc[i]['rec_K']:.1f}"
        doc.add_paragraph()

        # ===== 4. ANÁLISIS DE COSTOS =====
        doc.add_heading('4. ANÁLISIS DE COSTOS', level=1)
        costo_table = doc.add_table(rows=1, cols=5)
        costo_table.style = 'Table Grid'
        costo_headers = ['Zona', 'Costo N', 'Costo P', 'Costo K', 'Costo Total']
        for i, header in enumerate(costo_headers): costo_table.cell(0, i).text = header
        for i in range(min(10, len(resultados['gdf_completo']))):
            row = costo_table.add_row().cells
            row[0].text = str(resultados['gdf_completo'].iloc[i]['id_zona'])
            row[1].text = f"{resultados['gdf_completo'].iloc[i]['costo_costo_nitrogeno']:.2f}"
            row[2].text = f"{resultados['gdf_completo'].iloc[i]['costo_costo_fosforo']:.2f}"
            row[3].text = f"{resultados['gdf_completo'].iloc[i]['costo_costo_potasio']:.2f}"
            row[4].text = f"{resultados['gdf_completo'].iloc[i]['costo_costo_total']:.2f}"
        doc.add_paragraph()
        costo_total = resultados['gdf_completo']['costo_costo_total'].sum()
        costo_promedio = resultados['gdf_completo']['costo_costo_total'].mean()
        doc.add_paragraph(f'**Costo total estimado:** ${costo_total:.2f} USD')
        doc.add_paragraph(f'**Costo promedio por hectárea:** ${costo_promedio:.2f} USD/ha')
        doc.add_paragraph()

        # ===== 5. TEXTURA DEL SUELO =====
        doc.add_heading('5. TEXTURA DEL SUELO', level=1)
        text_table = doc.add_table(rows=1, cols=5)
        text_table.style = 'Table Grid'
        text_headers = ['Zona', 'Textura', 'Arena (%)', 'Limo (%)', 'Arcilla (%)']
        for i, header in enumerate(text_headers): text_table.cell(0, i).text = header
        for i in range(min(10, len(resultados['gdf_completo']))):
            row = text_table.add_row().cells
            row[0].text = str(resultados['gdf_completo'].iloc[i]['id_zona'])
            row[1].text = str(resultados['gdf_completo'].iloc[i]['textura_suelo'])
            row[2].text = f"{resultados['gdf_completo'].iloc[i]['arena']:.1f}"
            row[3].text = f"{resultados['gdf_completo'].iloc[i]['limo']:.1f}"
            row[4].text = f"{resultados['gdf_completo'].iloc[i]['arcilla']:.1f}"
        doc.add_paragraph()

        # ===== 6. PROYECCIONES DE COSECHA =====
        doc.add_heading('6. PROYECCIONES DE COSECHA', level=1)
        proy_table = doc.add_table(rows=1, cols=4)
        proy_table.style = 'Table Grid'
        proy_headers = ['Zona', 'Sin Fertilización', 'Con Fertilización', 'Incremento (%)']
        for i, header in enumerate(proy_headers): proy_table.cell(0, i).text = header
        for i in range(min(10, len(resultados['gdf_completo']))):
            row = proy_table.add_row().cells
            row[0].text = str(resultados['gdf_completo'].iloc[i]['id_zona'])
            row[1].text = f"{resultados['gdf_completo'].iloc[i]['proy_rendimiento_sin_fert']:.0f}"
            row[2].text = f"{resultados['gdf_completo'].iloc[i]['proy_rendimiento_con_fert']:.0f}"
            row[3].text = f"{resultados['gdf_completo'].iloc[i]['proy_incremento_esperado']:.1f}"
        doc.add_paragraph()
        rend_sin_total = resultados['gdf_completo']['proy_rendimiento_sin_fert'].sum()
        rend_con_total = resultados['gdf_completo']['proy_rendimiento_con_fert'].sum()
        incremento_prom = resultados['gdf_completo']['proy_incremento_esperado'].mean()
        doc.add_paragraph(f'**Rendimiento total sin fertilización:** {rend_sin_total:.0f} kg')
        doc.add_paragraph(f'**Rendimiento total con fertilización:** {rend_con_total:.0f} kg')
        doc.add_paragraph(f'**Incremento promedio esperado:** {incremento_prom:.1f}%')
        doc.add_paragraph()

        # ===== 7. TOPOGRAFÍA =====
        doc.add_heading('7. TOPOGRAFÍA', level=1)
        if 'dem_data' in resultados and resultados['dem_data']:
            dem = resultados['dem_data']
            doc.add_paragraph(f"**Fuente DEM:** {dem.get('fuente', 'N/A')}")
            if dem['Z'] is not None:
                doc.add_paragraph(f"Elevación: min {np.nanmin(dem['Z']):.1f} m, max {np.nanmax(dem['Z']):.1f} m, prom {np.nanmean(dem['Z']):.1f} m")
            if dem.get('pendientes') is not None:
                doc.add_paragraph(f"Pendiente promedio: {np.nanmean(dem['pendientes']):.1f}%")
        doc.add_paragraph()

        # ===== 8. METADATOS TÉCNICOS =====
        doc.add_heading('8. METADATOS TÉCNICOS', level=1)
        metadatos = [
            ('Generado por', 'Analizador Multi-Cultivo Satelital v6.1'),
            ('Fecha de generación', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ('Sistema de coordenadas', 'EPSG:4326 (WGS84)'),
            ('Número de zonas', str(len(resultados['gdf_completo']))),
            ('Resolución satelital', SATELITES_DISPONIBLES[satelite]['resolucion']),
            ('Resolución DEM', f'{resolucion_dem} m'),
            ('Intervalo curvas de nivel', f'{intervalo_curvas} m')
        ]
        for key, value in metadatos:
            p = doc.add_paragraph()
            run_key = p.add_run(f'{key}: '); run_key.bold = True
            p.add_run(value)

        # Guardar en memoria
        docx_output = io.BytesIO()
        doc.save(docx_output)
        docx_output.seek(0)
        return docx_output

    except Exception as e:
        st.error(f"❌ Error generando reporte: {str(e)}")
        return None

def generar_reporte_con_ia(resultados, cultivo, satelite, fecha_inicio, fecha_fin,
                           resolucion_dem, intervalo_curvas):
    """
    Genera un informe DOCX con análisis de IA (DeepSeek) e incluye mapas y gráficos.
    """
    try:
        from modules.ia_integration import (
            preparar_resumen_zonas,
            generar_analisis_fertilidad,
            generar_analisis_riesgo_hidrico,
            generar_recomendaciones_integradas
        )
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_SECTION
        import numpy as np
        from datetime import datetime
        import io
        import tempfile
        import os

        doc = Document()
        
        # ===== PORTADA =====
        title = doc.add_heading(f'REPORTE DE AMBIENTACIÓN AGRONÓMICA CON IA - {cultivo}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # ===== 1. INFORMACIÓN GENERAL =====
        doc.add_heading('1. INFORMACIÓN GENERAL', level=1)
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = 'Cultivo'; info_table.cell(0, 1).text = cultivo
        info_table.cell(1, 0).text = 'Área Total'; info_table.cell(1, 1).text = f'{resultados["area_total"]:.2f} ha'
        info_table.cell(2, 0).text = 'Zonas Analizadas'; info_table.cell(2, 1).text = str(len(resultados['gdf_completo']))
        info_table.cell(3, 0).text = 'Satélite'; info_table.cell(3, 1).text = satelite
        info_table.cell(4, 0).text = 'Período'; info_table.cell(4, 1).text = f'{fecha_inicio.strftime("%d/%m/%Y")} a {fecha_fin.strftime("%d/%m/%Y")}'
        info_table.cell(5, 0).text = 'Fuente Datos'; info_table.cell(5, 1).text = resultados['datos_satelitales']['fuente'] if resultados['datos_satelitales'] else 'N/A'

        # ===== PREPARAR DATOS PARA IA =====
        df_resumen, stats = preparar_resumen_zonas(resultados['gdf_completo'], cultivo)

        # Crear un directorio temporal para guardar las imágenes
        with tempfile.TemporaryDirectory() as tmpdir:
            # ===== 2. ANÁLISIS DE FERTILIDAD =====
            doc.add_heading('2. ANÁLISIS DE FERTILIDAD', level=1)
            doc.add_paragraph('**Resumen de parámetros de fertilidad por zona:**')
            fert_table = doc.add_table(rows=1, cols=7)
            fert_table.style = 'Table Grid'
            headers = ['Zona', 'Área (ha)', 'Índice NPK', 'NDVI', 'NDRE', 'Materia Org (%)', 'Humedad']
            for i, header in enumerate(headers): fert_table.cell(0, i).text = header
            for i in range(min(10, len(resultados['gdf_completo']))):
                row = fert_table.add_row().cells
                row[0].text = str(resultados['gdf_completo'].iloc[i]['id_zona'])
                row[1].text = f"{resultados['gdf_completo'].iloc[i]['area_ha']:.2f}"
                row[2].text = f"{resultados['gdf_completo'].iloc[i]['fert_npk_actual']:.3f}"
                row[3].text = f"{resultados['gdf_completo'].iloc[i]['fert_ndvi']:.3f}"
                row[4].text = f"{resultados['gdf_completo'].iloc[i]['fert_ndre']:.3f}"
                row[5].text = f"{resultados['gdf_completo'].iloc[i]['fert_materia_organica']:.1f}"
                row[6].text = f"{resultados['gdf_completo'].iloc[i]['fert_humedad_suelo']:.3f}"
            doc.add_paragraph()

            # Insertar mapa de fertilidad
            mapa_fert = crear_mapa_fertilidad(resultados['gdf_completo'], cultivo, satelite)
            if mapa_fert:
                img_path = os.path.join(tmpdir, 'mapa_fertilidad.png')
                with open(img_path, 'wb') as f:
                    f.write(mapa_fert.getvalue())
                doc.add_picture(img_path, width=Inches(6))
                doc.add_paragraph("Figura 1: Mapa de Índice de Fertilidad NPK", style='Caption')

            # Análisis interpretativo generado por IA
            doc.add_heading('2.1 Interpretación', level=2)
            analisis_fert = generar_analisis_fertilidad(df_resumen, stats, cultivo)
            doc.add_paragraph(analisis_fert)

            # ===== 3. RECOMENDACIONES NPK =====
            doc.add_heading('3. RECOMENDACIONES NPK', level=1)
            doc.add_paragraph('**Recomendaciones de fertilización por zona (kg/ha):**')
            npk_table = doc.add_table(rows=1, cols=4)
            npk_table.style = 'Table Grid'
            npk_headers = ['Zona', 'Nitrógeno (N)', 'Fósforo (P)', 'Potasio (K)']
            for i, header in enumerate(npk_headers): npk_table.cell(0, i).text = header
            for i in range(min(10, len(resultados['gdf_completo']))):
                row = npk_table.add_row().cells
                row[0].text = str(resultados['gdf_completo'].iloc[i]['id_zona'])
                row[1].text = f"{resultados['gdf_completo'].iloc[i]['rec_N']:.1f}"
                row[2].text = f"{resultados['gdf_completo'].iloc[i]['rec_P']:.1f}"
                row[3].text = f"{resultados['gdf_completo'].iloc[i]['rec_K']:.1f}"
            doc.add_paragraph()

            # Mapas de NPK (podemos insertar los tres o elegir uno)
            doc.add_heading('3.1 Mapas de recomendaciones', level=2)
            # Mapa de Nitrógeno
            mapa_n = crear_mapa_npk(resultados['gdf_completo'], cultivo, 'N')
            if mapa_n:
                img_path = os.path.join(tmpdir, 'mapa_n.png')
                with open(img_path, 'wb') as f:
                    f.write(mapa_n.getvalue())
                doc.add_picture(img_path, width=Inches(5))
                doc.add_paragraph("Figura 2: Recomendación de Nitrógeno (kg/ha)", style='Caption')
            # Mapa de Fósforo
            mapa_p = crear_mapa_npk(resultados['gdf_completo'], cultivo, 'P')
            if mapa_p:
                img_path = os.path.join(tmpdir, 'mapa_p.png')
                with open(img_path, 'wb') as f:
                    f.write(mapa_p.getvalue())
                doc.add_picture(img_path, width=Inches(5))
                doc.add_paragraph("Figura 3: Recomendación de Fósforo (kg/ha)", style='Caption')
            # Mapa de Potasio
            mapa_k = crear_mapa_npk(resultados['gdf_completo'], cultivo, 'K')
            if mapa_k:
                img_path = os.path.join(tmpdir, 'mapa_k.png')
                with open(img_path, 'wb') as f:
                    f.write(mapa_k.getvalue())
                doc.add_picture(img_path, width=Inches(5))
                doc.add_paragraph("Figura 4: Recomendación de Potasio (kg/ha)", style='Caption')

            # ===== 4. RIESGO HÍDRICO Y TOPOGRAFÍA =====
            doc.add_heading('4. RIESGO DE ENCHARCAMIENTO', level=1)
            if 'dem_data' in resultados and resultados['dem_data']:
                dem = resultados['dem_data']
                doc.add_paragraph(f"**Fuente DEM:** {dem.get('fuente', 'N/A')}")
                doc.add_paragraph(f"Elevación: min {np.nanmin(dem['Z']):.1f} m, max {np.nanmax(dem['Z']):.1f} m, prom {np.nanmean(dem['Z']):.1f} m")
                if dem.get('pendientes') is not None:
                    doc.add_paragraph(f"Pendiente promedio: {np.nanmean(dem['pendientes']):.1f}%")
                # Insertar mapa de pendientes o curvas de nivel
                if dem.get('pendientes') is not None:
                    # Generar mapa de pendientes (necesitamos una función que lo cree como imagen)
                    # Podemos usar la misma lógica que en la pestaña de topografía
                    fig, ax = plt.subplots(1, 1, figsize=(10, 8))
                    bounds = dem['bounds']
                    minx, miny, maxx, maxy = bounds
                    pendientes = dem['pendientes']
                    pendientes_plot = np.ma.masked_invalid(pendientes)
                    im = ax.imshow(pendientes_plot, extent=[minx, maxx, miny, maxy],
                                   origin='lower', cmap='RdYlGn_r', alpha=0.8,
                                   aspect='auto', vmin=0, vmax=30)
                    plt.colorbar(im, ax=ax, label='Pendiente (%)')
                    resultados['gdf_completo'].plot(ax=ax, color='none', edgecolor='black', linewidth=2)
                    ax.set_title('Mapa de Pendientes')
                    ax.set_xlabel('Longitud'); ax.set_ylabel('Latitud')
                    img_path = os.path.join(tmpdir, 'pendientes.png')
                    plt.savefig(img_path, dpi=150, bbox_inches='tight')
                    plt.close()
                    doc.add_picture(img_path, width=Inches(6))
                    doc.add_paragraph("Figura 5: Mapa de Pendientes", style='Caption')
            
            # Análisis de riesgo hídrico por IA
            analisis_agua = generar_analisis_riesgo_hidrico(df_resumen, stats, cultivo)
            doc.add_heading('4.1 Análisis de humedad y textura', level=2)
            doc.add_paragraph(analisis_agua)

            # ===== 5. COSTOS =====
            doc.add_heading('5. ANÁLISIS DE COSTOS', level=1)
            costo_table = doc.add_table(rows=1, cols=5)
            costo_table.style = 'Table Grid'
            costo_headers = ['Zona', 'Costo N', 'Costo P', 'Costo K', 'Costo Total']
            for i, header in enumerate(costo_headers): costo_table.cell(0, i).text = header
            for i in range(min(10, len(resultados['gdf_completo']))):
                row = costo_table.add_row().cells
                row[0].text = str(resultados['gdf_completo'].iloc[i]['id_zona'])
                row[1].text = f"{resultados['gdf_completo'].iloc[i]['costo_costo_nitrogeno']:.2f}"
                row[2].text = f"{resultados['gdf_completo'].iloc[i]['costo_costo_fosforo']:.2f}"
                row[3].text = f"{resultados['gdf_completo'].iloc[i]['costo_costo_potasio']:.2f}"
                row[4].text = f"{resultados['gdf_completo'].iloc[i]['costo_costo_total']:.2f}"
            doc.add_paragraph()
            costo_total = resultados['gdf_completo']['costo_costo_total'].sum()
            costo_promedio = resultados['gdf_completo']['costo_costo_total'].mean()
            doc.add_paragraph(f'**Costo total estimado:** ${costo_total:.2f} USD')
            doc.add_paragraph(f'**Costo promedio por hectárea:** ${costo_promedio:.2f} USD/ha')
            doc.add_paragraph()

            # Gráfico de costos
            costos_n = resultados['gdf_completo']['costo_costo_nitrogeno'].sum()
            costos_p = resultados['gdf_completo']['costo_costo_fosforo'].sum()
            costos_k = resultados['gdf_completo']['costo_costo_potasio'].sum()
            otros = costo_total - (costos_n + costos_p + costos_k)
            grafico_costos = crear_grafico_distribucion_costos(costos_n, costos_p, costos_k, otros, costo_total)
            if grafico_costos:
                img_path = os.path.join(tmpdir, 'costos.png')
                with open(img_path, 'wb') as f:
                    f.write(grafico_costos.getvalue())
                doc.add_picture(img_path, width=Inches(6))
                doc.add_paragraph("Figura 6: Distribución de costos de fertilización", style='Caption')

            # ===== 6. TEXTURA DEL SUELO =====
            doc.add_heading('6. TEXTURA DEL SUELO', level=1)
            text_table = doc.add_table(rows=1, cols=5)
            text_table.style = 'Table Grid'
            text_headers = ['Zona', 'Textura', 'Arena (%)', 'Limo (%)', 'Arcilla (%)']
            for i, header in enumerate(text_headers): text_table.cell(0, i).text = header
            for i in range(min(10, len(resultados['gdf_completo']))):
                row = text_table.add_row().cells
                row[0].text = str(resultados['gdf_completo'].iloc[i]['id_zona'])
                row[1].text = str(resultados['gdf_completo'].iloc[i]['textura_suelo'])
                row[2].text = f"{resultados['gdf_completo'].iloc[i]['arena']:.1f}"
                row[3].text = f"{resultados['gdf_completo'].iloc[i]['limo']:.1f}"
                row[4].text = f"{resultados['gdf_completo'].iloc[i]['arcilla']:.1f}"
            doc.add_paragraph()

            # Mapa de texturas
            mapa_text = crear_mapa_texturas(resultados['gdf_completo'], cultivo)
            if mapa_text:
                img_path = os.path.join(tmpdir, 'texturas.png')
                with open(img_path, 'wb') as f:
                    f.write(mapa_text.getvalue())
                doc.add_picture(img_path, width=Inches(6))
                doc.add_paragraph("Figura 7: Mapa de texturas del suelo", style='Caption')

            # Gráfico de composición granulométrica
            arena_prom = resultados['gdf_completo']['arena'].mean()
            limo_prom = resultados['gdf_completo']['limo'].mean()
            arcilla_prom = resultados['gdf_completo']['arcilla'].mean()
            textura_dist = resultados['gdf_completo']['textura_suelo'].value_counts()
            grafico_textura = crear_grafico_composicion_textura(arena_prom, limo_prom, arcilla_prom, textura_dist)
            if grafico_textura:
                img_path = os.path.join(tmpdir, 'textura_pie.png')
                with open(img_path, 'wb') as f:
                    f.write(grafico_textura.getvalue())
                doc.add_picture(img_path, width=Inches(6))
                doc.add_paragraph("Figura 8: Composición granulométrica promedio", style='Caption')

            # ===== 7. PROYECCIONES DE COSECHA =====
            doc.add_heading('7. PROYECCIONES DE COSECHA', level=1)
            proy_table = doc.add_table(rows=1, cols=4)
            proy_table.style = 'Table Grid'
            proy_headers = ['Zona', 'Sin Fertilización', 'Con Fertilización', 'Incremento (%)']
            for i, header in enumerate(proy_headers): proy_table.cell(0, i).text = header
            for i in range(min(10, len(resultados['gdf_completo']))):
                row = proy_table.add_row().cells
                row[0].text = str(resultados['gdf_completo'].iloc[i]['id_zona'])
                row[1].text = f"{resultados['gdf_completo'].iloc[i]['proy_rendimiento_sin_fert']:.0f}"
                row[2].text = f"{resultados['gdf_completo'].iloc[i]['proy_rendimiento_con_fert']:.0f}"
                row[3].text = f"{resultados['gdf_completo'].iloc[i]['proy_incremento_esperado']:.1f}"
            doc.add_paragraph()
            rend_sin_total = resultados['gdf_completo']['proy_rendimiento_sin_fert'].sum()
            rend_con_total = resultados['gdf_completo']['proy_rendimiento_con_fert'].sum()
            incremento_prom = resultados['gdf_completo']['proy_incremento_esperado'].mean()
            doc.add_paragraph(f'**Rendimiento total sin fertilización:** {rend_sin_total:.0f} kg')
            doc.add_paragraph(f'**Rendimiento total con fertilización:** {rend_con_total:.0f} kg')
            doc.add_paragraph(f'**Incremento promedio esperado:** {incremento_prom:.1f}%')
            doc.add_paragraph()

            # Gráfico de proyecciones
            zonas_ids = resultados['gdf_completo']['id_zona'].astype(str).tolist()
            sin_fert = resultados['gdf_completo']['proy_rendimiento_sin_fert'].tolist()
            con_fert = resultados['gdf_completo']['proy_rendimiento_con_fert'].tolist()
            grafico_proy = crear_grafico_proyecciones_rendimiento(zonas_ids, sin_fert, con_fert)
            if grafico_proy:
                img_path = os.path.join(tmpdir, 'proyecciones.png')
                with open(img_path, 'wb') as f:
                    f.write(grafico_proy.getvalue())
                doc.add_picture(img_path, width=Inches(6))
                doc.add_paragraph("Figura 9: Proyecciones de rendimiento por zona", style='Caption')

            # ===== 8. POTENCIAL DE COSECHA =====
            doc.add_heading('8. POTENCIAL DE COSECHA', level=1)
            # Mapa de potencial base
            mapa_pot_base = crear_mapa_potencial_cosecha(resultados['gdf_completo'], cultivo)
            if mapa_pot_base:
                img_path = os.path.join(tmpdir, 'potencial_base.png')
                with open(img_path, 'wb') as f:
                    f.write(mapa_pot_base.getvalue())
                doc.add_picture(img_path, width=Inches(5))
                doc.add_paragraph("Figura 10: Potencial base de cosecha (sin fertilización)", style='Caption')
            # Mapa de potencial con recomendaciones
            mapa_pot_rec = crear_mapa_potencial_con_recomendaciones(resultados['gdf_completo'], cultivo)
            if mapa_pot_rec:
                img_path = os.path.join(tmpdir, 'potencial_rec.png')
                with open(img_path, 'wb') as f:
                    f.write(mapa_pot_rec.getvalue())
                doc.add_picture(img_path, width=Inches(5))
                doc.add_paragraph("Figura 11: Potencial con recomendaciones", style='Caption')
            # Gráfico comparativo
            grafico_comp = crear_grafico_comparativo_potencial(resultados['gdf_completo'], cultivo)
            if grafico_comp:
                img_path = os.path.join(tmpdir, 'comparativo_potencial.png')
                with open(img_path, 'wb') as f:
                    f.write(grafico_comp.getvalue())
                doc.add_picture(img_path, width=Inches(6))
                doc.add_paragraph("Figura 12: Comparativa de potencial", style='Caption')

            # ===== 9. RECOMENDACIONES INTEGRADAS (IA) =====
            doc.add_heading('9. RECOMENDACIONES DE MANEJO', level=1)
            recomendaciones_ia = generar_recomendaciones_integradas(df_resumen, stats, cultivo)
            doc.add_paragraph(recomendaciones_ia)

            # ===== 10. METADATOS TÉCNICOS =====
            doc.add_heading('10. METADATOS TÉCNICOS', level=1)
            metadatos = [
                ('Generado por', 'Analizador Multi-Cultivo Satelital v6.1 con IA DeepSeek'),
                ('Fecha de generación', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                ('Sistema de coordenadas', 'EPSG:4326 (WGS84)'),
                ('Número de zonas', str(len(resultados['gdf_completo']))),
                ('Resolución satelital', SATELITES_DISPONIBLES[satelite]['resolucion']),
                ('Resolución DEM', f'{resolucion_dem} m'),
                ('Intervalo curvas de nivel', f'{intervalo_curvas} m')
            ]
            for key, value in metadatos:
                p = doc.add_paragraph()
                run_key = p.add_run(f'{key}: '); run_key.bold = True
                p.add_run(value)

        # Guardar en memoria
        docx_output = io.BytesIO()
        doc.save(docx_output)
        docx_output.seek(0)
        return docx_output

    except Exception as e:
        st.error(f"❌ Error generando reporte con IA: {str(e)}")
        import traceback
        traceback.print_exc()
        return None
# ===== INTERFAZ PRINCIPAL =====
st.title("ANALIZADOR MULTI-CULTIVO SATELITAL")

if uploaded_file:
    with st.spinner("Cargando parcela..."):
        try:
            gdf = cargar_archivo_parcela(uploaded_file)
            if gdf is not None:
                st.success(f"✅ Parcela cargada exitosamente: {len(gdf)} polígono(s)")
                area_total = calcular_superficie(gdf)
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**📊 INFORMACIÓN DE LA PARCELA:**")
                    st.write(f"- Polígonos: {len(gdf)}")
                    st.write(f"- Área total: {area_total:.1f} ha")
                    st.write(f"- CRS: {gdf.crs}")
                    st.write(f"- Formato: {uploaded_file.name.split('.')[-1].upper()}")
                    fig, ax = plt.subplots(figsize=(8, 6))
                    gdf.plot(ax=ax, color='lightgreen', edgecolor='darkgreen', alpha=0.7)
                    ax.set_title(f"Parcela: {uploaded_file.name}")
                    ax.set_xlabel("Longitud"); ax.set_ylabel("Latitud"); ax.grid(True, alpha=0.3)
                    st.pyplot(fig)
                    buf_vista = io.BytesIO()
                    plt.savefig(buf_vista, format='png', dpi=150, bbox_inches='tight')
                    buf_vista.seek(0)
                    crear_boton_descarga_tiff(
                        buf_vista, gdf, f"vista_previa_{cultivo}",
                        "📥 Descargar Vista Previa TIFF", cultivo
                    )
                with col2:
                    st.write("**🎯 CONFIGURACIÓN**")
                    st.write(f"- Cultivo: {ICONOS_CULTIVOS[cultivo]} {cultivo}")
                    st.write(f"- Variedad: {variedad}")
                    st.write(f"- Zonas: {n_divisiones}")
                    st.write(f"- Satélite: {SATELITES_DISPONIBLES[satelite_seleccionado]['nombre']}")
                    st.write(f"- Período: {fecha_inicio} a {fecha_fin}")
                    st.write(f"- Intervalo curvas: {intervalo_curvas} m")
                    st.write(f"- Resolución DEM: {resolucion_dem} m")
                    if satelite_seleccionado in ['SENTINEL-2_GEE', 'LANDSAT-8_GEE', 'LANDSAT-9_GEE']:
                        if st.session_state.gee_authenticated:
                            st.success("✅ GEE autenticado")
                        else:
                            st.error("❌ GEE no autenticado - usando datos simulados")
                
                if st.button("🚀 EJECUTAR ANÁLISIS COMPLETO", type="primary", use_container_width=True):
                    with st.spinner("Ejecutando análisis completo..."):
                        resultados = ejecutar_analisis_completo(
                            gdf, cultivo, n_divisiones, 
                            satelite_seleccionado, fecha_inicio, fecha_fin,
                            intervalo_curvas, resolucion_dem
                        )
                        if resultados['exitoso']:
                            st.session_state.resultados_todos = resultados
                            st.session_state.analisis_completado = True
                            st.success("✅ Análisis completado exitosamente!")
                            st.rerun()
                        else:
                            st.error("❌ Error en el análisis completo")
            else:
                st.error("❌ Error al cargar la parcela. Verifica el formato del archivo.")
        except Exception as e:
            st.error(f"❌ Error en el análisis: {str(e)}")
            import traceback
            traceback.print_exc()

if st.session_state.analisis_completado and 'resultados_todos' in st.session_state:
    resultados = st.session_state.resultados_todos

    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs([
        "📊 Fertilidad Actual",
        "🧪 Recomendaciones NPK",
        "💰 Análisis de Costos",
        "🏗️ Textura del Suelo",
        "📈 Proyecciones",
        "🎯 Potencial de Cosecha",
        "🏔️ Curvas de Nivel y 3D",
        "🌍 Visualización NDVI+NDRE",
        "🦠 Detección YOLO"
    ])

    with tab1:
        st.subheader("FERTILIDAD ACTUAL")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            npk_prom = resultados['gdf_completo']['fert_npk_actual'].mean()
            st.metric("Índice NPK Promedio", f"{npk_prom:.3f}")
        with col2:
            ndvi_prom = resultados['gdf_completo']['fert_ndvi'].mean()
            st.metric("NDVI Promedio", f"{ndvi_prom:.3f}")
        with col3:
            mo_prom = resultados['gdf_completo']['fert_materia_organica'].mean()
            st.metric("Materia Orgánica", f"{mo_prom:.1f}%")
        with col4:
            hum_prom = resultados['gdf_completo']['fert_humedad_suelo'].mean()
            st.metric("Humedad Suelo", f"{hum_prom:.3f}")
        st.subheader("🗺️ MAPA DE FERTILIDAD")
        mapa_fert = crear_mapa_fertilidad(resultados['gdf_completo'], cultivo, satelite_seleccionado)
        if mapa_fert:
            st.image(mapa_fert, use_container_width=True)
            crear_boton_descarga_tiff(
                mapa_fert, resultados['gdf_completo'],
                f"mapa_fertilidad_{cultivo}",
                "📥 Descargar Mapa de Fertilidad TIFF", cultivo
            )
        st.subheader("📋 TABLA DE RESULTADOS")
        columnas_fert = ['id_zona', 'area_ha', 'fert_npk_actual', 'fert_ndvi',
                         'fert_ndre', 'fert_materia_organica', 'fert_humedad_suelo']
        tabla_fert = resultados['gdf_completo'][columnas_fert].copy()
        tabla_fert.columns = ['Zona', 'Área (ha)', 'Índice NPK', 'NDVI',
                              'NDRE', 'Materia Org (%)', 'Humedad']
        st.dataframe(tabla_fert)

    with tab2:
        st.subheader("RECOMENDACIONES NPK")
        col1, col2, col3 = st.columns(3)
        with col1:
            n_prom = resultados['gdf_completo']['rec_N'].mean()
            st.metric("Nitrógeno Promedio", f"{n_prom:.1f} kg/ha")
        with col2:
            p_prom = resultados['gdf_completo']['rec_P'].mean()
            st.metric("Fósforo Promedio", f"{p_prom:.1f} kg/ha")
        with col3:
            k_prom = resultados['gdf_completo']['rec_K'].mean()
            st.metric("Potasio Promedio", f"{k_prom:.1f} kg/ha")
        st.subheader("🗺️ MAPAS DE RECOMENDACIONES")
        col_n, col_p, col_k = st.columns(3)
        with col_n:
            mapa_n = crear_mapa_npk(resultados['gdf_completo'], cultivo, 'N')
            if mapa_n:
                st.image(mapa_n, use_container_width=True)
                st.caption("Nitrógeno (N)")
                crear_boton_descarga_tiff(
                    mapa_n, resultados['gdf_completo'],
                    f"mapa_nitrogeno_{cultivo}",
                    "📥 Descargar Mapa N TIFF", cultivo
                )
        with col_p:
            mapa_p = crear_mapa_npk(resultados['gdf_completo'], cultivo, 'P')
            if mapa_p:
                st.image(mapa_p, use_container_width=True)
                st.caption("Fósforo (P)")
                crear_boton_descarga_tiff(
                    mapa_p, resultados['gdf_completo'],
                    f"mapa_fosforo_{cultivo}",
                    "📥 Descargar Mapa P TIFF", cultivo
                )
        with col_k:
            mapa_k = crear_mapa_npk(resultados['gdf_completo'], cultivo, 'K')
            if mapa_k:
                st.image(mapa_k, use_container_width=True)
                st.caption("Potasio (K)")
                crear_boton_descarga_tiff(
                    mapa_k, resultados['gdf_completo'],
                    f"mapa_potasio_{cultivo}",
                    "📥 Descargar Mapa K TIFF", cultivo
                )
        st.subheader("📋 TABLA DE RECOMENDACIONES")
        columnas_npk = ['id_zona', 'area_ha', 'rec_N', 'rec_P', 'rec_K']
        tabla_npk = resultados['gdf_completo'][columnas_npk].copy()
        tabla_npk.columns = ['Zona', 'Área (ha)', 'Nitrógeno (kg/ha)',
                             'Fósforo (kg/ha)', 'Potasio (kg/ha)']
        st.dataframe(tabla_npk)

    with tab3:
        st.subheader("ANÁLISIS DE COSTOS")
        costo_total = resultados['gdf_completo']['costo_costo_total'].sum()
        costo_prom = resultados['gdf_completo']['costo_costo_total'].mean()
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Costo Total Estimado", f"${costo_total:.2f} USD")
        with col2:
            st.metric("Costo Promedio por ha", f"${costo_prom:.2f} USD/ha")
        with col3:
            inversion_ha = costo_total / resultados['area_total'] if resultados['area_total'] > 0 else 0
            st.metric("Inversión por ha", f"${inversion_ha:.2f} USD/ha")
        st.subheader("📊 DISTRIBUCIÓN DE COSTOS")
        costos_n = resultados['gdf_completo']['costo_costo_nitrogeno'].sum()
        costos_p = resultados['gdf_completo']['costo_costo_fosforo'].sum()
        costos_k = resultados['gdf_completo']['costo_costo_potasio'].sum()
        otros = costo_total - (costos_n + costos_p + costos_k)
        grafico_costos = crear_grafico_distribucion_costos(costos_n, costos_p, costos_k, otros, costo_total)
        if grafico_costos:
            st.image(grafico_costos, use_container_width=True)
            st.download_button(
                label="📥 Descargar Gráfico de Costos PNG",
                data=grafico_costos,
                file_name=f"grafico_costos_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                mime="image/png"
            )
        st.subheader("📋 TABLA DE COSTOS POR ZONA")
        columnas_costos = ['id_zona', 'area_ha', 'costo_costo_nitrogeno', 'costo_costo_fosforo',
                           'costo_costo_potasio', 'costo_costo_total']
        tabla_costos = resultados['gdf_completo'][columnas_costos].copy()
        tabla_costos.columns = ['Zona', 'Área (ha)', 'Costo N (USD)', 'Costo P (USD)',
                                'Costo K (USD)', 'Total (USD)']
        st.dataframe(tabla_costos)

    with tab4:
        st.subheader("TEXTURA DEL SUELO")
        textura_pred = resultados['gdf_completo']['textura_suelo'].mode()[0] if len(resultados['gdf_completo']) > 0 else "N/D"
        arena_prom = resultados['gdf_completo']['arena'].mean()
        limo_prom = resultados['gdf_completo']['limo'].mean()
        arcilla_prom = resultados['gdf_completo']['arcilla'].mean()
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Textura Predominante", textura_pred)
        with col2:
            st.metric("Arena Promedio", f"{arena_prom:.1f}%")
        with col3:
            st.metric("Limo Promedio", f"{limo_prom:.1f}%")
        with col4:
            st.metric("Arcilla Promedio", f"{arcilla_prom:.1f}%")
        st.subheader("🗺️ MAPA DE TEXTURAS")
        mapa_text = crear_mapa_texturas(resultados['gdf_completo'], cultivo)
        if mapa_text:
            st.image(mapa_text, use_container_width=True)
            crear_boton_descarga_tiff(
                mapa_text, resultados['gdf_completo'],
                f"mapa_texturas_{cultivo}",
                "📥 Descargar Mapa de Texturas TIFF", cultivo
            )
        st.subheader("📊 COMPOSICIÓN GRANULOMÉTRICA")
        textura_dist = resultados['gdf_completo']['textura_suelo'].value_counts()
        grafico_textura = crear_grafico_composicion_textura(arena_prom, limo_prom, arcilla_prom, textura_dist)
        if grafico_textura:
            st.image(grafico_textura, use_container_width=True)
            st.download_button(
                label="📥 Descargar Gráfico de Textura PNG",
                data=grafico_textura,
                file_name=f"grafico_textura_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                mime="image/png"
            )
        st.subheader("📋 TABLA DE TEXTURAS POR ZONA")
        columnas_text = ['id_zona', 'area_ha', 'textura_suelo', 'arena', 'limo', 'arcilla']
        tabla_text = resultados['gdf_completo'][columnas_text].copy()
        tabla_text.columns = ['Zona', 'Área (ha)', 'Textura', 'Arena (%)', 'Limo (%)', 'Arcilla (%)']
        st.dataframe(tabla_text)

    with tab5:
        st.subheader("PROYECCIONES DE COSECHA")
        rend_sin = resultados['gdf_completo']['proy_rendimiento_sin_fert'].sum()
        rend_con = resultados['gdf_completo']['proy_rendimiento_con_fert'].sum()
        incremento = ((rend_con - rend_sin) / rend_sin * 100) if rend_sin > 0 else 0
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Rendimiento sin Fertilización", f"{rend_sin:.0f} kg")
        with col2:
            st.metric("Rendimiento con Fertilización", f"{rend_con:.0f} kg")
        with col3:
            st.metric("Incremento Esperado", f"{incremento:.1f}%")
        st.subheader("📈 GRÁFICO DE PROYECCIONES")
        zonas_ids = resultados['gdf_completo']['id_zona'].astype(str).tolist()
        sin_fert = resultados['gdf_completo']['proy_rendimiento_sin_fert'].tolist()
        con_fert = resultados['gdf_completo']['proy_rendimiento_con_fert'].tolist()
        grafico_proyecciones = crear_grafico_proyecciones_rendimiento(zonas_ids, sin_fert, con_fert)
        if grafico_proyecciones:
            st.image(grafico_proyecciones, use_container_width=True)
            st.download_button(
                label="📥 Descargar Gráfico de Proyecciones PNG",
                data=grafico_proyecciones,
                file_name=f"grafico_proyecciones_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                mime="image/png"
            )
        st.subheader("📋 TABLA DE PROYECCIONES")
        columnas_proy = ['id_zona', 'area_ha', 'proy_rendimiento_sin_fert', 'proy_rendimiento_con_fert', 'proy_incremento_esperado']
        tabla_proy = resultados['gdf_completo'][columnas_proy].copy()
        tabla_proy.columns = ['Zona', 'Área (ha)', 'Sin Fertilización (kg)', 'Con Fertilización (kg)', 'Incremento (%)']
        st.dataframe(tabla_proy)

    with tab6:
        st.subheader("🎯 POTENCIAL DE COSECHA")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            potencial_prom = resultados['gdf_completo']['proy_rendimiento_sin_fert'].mean()
            st.metric("Potencial Promedio", f"{potencial_prom:.0f} kg/ha")
        with col2:
            potencial_max = resultados['gdf_completo']['proy_rendimiento_sin_fert'].max()
            st.metric("Potencial Máximo", f"{potencial_max:.0f} kg/ha")
        with col3:
            potencial_min = resultados['gdf_completo']['proy_rendimiento_sin_fert'].min()
            st.metric("Potencial Mínimo", f"{potencial_min:.0f} kg/ha")
        with col4:
            variabilidad = (potencial_max - potencial_min) / potencial_prom * 100
            st.metric("Variabilidad", f"{variabilidad:.1f}%")
        st.subheader("🗺️ MAPA DE POTENCIAL DE COSECHA")
        col_pot1, col_pot2 = st.columns(2)
        with col_pot1:
            mapa_potencial = crear_mapa_potencial_cosecha(resultados['gdf_completo'], cultivo)
            if mapa_potencial:
                st.image(mapa_potencial, use_container_width=True)
                st.caption("**Potencial Base (sin fertilización)**")
                crear_boton_descarga_tiff(
                    mapa_potencial, resultados['gdf_completo'],
                    f"mapa_potencial_base_{cultivo}",
                    "📥 Descargar Mapa Potencial Base TIFF", cultivo
                )
        with col_pot2:
            mapa_potencial_rec = crear_mapa_potencial_con_recomendaciones(resultados['gdf_completo'], cultivo)
            if mapa_potencial_rec:
                st.image(mapa_potencial_rec, use_container_width=True)
                st.caption("**Potencial con Recomendaciones**")
                crear_boton_descarga_tiff(
                    mapa_potencial_rec, resultados['gdf_completo'],
                    f"mapa_potencial_recomendaciones_{cultivo}",
                    "📥 Descargar Mapa Potencial con Recomendaciones TIFF", cultivo
                )
        st.subheader("📊 COMPARATIVA DE POTENCIAL")
        grafico_comparativo = crear_grafico_comparativo_potencial(resultados['gdf_completo'], cultivo)
        if grafico_comparativo:
            st.image(grafico_comparativo, use_container_width=True)
            st.download_button(
                label="📥 Descargar Gráfico Comparativo PNG",
                data=grafico_comparativo,
                file_name=f"grafico_comparativo_potencial_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                mime="image/png"
            )
        st.subheader("📋 ANÁLISIS POR ZONAS DE POTENCIAL")
        gdf_analisis = resultados['gdf_completo'].copy()
        gdf_analisis['potencial_categoria'] = pd.qcut(
            gdf_analisis['proy_rendimiento_sin_fert'],
            q=3,
            labels=['Bajo', 'Medio', 'Alto']
        )
        st.write("**Categorías de Potencial:**")
        categorias = gdf_analisis.groupby('potencial_categoria').agg({
            'id_zona': 'count',
            'area_ha': 'sum',
            'proy_rendimiento_sin_fert': ['mean', 'min', 'max'],
            'proy_incremento_esperado': 'mean'
        }).round(1)
        st.dataframe(categorias)
        st.subheader("🎯 RECOMENDACIONES POR CATEGORÍA DE POTENCIAL")
        col_rec1, col_rec2, col_rec3 = st.columns(3)
        with col_rec1:
            st.markdown("""
            **🔴 Zonas de POTENCIAL BAJO:**
            - Analizar causas: suelo compactado, drenaje, pH
            - Considerar enmiendas orgánicas
            - Evaluar cambio de cultivo/variedad
            - Priorizar en programas de mejora
            """)
        with col_rec2:
            st.markdown("""
            **🟡 Zonas de POTENCIAL MEDIO:**
            - Ajustar fertilización según análisis NPK
            - Mejorar prácticas de manejo
            - Implementar riego optimizado
            - Monitorear evolución temporal
            """)
        with col_rec3:
            st.markdown("""
            **🟢 Zonas de POTENCIAL ALTO:**
            - Mantener prácticas actuales
            - Optimizar cosecha y postcosecha
            - Considerar intensificación sostenible
            - Usar como referencia/control
            """)
        st.subheader("📋 TABLA DETALLADA DE POTENCIAL")
        columnas_potencial = [
            'id_zona', 'area_ha', 'potencial_categoria',
            'proy_rendimiento_sin_fert', 'proy_rendimiento_con_fert',
            'proy_incremento_esperado', 'fert_npk_actual'
        ]
        tabla_potencial = gdf_analisis[columnas_potencial].copy()
        tabla_potencial.columns = [
            'Zona', 'Área (ha)', 'Categoría',
            'Potencial Base (kg/ha)', 'Potencial Mejorado (kg/ha)',
            'Incremento (%)', 'Índice Fertilidad'
        ]
        st.dataframe(tabla_potencial.sort_values('Potencial Base (kg/ha)', ascending=False))

    with tab7:
        st.subheader("🏔️ ANÁLISIS TOPOGRÁFICO Y CURVAS DE NIVEL")

        # Verificar si existen datos topográficos válidos
        if ('dem_data' in resultados and resultados['dem_data'] and
            resultados['dem_data']['Z'] is not None and
            not np.all(np.isnan(resultados['dem_data']['Z']))):

            dem_data = resultados['dem_data']

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                elev_min = np.nanmin(dem_data['Z'])
                st.metric("Elevación Mínima", f"{elev_min:.1f} m" if not np.isnan(elev_min) else "N/A")
            with col2:
                elev_max = np.nanmax(dem_data['Z'])
                st.metric("Elevación Máxima", f"{elev_max:.1f} m" if not np.isnan(elev_max) else "N/A")
            with col3:
                elev_prom = np.nanmean(dem_data['Z'])
                st.metric("Elevación Promedio", f"{elev_prom:.1f} m" if not np.isnan(elev_prom) else "N/A")
            with col4:
                fuente = dem_data.get('fuente', 'Desconocida')
                st.metric("Fuente DEM", fuente)

            visualizacion = st.radio(
                "Tipo de visualización:",
                ["Mapa Interactivo (Folium)", "Mapa de Pendientes", "Curvas de Nivel (estático)", "Modelo 3D"],
                horizontal=True
            )

            if visualizacion == "Mapa Interactivo (Folium)":
                if FOLIUM_OK and dem_data.get('curvas_con_elevacion'):
                    st.subheader("🗺️ Mapa Interactivo de Curvas de Nivel")
                    m = mapa_curvas_coloreadas(resultados['gdf_completo'], dem_data['curvas_con_elevacion'])
                    if m:
                        if FOLIUM_STATIC_OK:
                            folium_static(m, width=1000, height=600)
                        else:
                            st.components.v1.html(m._repr_html_(), width=1000, height=600)
                    else:
                        st.error("No se pudo generar el mapa interactivo.")
                else:
                    if not FOLIUM_OK:
                        st.warning("⚠️ Folium no está instalado. No se puede mostrar el mapa interactivo.")
                    elif not dem_data.get('curvas_con_elevacion'):
                        st.warning("⚠️ No hay curvas de nivel generadas para esta área.")

            elif visualizacion == "Mapa de Pendientes":
                st.subheader("📉 MAPA DE PENDIENTES")
                if dem_data.get('pendientes') is not None:
                    fig, ax = plt.subplots(1, 1, figsize=(12, 8))

                    # Usar imshow para un mapa continuo de pendientes
                    bounds = dem_data['bounds']
                    minx, miny, maxx, maxy = bounds

                    # Crear una máscara para valores NaN
                    pendientes = dem_data['pendientes']
                    pendientes_plot = np.ma.masked_invalid(pendientes)

                    im = ax.imshow(pendientes_plot, extent=[minx, maxx, miny, maxy],
                                   origin='lower', cmap='RdYlGn_r', alpha=0.8,
                                   aspect='auto', vmin=0, vmax=30)
                    plt.colorbar(im, ax=ax, label='Pendiente (%)')

                    # Superponer el polígono de la parcela
                    resultados['gdf_completo'].plot(ax=ax, color='none', edgecolor='black', linewidth=2)

                    ax.set_title(f'Mapa de Pendientes - {fuente}')
                    ax.set_xlabel('Longitud'); ax.set_ylabel('Latitud')
                    st.pyplot(fig)
                    buf = io.BytesIO()
                    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
                    buf.seek(0)
                    st.download_button(
                        label="📥 Descargar Mapa de Pendientes PNG",
                        data=buf,
                        file_name=f"pendientes_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                        mime="image/png"
                    )
                else:
                    st.info("No hay datos de pendiente disponibles.")

            elif visualizacion == "Curvas de Nivel (estático)":
                st.subheader("⛰️ MAPA DE CURVAS DE NIVEL")
                if dem_data['Z'] is not None and not np.all(np.isnan(dem_data['Z'])):
                    fig, ax = plt.subplots(1, 1, figsize=(12, 8))

                    # Dibujar el fondo de elevación (siempre)
                    contourf = ax.contourf(dem_data['X'], dem_data['Y'], dem_data['Z'],
                                            levels=20, cmap='terrain', alpha=0.7)
                    plt.colorbar(contourf, ax=ax, label='Elevación (m)')

                    # Superponer curvas de nivel si existen
                    if dem_data.get('curvas_nivel') and len(dem_data['curvas_nivel']) > 0:
                        for line, elev in zip(dem_data['curvas_nivel'], dem_data['elevaciones']):
                            x, y = line.xy
                            ax.plot(x, y, 'b-', linewidth=0.8, alpha=0.7)
                            if len(x) > 0:
                                mid = len(x)//2
                                ax.text(x[mid], y[mid], f'{elev:.0f}', fontsize=7,
                                        bbox=dict(boxstyle="round,pad=0.2", fc='white', alpha=0.7))
                    else:
                        st.info("ℹ️ No se generaron curvas de nivel, solo se muestra el relieve.")

                    # Dibujar el contorno de la parcela
                    resultados['gdf_completo'].plot(ax=ax, color='none', edgecolor='black', linewidth=2)

                    ax.set_title(f'Curvas de Nivel - {dem_data.get("fuente", "Desconocida")}')
                    ax.set_xlabel('Longitud')
                    ax.set_ylabel('Latitud')
                    st.pyplot(fig)

                    # Botón de descarga
                    buf = io.BytesIO()
                    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
                    buf.seek(0)
                    st.download_button(
                        label="📥 Descargar Mapa de Curvas PNG",
                        data=buf,
                        file_name=f"curvas_nivel_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                        mime="image/png"
                    )
                else:
                    st.warning("No hay datos de elevación válidos para mostrar.")

            elif visualizacion == "Modelo 3D":
                st.subheader("🎨 VISUALIZACIÓN 3D DEL TERRENO")
                fig = plt.figure(figsize=(14, 10))
                ax = fig.add_subplot(111, projection='3d')
                step = max(1, dem_data['X'].shape[0] // 50)
                X_s = dem_data['X'][::step, ::step]
                Y_s = dem_data['Y'][::step, ::step]
                Z_s = dem_data['Z'][::step, ::step]
                surf = ax.plot_surface(X_s, Y_s, Z_s, cmap='terrain', alpha=0.8,
                                       linewidth=0, antialiased=True)
                ax.set_xlabel('Longitud'); ax.set_ylabel('Latitud'); ax.set_zlabel('Elevación (m)')
                ax.set_title(f'Modelo 3D del Terreno - {fuente}')
                fig.colorbar(surf, ax=ax, shrink=0.5, aspect=5, label='Elevación (m)')
                ax.view_init(elev=30, azim=45)
                st.pyplot(fig)

        else:
            st.info("ℹ️ No hay datos topográficos disponibles para esta parcela.")

    with tab8:
        st.subheader("🌱 VISUALIZACIÓN NDVI + NDRE")
        col_info1, col_info2 = st.columns(2)
        with col_info1:
            st.markdown("""
            ### 🌱 **NDVI**
            - **Fórmula:** (NIR - Rojo) / (NIR + Rojo)
            - **Rango:** -1.0 a 1.0
            - **Interpretación:**
              * < 0.1: Suelo desnudo/agua
              * 0.2-0.3: Vegetación escasa
              * 0.4-0.6: Vegetación moderada
              * > 0.7: Vegetación densa y saludable
            """)
        with col_info2:
            st.markdown("""
            ### 🌿 **NDRE**
            - **Fórmula:** (NIR - Borde Rojo) / (NIR + Borde Rojo)
            - **Rango:** -0.5 a 0.8
            - **Ventajas:**
              * Más sensible a clorofila en capas internas
              * Menos saturación en vegetación densa
              * Mejor para monitoreo de nitrógeno
            """)

        st.subheader("🛰️ Generar Mapas Estáticos")
        if satelite_seleccionado in ['SENTINEL-2_GEE', 'LANDSAT-8_GEE', 'LANDSAT-9_GEE']:
            if st.session_state.gee_authenticated:
                st.success(f"✅ Google Earth Engine autenticado - {SATELITES_DISPONIBLES[satelite_seleccionado]['nombre']}")
                if st.button("🔄 Generar Mapas NDVI + NDRE", type="primary", use_container_width=True):
                    with st.spinner("Descargando imágenes desde Google Earth Engine..."):
                        resultados_indices, mensaje = visualizar_indices_gee_estatico(
                            resultados['gdf_dividido'], satelite_seleccionado, fecha_inicio, fecha_fin
                        )
                    if resultados_indices:
                        st.session_state.indices_data = resultados_indices
                        st.session_state.indices_message = mensaje
                        st.success(mensaje)
                    else:
                        st.error(mensaje)

                if 'indices_data' in st.session_state:
                    indices_data = st.session_state.indices_data
                    st.subheader("🗺️ Mapas Generados")
                    col_map1, col_map2 = st.columns(2)
                    with col_map1:
                        st.image(indices_data['ndvi_bytes'], caption="Mapa NDVI", use_container_width=True)
                        ndvi_tiff_buffer, ndvi_tiff_filename = exportar_mapa_tiff(
                            indices_data['ndvi_bytes'], resultados['gdf_dividido'],
                            f"ndvi_{cultivo}", cultivo
                        )
                        if ndvi_tiff_buffer:
                            st.download_button(
                                label="📥 Descargar NDVI (TIFF)",
                                data=ndvi_tiff_buffer, file_name=ndvi_tiff_filename,
                                mime="image/tiff", use_container_width=True
                            )
                    with col_map2:
                        st.image(indices_data['ndre_bytes'], caption="Mapa NDRE", use_container_width=True)
                        ndre_tiff_buffer, ndre_tiff_filename = exportar_mapa_tiff(
                            indices_data['ndre_bytes'], resultados['gdf_dividido'],
                            f"ndre_{cultivo}", cultivo
                        )
                        if ndre_tiff_buffer:
                            st.download_button(
                                label="📥 Descargar NDRE (TIFF)",
                                data=ndre_tiff_buffer, file_name=ndre_tiff_filename,
                                mime="image/tiff", use_container_width=True
                            )
                    st.subheader("📊 Información Técnica")
                    info_col1, info_col2 = st.columns(2)
                    with info_col1:
                        fecha_str = datetime.fromtimestamp(indices_data['image_date']/1000).strftime('%Y-%m-%d') if indices_data['image_date'] else 'N/A'
                        st.markdown(f"""
                        **🌱 NDVI:**
                        - Fuente: {indices_data['title']}
                        - Fecha imagen: {fecha_str}
                        - Cobertura nubes: {indices_data['cloud_percent']}%
                        - ID: {indices_data['image_id']}
                        """)
                    with info_col2:
                        st.markdown("""
                        **🎯 Guía de Interpretación:**
                        - **NDVI > 0.7**: Vegetación muy densa y saludable
                        - **NDVI 0.4-0.7**: Vegetación en buen estado
                        - **NDVI 0.2-0.4**: Vegetación escasa o estresada
                        - **NDVI < 0.2**: Suelo desnudo o vegetación muy estresada
                        """)
                    st.subheader("📦 Descargar Todo")
                    zip_buffer = BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
                        if ndvi_tiff_buffer:
                            zip_file.writestr(
                                f"NDVI_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.tiff",
                                ndvi_tiff_buffer.getvalue()
                            )
                        if ndre_tiff_buffer:
                            zip_file.writestr(
                                f"NDRE_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.tiff",
                                ndre_tiff_buffer.getvalue()
                            )
                        bounds = resultados['gdf_dividido'].total_bounds
                        fecha_img = datetime.fromtimestamp(indices_data['image_date']/1000).strftime('%Y-%m-%d') if indices_data['image_date'] else 'N/A'
                        info_text = f"""INFORMACIÓN TÉCNICA - MAPAS NDVI + NDRE
Cultivo: {cultivo}
Satélite: {indices_data['title']}
Fecha generación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Fecha imagen: {fecha_img}
Cobertura nubes: {indices_data['cloud_percent']}%
ID Imagen: {indices_data['image_id']}
Coordenadas: [{bounds[0]:.6f}, {bounds[1]:.6f}, {bounds[2]:.6f}, {bounds[3]:.6f}]

ESCALAS DE COLOR:
- NDVI: -0.2 (rojo) a 0.8 (verde)
- NDRE: -0.1 (azul) a 0.6 (verde)

INTERPRETACIÓN:
- NDVI > 0.7: Vegetación muy densa
- NDVI 0.4-0.7: Vegetación saludable
- NDVI < 0.2: Posible estrés o suelo desnudo
- NDRE óptimo: 0.3-0.5
"""
                        zip_file.writestr(
                            f"INFO_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                            info_text
                        )
                    zip_buffer.seek(0)
                    st.download_button(
                        label="📥 Descargar Paquete Completo (ZIP)",
                        data=zip_buffer,
                        file_name=f"mapas_ndvi_ndre_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.zip",
                        mime="application/zip",
                        use_container_width=True
                    )
                else:
                    st.info("👆 Haz clic en 'Generar Mapas NDVI + NDRE' para crear las imágenes")
            else:
                st.error("❌ Google Earth Engine no está autenticado")
                st.info("Para generar mapas NDVI+NDRE desde GEE, configura el secret `GEE_SERVICE_ACCOUNT`")
        else:
            st.warning("⚠️ Para visualizaciones NDVI+NDRE, selecciona una fuente GEE")
            st.info("Fuentes GEE disponibles: SENTINEL-2_GEE, LANDSAT-8_GEE, LANDSAT-9_GEE")

        st.markdown("---")
        st.subheader("🗺️ Exportar GeoJSON de la Parcela")
        if st.button("📤 Generar GeoJSON de Parcela", use_container_width=True):
            with st.spinner("Generando GeoJSON..."):
                geojson_data, nombre_geojson = exportar_a_geojson(
                    resultados['gdf_completo'],
                    f"parcela_{cultivo}"
                )
                if geojson_data:
                    st.session_state.geojson_data = geojson_data
                    st.session_state.nombre_geojson = nombre_geojson
                    st.success(f"✅ GeoJSON generado: {nombre_geojson}")
                    st.rerun()
        if 'geojson_data' in st.session_state and st.session_state.geojson_data:
            col_geo1, col_geo2 = st.columns(2)
            with col_geo1:
                st.download_button(
                    label="📥 Descargar GeoJSON",
                    data=st.session_state.geojson_data,
                    file_name=st.session_state.nombre_geojson,
                    mime="application/json",
                    use_container_width=True
                )
            with col_geo2:
                if st.button("👁️ Previsualizar GeoJSON", use_container_width=True):
                    try:
                        geojson_dict = json.loads(st.session_state.geojson_data)
                        st.json(geojson_dict, expanded=False)
                    except:
                        st.warning("No se pudo mostrar la previsualización")

    with tab9:
        st.subheader("🦠 DETECCIÓN DE PLAGAS/ENFERMEDADES CON YOLO")
        col_yolo1, col_yolo2 = st.columns([2, 1])
        with col_yolo1:
            fuente_imagen = st.radio(
                "Fuente de imagen para análisis:",
                ["Subir imagen de campo", "Generar imagen simulada", "Usar imagen satelital GEE"],
                horizontal=True
            )
        with col_yolo2:
            confianza = st.slider("Confianza mínima", 0.3, 0.9, 0.5, 0.05)

        if 'modelo_yolo' not in st.session_state or st.session_state.modelo_yolo is None:
            with st.spinner("Cargando modelo YOLO..."):
                st.session_state.modelo_yolo = cargar_modelo_yolo()

        if fuente_imagen == "Subir imagen de campo":
            uploaded_image = st.file_uploader(
                "Sube imagen de campo/dron",
                type=['jpg', 'jpeg', 'png', 'bmp'],
                help="Imágenes de cultivo para detección de plagas"
            )
            if uploaded_image and st.button("🔍 Analizar con YOLO", type="primary"):
                with st.spinner("Procesando imagen con YOLO..."):
                    detecciones, imagen_resultado = detectar_plagas_yolo(
                        uploaded_image,
                        st.session_state.modelo_yolo,
                        confianza_minima=confianza
                    )
                    if imagen_resultado is not None:
                        col_res1, col_res2 = st.columns(2)
                        with col_res1:
                            st.subheader("📷 Imagen Analizada")
                            st.image(uploaded_image, caption="Imagen original", use_container_width=True)
                        with col_res2:
                            st.subheader("🎯 Detecciones YOLO")
                            st.image(imagen_resultado, caption="Detecciones", use_container_width=True)
                        if detecciones:
                            st.subheader("📊 Estadísticas de Detección")
                            df_detecciones = pd.DataFrame(detecciones)
                            col_stats1, col_stats2, col_stats3 = st.columns(3)
                            with col_stats1:
                                st.metric("Total detecciones", len(detecciones))
                            with col_stats2:
                                clases_unicas = df_detecciones['clase'].nunique()
                                st.metric("Tipos encontrados", clases_unicas)
                            with col_stats3:
                                conf_prom = df_detecciones['confianza'].mean()
                                st.metric("Confianza promedio", f"{conf_prom:.2f}")
                            st.dataframe(df_detecciones)
                            reporte = generar_reporte_plagas(detecciones, cultivo)
                            st.markdown(reporte)
                        else:
                            st.info("ℹ️ No se detectaron plagas/enfermedades con la confianza seleccionada")
        elif fuente_imagen == "Generar imagen simulada":
            if st.button("🔄 Generar y Analizar Imagen Simulada", type="primary"):
                with st.spinner("Generando imagen de simulación..."):
                    imagen_simulada = analizar_imagen_dron(
                        resultados['gdf_dividido'],
                        fecha_fin
                    )
                    if imagen_simulada:
                        detecciones, imagen_resultado = detectar_plagas_yolo(
                            imagen_simulada,
                            st.session_state.modelo_yolo,
                            confianza_minima=confianza
                        )
                        if imagen_resultado is not None:
                            col_res1, col_res2 = st.columns(2)
                            with col_res1:
                                st.subheader("📷 Imagen Simulada")
                                st.image(imagen_simulada, caption="Imagen simulada", use_container_width=True)
                            with col_res2:
                                st.subheader("🎯 Detecciones YOLO")
                                st.image(imagen_resultado, caption="Detecciones", use_container_width=True)
                            if detecciones:
                                st.subheader("📊 Estadísticas de Detección")
                                df_detecciones = pd.DataFrame(detecciones)
                                col_stats1, col_stats2, col_stats3 = st.columns(3)
                                with col_stats1:
                                    st.metric("Total detecciones", len(detecciones))
                                with col_stats2:
                                    clases_unicas = df_detecciones['clase'].nunique()
                                    st.metric("Tipos encontrados", clases_unicas)
                                with col_stats3:
                                    conf_prom = df_detecciones['confianza'].mean()
                                    st.metric("Confianza promedio", f"{conf_prom:.2f}")
                                st.dataframe(df_detecciones)
                                reporte = generar_reporte_plagas(detecciones, cultivo)
                                st.markdown(reporte)
        elif fuente_imagen == "Usar imagen satelital GEE":
            if st.session_state.gee_authenticated:
                if st.button("📡 Descargar y Analizar Imagen GEE", type="primary"):
                    st.info("🛠️ Funcionalidad en desarrollo - Usando simulación")
                    imagen_simulada = analizar_imagen_dron(
                        resultados['gdf_dividido'],
                        fecha_fin
                    )
                    if imagen_simulada:
                        detecciones, imagen_resultado = detectar_plagas_yolo(
                            imagen_simulada,
                            st.session_state.modelo_yolo,
                            confianza_minima=confianza
                        )
                        if imagen_resultado:
                            st.image(imagen_resultado, caption="Detecciones (simuladas)", use_container_width=True)
            else:
                st.warning("⚠️ Necesitas autenticación GEE para esta función")

    st.markdown("---")
    st.subheader("💾 EXPORTAR RESULTADOS")
    col_exp1, col_exp2, col_exp3 = st.columns(3)
    with col_exp1:
        st.markdown("**GeoJSON**")
        if st.button("📤 Generar GeoJSON", key="generate_geojson"):
            with st.spinner("Generando GeoJSON..."):
                geojson_data, nombre_geojson = exportar_a_geojson(
                    resultados['gdf_completo'],
                    f"analisis_{cultivo}"
                )
                if geojson_data:
                    st.session_state.geojson_data = geojson_data
                    st.session_state.nombre_geojson = nombre_geojson
                    st.success("✅ GeoJSON generado correctamente")
                    st.rerun()
        if 'geojson_data' in st.session_state and st.session_state.geojson_data:
            st.download_button(
                label="📥 Descargar GeoJSON",
                data=st.session_state.geojson_data,
                file_name=st.session_state.nombre_geojson,
                mime="application/json",
                key="geojson_download"
            )
    with col_exp2:
        st.markdown("**Reporte DOCX**")
        # Botón para reporte estándar
        if st.button("📄 Generar Reporte Completo", key="generate_report"):
            with st.spinner("Generando reporte DOCX..."):
                reporte = generar_reporte_completo(
                    resultados,
                    cultivo,
                    satelite_seleccionado,
                    fecha_inicio,
                    fecha_fin,
                    resolucion_dem,
                    intervalo_curvas
                )
                if reporte:
                    st.session_state.reporte_completo = reporte
                    st.session_state.nombre_reporte = f"reporte_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                    st.success("✅ Reporte generado correctamente")
                    st.rerun()
        if 'reporte_completo' in st.session_state and st.session_state.reporte_completo:
            st.download_button(
                label="📥 Descargar Reporte DOCX",
                data=st.session_state.reporte_completo,
                file_name=st.session_state.nombre_reporte,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="report_download"
            )
        # Botón para reporte con IA
        st.markdown("---")
        if st.button("🤖 Generar Reporte con IA", key="generate_ia_report"):
            with st.spinner("Generando informe con análisis de IA..."):
                reporte_ia = generar_reporte_con_ia(
                    resultados, cultivo, satelite_seleccionado, fecha_inicio, fecha_fin,
                    resolucion_dem, intervalo_curvas
                )
                if reporte_ia:
                    st.session_state.reporte_ia = reporte_ia
                    st.session_state.nombre_reporte_ia = f"reporte_ia_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                    st.success("✅ Reporte con IA generado")
                    st.rerun()
        if 'reporte_ia' in st.session_state and st.session_state.reporte_ia:
            st.download_button(
                label="📥 Descargar Reporte con IA",
                data=st.session_state.reporte_ia,
                file_name=st.session_state.nombre_reporte_ia,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="ia_report_download"
            )
    with col_exp3:
        st.markdown("**Limpiar Resultados**")
        if st.button("🗑️ Limpiar Resultados", use_container_width=True):
            for key in list(st.session_state.keys()):
                if key not in ['gee_authenticated', 'gee_project']:
                    del st.session_state[key]
            st.rerun()

st.markdown("---")
col_footer1, col_footer2, col_footer3 = st.columns(3)
with col_footer1:
    st.markdown("""
    📡 **Fuentes de Datos:**  
    NASA POWER API  
    Google Earth Engine  
    Sentinel-2 (ESA)  
    Landsat-8/9 (USGS)  
    SRTM 30m (OpenTopography)  
    Open Topo Data API  
    Datos simulados
    """)
with col_footer2:
    st.markdown("""
    🛠️ **Tecnologías:**  
    Streamlit  
    GeoPandas  
    Google Earth Engine API  
    Matplotlib  
    Rasterio / scikit-image  
    Folium / Branca  
    Python-DOCX
    """)
with col_footer3:
    st.markdown("""
    📞 **Soporte:**  
    Versión: 6.1 - Fuente alternativa DEM (Open Topo Data)  
    Última actualización: Febrero 2026  
    Martin Ernesto Cano  
    mawucano@gmail.com | +5493525 532313
    """)

st.markdown(
    '<div style="text-align: center; padding: 20px; margin-top: 20px; border-top: 1px solid #3b82f6;">'
    '<p style="color: #94a3b8; margin: 0;">© 2026 Analizador Multi-Cultivo Satelital. Todos los derechos reservados.</p>'
    '</div>',
    unsafe_allow_html=True
)
