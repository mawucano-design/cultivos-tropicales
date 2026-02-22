# modules/generar_reporte.py
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from .ia_integration import (
    generar_analisis_fertilidad,
    generar_analisis_riesgo_hidrico,
    generar_recomendaciones_integradas,
    preparar_resumen_zonas
)

def crear_docx_con_ia(output_path, resultados, cultivo, satelite, fecha_inicio, fecha_fin):
    """
    Versión mejorada que incluye análisis generado por IA.
    """
    doc = Document()
    
    # Título
    title = doc.add_heading(f'REPORTE DE AMBIENTACIÓN AGRONÓMICA - {cultivo}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    # ===== SECCIÓN 1: INTRODUCCIÓN (generada por IA) =====
    doc.add_heading('1. INTRODUCCIÓN', level=1)
    # Preparamos datos para IA
    df_resumen, stats = preparar_resumen_zonas(resultados['gdf_completo'], cultivo)
    
    # Podemos pedir a la IA que genere una introducción personalizada
    prompt_intro = f"""
    Redacta una introducción profesional para un informe agronómico de un lote de {cultivo} de {stats['area_total']:.2f} ha.
    Menciona que se utilizaron imágenes satelitales ({satelite}), modelo digital de elevación y análisis de suelo.
    El objetivo es caracterizar la heterogeneidad del lote y orientar prácticas de manejo específicas.
    """
    intro_text = llamar_deepseek(prompt_intro, temperatura=0.5, max_tokens=500)
    doc.add_paragraph(intro_text)
    
    # ===== SECCIÓN 2: ANÁLISIS DE FERTILIDAD =====
    doc.add_heading('2. ANÁLISIS DE FERTILIDAD', level=1)
    # Tabla de fertilidad (como ya tienes)
    # ...
    # Luego el análisis de IA
    analisis_fert = generar_analisis_fertilidad(df_resumen, stats, cultivo)
    doc.add_heading('2.1 Interpretación', level=2)
    doc.add_paragraph(analisis_fert)
    
    # ===== SECCIÓN 3: RIESGO HÍDRICO Y TOPOGRAFÍA =====
    doc.add_heading('3. RIESGO DE ENCHARCAMIENTO', level=1)
    # Si hay datos topográficos, mostrarlos...
    # ...
    analisis_agua = generar_analisis_riesgo_hidrico(df_resumen, stats, cultivo)
    doc.add_heading('3.1 Análisis de humedad y textura', level=2)
    doc.add_paragraph(analisis_agua)
    
    # ===== SECCIÓN 4: RECOMENDACIONES INTEGRADAS =====
    doc.add_heading('4. RECOMENDACIONES DE MANEJO', level=1)
    recomendaciones = generar_recomendaciones_integradas(df_resumen, stats, cultivo)
    doc.add_paragraph(recomendaciones)
    
    # ... resto del informe (costos, proyecciones, etc.) como antes ...
    
    doc.save(output_path)
    return doc
